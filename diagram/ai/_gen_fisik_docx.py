# /// script
# requires-python = ">=3.11"
# dependencies = ["python-docx>=1.1.2"]
# ///
"""Generate 4.3.3 Perancangan Fisik Basis Data .docx aligned to ERD (6 entitas).
Tables: tb_user, tb_klien, tb_proyek, tb_penugasan_proyek, tb_tugas, tb_bukti,
tb_eskalasi. Style: Times New Roman, 5-col schema tables."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Twips
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = r"D:\__CODING\05-MyProjects\__IRENE\shi-crm\diagram\ai\4.3.3_PERANCANGAN_FISIK.docx"
FONT = "Times New Roman"

# Column widths in twips (DXA). 5 columns: No | Nama Field | Tipe Data | Panjang | Constraint
COL_W = [600, 2400, 1800, 3360, 1500]

TABLES = [
    {
        "num": "4.6", "title": "User", "nama": "users", "pk": "id", "fk": [],
        "rows": [
            ("id", "SERIAL", "-", "PK"),
            ("name", "VARCHAR", "255", "NOT NULL"),
            ("email", "VARCHAR", "255", "Unique, NOT NULL"),
            ("role", "VARCHAR", "50", "NOT NULL, Default: technician"),
            ("password_hash", "VARCHAR", "255", "NOT NULL"),
            ("created_at", "TIMESTAMP", "-", "Default: CURRENT_TIMESTAMP"),
        ],
    },
    {
        "num": "4.7", "title": "Klien", "nama": "clients", "pk": "id",
        "fk": ["created_by (terhubung dengan users)"],
        "rows": [
            ("id", "SERIAL", "-", "PK"),
            ("name", "VARCHAR", "255", "NOT NULL"),
            ("address", "TEXT", "-", ""),
            ("phone", "VARCHAR", "50", ""),
            ("email", "VARCHAR", "255", ""),
            ("notes", "TEXT", "-", ""),
            ("latitude", "DECIMAL", "10,7", ""),
            ("longitude", "DECIMAL", "10,7", ""),
            ("photo_path", "VARCHAR", "1000", ""),
            ("photo_name", "VARCHAR", "500", ""),
            ("created_by", "INT", "-", "FK, NOT NULL"),
            ("created_at", "TIMESTAMP", "-", "Default: CURRENT_TIMESTAMP"),
            ("updated_at", "TIMESTAMP", "-", "Default: CURRENT_TIMESTAMP"),
        ],
    },
    {
        "num": "4.8", "title": "Proyek", "nama": "projects", "pk": "id",
        "fk": [
            "client_id (terhubung dengan clients)",
            "created_by (terhubung dengan users)",
            "survey_approved_by (terhubung dengan users)",
        ],
        "rows": [
            ("id", "SERIAL", "-", "PK"),
            ("project_code", "VARCHAR", "12", "Unique, NOT NULL"),
            ("name", "VARCHAR", "255", "NOT NULL"),
            ("description", "TEXT", "-", ""),
            ("client_id", "INT", "-", "FK"),
            ("start_date", "DATE", "-", "NOT NULL"),
            ("end_date", "DATE", "-", "NOT NULL"),
            ("duration", "INT", "-", "Generated"),
            ("status", "VARCHAR", "50", "Default: active"),
            ("phase", "VARCHAR", "50", "Default: survey"),
            ("category", "VARCHAR", "50", "Default: instalasi"),
            ("project_value", "DECIMAL", "15,2", "Default: 0"),
            ("survey_approved", "BOOLEAN", "-", "Default: FALSE"),
            ("survey_approved_by", "INT", "-", "FK"),
            ("survey_approved_at", "TIMESTAMP", "-", ""),
            ("target_description", "TEXT", "-", ""),
            ("created_by", "INT", "-", "FK, NOT NULL"),
            ("created_at", "TIMESTAMP", "-", "Default: CURRENT_TIMESTAMP"),
            ("updated_at", "TIMESTAMP", "-", "Default: CURRENT_TIMESTAMP"),
        ],
    },
    {
        "num": "4.9", "title": "Penugasan Proyek", "nama": "project_assignments",
        "pk": "(project_id, user_id)",
        "fk": [
            "project_id (terhubung dengan projects)",
            "user_id (terhubung dengan users)",
        ],
        "rows": [
            ("project_id", "INT", "-", "PK, FK"),
            ("user_id", "INT", "-", "PK, FK"),
            ("assigned_at", "TIMESTAMP", "-", "Default: CURRENT_TIMESTAMP"),
        ],
    },
    {
        "num": "4.10", "title": "Tugas", "nama": "tasks", "pk": "id",
        "fk": [
            "project_id (terhubung dengan projects)",
            "assigned_to (terhubung dengan users)",
            "depends_on (terhubung dengan tasks)",
            "created_by (terhubung dengan users)",
        ],
        "rows": [
            ("id", "SERIAL", "-", "PK"),
            ("project_id", "INT", "-", "FK, NOT NULL"),
            ("name", "VARCHAR", "500", "NOT NULL"),
            ("description", "TEXT", "-", ""),
            ("assigned_to", "INT", "-", "FK"),
            ("status", "VARCHAR", "50", "Default: to_do"),
            ("due_date", "DATE", "-", ""),
            ("timeline_start", "DATE", "-", ""),
            ("timeline_end", "DATE", "-", ""),
            ("notes", "TEXT", "-", ""),
            ("sort_order", "INT", "-", "Default: 0"),
            ("is_survey_task", "BOOLEAN", "-", "Default: FALSE"),
            ("timer_started_at", "TIMESTAMP", "-", ""),
            ("time_spent_seconds", "INT", "-", "Default: 0"),
            ("is_tracking", "BOOLEAN", "-", "Default: FALSE"),
            ("estimated_hours", "DECIMAL", "5,1", ""),
            ("depends_on", "INT", "-", "FK"),
            ("status_changed_at", "TIMESTAMP", "-", "Default: NOW()"),
            ("created_by", "INT", "-", "FK, NOT NULL"),
            ("created_at", "TIMESTAMP", "-", "Default: CURRENT_TIMESTAMP"),
            ("updated_at", "TIMESTAMP", "-", "Default: CURRENT_TIMESTAMP"),
        ],
    },
    {
        "num": "4.11", "title": "Bukti Tugas", "nama": "task_evidence", "pk": "id",
        "fk": [
            "task_id (terhubung dengan tasks)",
            "uploaded_by (terhubung dengan users)",
        ],
        "rows": [
            ("id", "SERIAL", "-", "PK"),
            ("task_id", "INT", "-", "FK, NOT NULL"),
            ("file_path", "VARCHAR", "1000", "NOT NULL"),
            ("file_name", "VARCHAR", "500", "NOT NULL"),
            ("file_type", "VARCHAR", "50", "NOT NULL"),
            ("file_size", "INT", "-", "Default: 0"),
            ("description", "TEXT", "-", ""),
            ("uploaded_by", "INT", "-", "FK, NOT NULL"),
            ("uploaded_at", "TIMESTAMP", "-", "Default: CURRENT_TIMESTAMP"),
        ],
    },
    {
        "num": "4.12", "title": "Daily Report", "nama": "daily_reports", "pk": "id",
        "fk": [
            "project_id (terhubung dengan projects)",
            "task_id (terhubung dengan tasks)",
            "created_by (terhubung dengan users)",
        ],
        "rows": [
            ("id", "SERIAL", "-", "PK"),
            ("project_id", "INT", "-", "FK, NOT NULL"),
            ("task_id", "INT", "-", "FK"),
            ("report_date", "DATE", "-", "NOT NULL"),
            ("progress_percentage", "DECIMAL", "5,2", "NOT NULL, Range: 0-100"),
            ("constraints", "TEXT", "-", ""),
            ("created_by", "INT", "-", "FK, NOT NULL"),
            ("created_at", "TIMESTAMP", "-", "Default: CURRENT_TIMESTAMP"),
        ],
    },
    {
        "num": "4.13", "title": "Kesehatan Proyek", "nama": "project_health",
        "pk": "project_id",
        "fk": ["project_id (terhubung dengan projects)"],
        "rows": [
            ("project_id", "INT", "-", "PK, FK"),
            ("spi_value", "DECIMAL", "6,4", ""),
            ("status", "VARCHAR", "50", "CHECK: green|amber|red"),
            ("actual_progress", "DECIMAL", "5,2", ""),
            ("planned_progress", "DECIMAL", "5,2", ""),
            ("total_tasks", "INT", "-", "Default: 0"),
            ("completed_tasks", "INT", "-", "Default: 0"),
            ("working_tasks", "INT", "-", "Default: 0"),
            ("overtime_tasks", "INT", "-", "Default: 0"),
            ("overdue_tasks", "INT", "-", "Default: 0"),
            ("last_updated", "TIMESTAMP", "-", "Default: CURRENT_TIMESTAMP"),
        ],
    },
    {
        "num": "4.14", "title": "Aktivitas Tugas", "nama": "task_activities",
        "pk": "id",
        "fk": [
            "task_id (terhubung dengan tasks)",
            "user_id (terhubung dengan users)",
        ],
        "rows": [
            ("id", "SERIAL", "-", "PK"),
            ("task_id", "INT", "-", "FK, NOT NULL"),
            ("user_id", "INT", "-", "FK, NOT NULL"),
            ("message", "TEXT", "-", "NOT NULL"),
            ("activity_type", "VARCHAR", "50", "Default: note"),
            ("file_path", "VARCHAR", "1000", ""),
            ("file_name", "VARCHAR", "500", ""),
            ("file_type", "VARCHAR", "50", ""),
            ("file_size", "INT", "-", "Default: 0"),
            ("created_at", "TIMESTAMP", "-", "Default: CURRENT_TIMESTAMP"),
        ],
    },
    {
        "num": "4.15", "title": "Eskalasi", "nama": "escalations", "pk": "id",
        "fk": [
            "task_id (terhubung dengan tasks)",
            "project_id (terhubung dengan projects)",
            "reported_by (terhubung dengan users)",
            "resolved_by (terhubung dengan users)",
        ],
        "rows": [
            ("id", "SERIAL", "-", "PK"),
            ("task_id", "INT", "-", "FK, NOT NULL"),
            ("project_id", "INT", "-", "FK, NOT NULL"),
            ("reported_by", "INT", "-", "FK, NOT NULL"),
            ("title", "VARCHAR", "500", "NOT NULL"),
            ("description", "TEXT", "-", "NOT NULL"),
            ("status", "VARCHAR", "50", "Default: open"),
            ("priority", "VARCHAR", "50", "Default: medium"),
            ("file_path", "VARCHAR", "1000", ""),
            ("file_name", "VARCHAR", "500", ""),
            ("file_type", "VARCHAR", "50", ""),
            ("file_size", "INT", "-", "Default: 0"),
            ("resolved_by", "INT", "-", "FK"),
            ("resolved_at", "TIMESTAMP", "-", ""),
            ("resolution_notes", "TEXT", "-", ""),
            ("created_at", "TIMESTAMP", "-", "Default: CURRENT_TIMESTAMP"),
            ("updated_at", "TIMESTAMP", "-", "Default: CURRENT_TIMESTAMP"),
        ],
    },
    {
        "num": "4.16", "title": "Audit Log", "nama": "audit_log", "pk": "id",
        "fk": ["changed_by (terhubung dengan users)"],
        "rows": [
            ("id", "SERIAL", "-", "PK"),
            ("entity_type", "VARCHAR", "50", "NOT NULL"),
            ("entity_id", "INT", "-", "NOT NULL"),
            ("entity_name", "VARCHAR", "255", ""),
            ("action", "VARCHAR", "50", "NOT NULL"),
            ("field_name", "VARCHAR", "100", ""),
            ("old_value", "TEXT", "-", ""),
            ("new_value", "TEXT", "-", ""),
            ("changed_by", "INT", "-", "FK, NOT NULL"),
            ("changed_by_name", "VARCHAR", "255", ""),
            ("created_at", "TIMESTAMP", "-", "Default: CURRENT_TIMESTAMP"),
        ],
    },
]

INTRO = (
    "Perancangan model fisik basis data merupakan tahap akhir dalam proses desain "
    "basis data. Pada tahap ini, skema logis atau model Entity Relationship Diagram "
    "(ERD) yang telah dirancang sebelumnya diwujudkan menjadi basis data nyata "
    "menggunakan perangkat lunak manajemen basis data (DBMS) yang dipilih. Tujuan "
    "utama dari perancangan fisik ini adalah untuk mencapai efisiensi dalam "
    "pemrosesan dan pengelolaan data. Berikut adalah perancangan model fisik basis "
    "data untuk sistem manajemen proyek di PT Smart Home Inovasi (SHI)."
)


def set_cell_shading(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "000000")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def style_run(run, *, bold: bool = False, size: int = 11, italic: bool = False) -> None:
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        rpr.append(r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        r_fonts.set(qn(attr), FONT)


def write_cell(cell, text: str, *, bold: bool = False, center: bool = False,
               header: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    style_run(run, bold=bold or header, size=11)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_borders(cell)
    if header:
        set_cell_shading(cell, "D9E2F3")


def add_para(doc, text: str, *, bold: bool = False, italic: bool = False,
             center: bool = False, justify: bool = False, indent_first: float = 0.0,
             indent_left: float = 0.0, size: int = 11, space_before: int = 0,
             space_after: int = 6, heading_level: int | None = None):
    p = doc.add_paragraph()
    if heading_level is not None:
        p.style = doc.styles[f"Heading {heading_level}"]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    if indent_first:
        pf.first_line_indent = Cm(indent_first)
    if indent_left:
        pf.left_indent = Cm(indent_left)
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    style_run(run, bold=bold, italic=italic, size=size)
    return p


def set_col_widths(table, widths_twips: list[int]) -> None:
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Twips(widths_twips[idx])


def build_table(doc, t: dict) -> None:
    rows = t["rows"]
    table = doc.add_table(rows=len(rows) + 1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    headers = ("No", "Nama Field", "Tipe Data", "Panjang Karakter", "Constraint")
    for i, h in enumerate(headers):
        write_cell(table.rows[0].cells[i], h, header=True, center=True)

    for r_idx, (field, dtype, length, constraint) in enumerate(rows, start=1):
        write_cell(table.rows[r_idx].cells[0], str(r_idx), center=True)
        write_cell(table.rows[r_idx].cells[1], field)
        write_cell(table.rows[r_idx].cells[2], dtype)
        write_cell(table.rows[r_idx].cells[3], length)
        write_cell(table.rows[r_idx].cells[4], constraint, center=True)

    set_col_widths(table, COL_W)


def main() -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        rpr.append(r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        r_fonts.set(qn(attr), FONT)

    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    add_para(doc, "4.3.3  Perancangan Fisik Basis Data", bold=True, size=12,
             space_before=12, space_after=12)

    add_para(doc, INTRO, justify=True, indent_first=1.27, size=12, space_after=12)

    for idx, t in enumerate(TABLES, start=1):
        add_para(
            doc, f"{idx}.  Tabel {t['title']}", bold=True, size=12,
            indent_left=0.75, space_before=10, space_after=4,
        )
        add_para(doc, f"Nama Tabel: {t['nama']}", indent_left=1.5, size=12,
                 space_after=2)
        add_para(doc, f"Primary Key: {t['pk']}", indent_left=1.5, size=12,
                 space_after=2)
        for fk in t["fk"]:
            add_para(doc, f"Foreign Key: {fk}", indent_left=1.5, size=12,
                     space_after=2)
        add_para(doc, f"Tabel {t['num']} {t['title']}", italic=True, center=True,
                 size=11, space_before=6, space_after=4)
        build_table(doc, t)
        add_para(doc, "", size=11, space_after=0)

    doc.save(OUT)
    print(f"OK {OUT}")


if __name__ == "__main__":
    main()
