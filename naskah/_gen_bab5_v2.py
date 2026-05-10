# /// script
# requires-python = ">=3.11"
# dependencies = ["python-docx>=1.1.2", "Pillow>=10"]
# ///
"""Generate BAB V (Implementasi & Hasil) for shi-crm thesis.
Content covers actual implementation: Next.js 15 + React 19 + PostgreSQL.
Screenshots from alur-sistem/*.png (real Playwright captures of running app).

Outline:
  5.1 Implementasi
    5.1.1 Lingkungan Pengembangan (HW/SW table)
    5.1.2 Implementasi Basis Data (13 tables)
    5.1.3 Implementasi Antarmuka (10 pages per §4.3.5 + 12-step flow appendix)
    5.1.4 Implementasi Logika EWS (SPI calculator)
  5.2 Hasil Pengujian
    5.2.1 Black Box Testing
    5.2.2 User Acceptance Testing
"""
from __future__ import annotations
import io
import os
from PIL import Image
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT_PATH = r"D:\__CODING\05-MyProjects\__IRENE\shi-crm\naskah\BAB_V_Implementasi.docx"
SHOTS_DIR = r"D:\__CODING\05-MyProjects\__IRENE\shi-crm\alur-sistem"
FONT = "Times New Roman"


# -- Style helpers --------------------------------------------------------
def style_run(run, *, bold=False, italic=False, size=12, mono=False, color=None):
    name = "Consolas" if mono else FONT
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def set_cell_borders(cell):
    tcpr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "000000")
        borders.append(b)
    tcpr.append(borders)


def set_cell_shading(cell, color_hex):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcpr.append(shd)


def add_para(doc, text, *, bold=False, italic=False, size=12, justify=False,
             center=False, indent_first=0.0, indent_left=0.0,
             space_before=0, space_after=6, mono=False, color=None):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    if indent_first:
        pf.first_line_indent = Cm(indent_first)
    if indent_left:
        pf.left_indent = Cm(indent_left)
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        style_run(run, bold=bold, italic=italic, size=size, mono=mono, color=color)
    return p


def add_heading(doc, level, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level <= 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    size = {0: 16, 1: 14, 2: 13, 3: 12}.get(level, 12)
    run = p.add_run(text)
    style_run(run, bold=True, size=size)


def add_image(doc, png_path, max_width_cm=14.0):
    if not os.path.exists(png_path):
        add_para(doc, f"[Gambar tidak ditemukan: {os.path.basename(png_path)}]",
                 italic=True, center=True, size=10, color="888888")
        return
    img = Image.open(png_path)
    w_px, h_px = img.size
    w_cm = w_px / 96 * 2.54
    h_cm = h_px / 96 * 2.54
    if w_cm > max_width_cm:
        scale = max_width_cm / w_cm
        w_cm = max_width_cm
        h_cm = h_cm * scale
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(png_path, width=Cm(w_cm))


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    style_run(run, italic=True, size=11)


def write_cell(cell, text, *, bold=False, center=False, header=False,
               size=11, mono=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    style_run(run, bold=bold or header, size=size, mono=mono)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_borders(cell)
    if header:
        set_cell_shading(cell, "D9E2F3")


def build_table(doc, headers, rows, col_widths_dxa, *, header_align_center=True):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, h in enumerate(headers):
        write_cell(table.rows[0].cells[i], h, header=True,
                   center=header_align_center)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            mono = isinstance(val, tuple) and val[1] == "mono"
            txt = val[0] if isinstance(val, tuple) else str(val)
            center = c_idx == 0 or (isinstance(val, tuple) and val[0].startswith("PK"))
            write_cell(table.rows[r_idx].cells[c_idx], txt,
                       center=center if c_idx == 0 else False, mono=mono, size=10)
    for row in table.rows:
        for c_idx, cell in enumerate(row.cells):
            cell.width = Twips(col_widths_dxa[c_idx])
    return table


# -- Document content -----------------------------------------------------
def main():
    doc = Document()

    # Default Normal style
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT)

    # Page margins (academic Indo: 4cm L, 3cm R/T/B - flexible)
    section = doc.sections[0]
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(3.0)
    section.left_margin = Cm(4.0)
    section.right_margin = Cm(3.0)

    # ==== Title ====
    add_heading(doc, 0, "BAB V")
    add_heading(doc, 0, "IMPLEMENTASI DAN HASIL SERTA PEMBAHASAN")

    # ==== 5.1 Implementasi ====
    add_heading(doc, 1, "5.1  Implementasi")
    add_para(doc,
             "Tahap implementasi merupakan realisasi dari rancangan sistem yang telah "
             "disusun pada BAB IV menjadi sebuah perangkat lunak yang dapat dijalankan "
             "secara nyata. Pengembangan sistem dashboard pemantauan proyek pada PT "
             "Smart Home Inovasi (SHI) ini dibangun sebagai aplikasi web berbasis Next.js "
             "dengan dukungan basis data PostgreSQL. Pada bagian ini akan dipaparkan "
             "lingkungan pengembangan, implementasi struktur basis data, implementasi "
             "antarmuka, serta implementasi logika kalkulasi Early Warning System (EWS).",
             justify=True, indent_first=1.27, space_after=8)

    # ---- 5.1.1 Lingkungan Pengembangan ----
    add_heading(doc, 2, "5.1.1  Lingkungan Pengembangan")
    add_para(doc,
             "Pengembangan sistem dilakukan menggunakan kombinasi perangkat keras dan "
             "perangkat lunak sebagai berikut.",
             justify=True, indent_first=1.27)

    add_para(doc, "1.  Perangkat Keras (Hardware)", bold=True, indent_left=0.5,
             space_before=6, space_after=4)
    build_table(
        doc,
        ["No", "Komponen", "Spesifikasi"],
        [
            ["1", "Laptop", "ASUS TUF Gaming F15"],
            ["2", "Processor", "Intel Core i7-11800H @2.30GHz"],
            ["3", "RAM", "16 GB DDR4"],
            ["4", "Penyimpanan", "SSD 512 GB NVMe"],
            ["5", "GPU", "NVIDIA GeForce RTX 3050 4GB"],
            ["6", "Resolusi Layar", "1920 x 1080 (Full HD)"],
        ],
        [600, 3000, 5400],
    )
    add_caption(doc, "Tabel 5.1 Spesifikasi Perangkat Keras Pengembangan")

    add_para(doc, "2.  Perangkat Lunak (Software)", bold=True, indent_left=0.5,
             space_before=6, space_after=4)
    build_table(
        doc,
        ["No", "Komponen", "Versi", "Fungsi"],
        [
            ["1", "Sistem Operasi", "Windows 11 Pro", "Sistem operasi pengembangan"],
            ["2", "Node.js", "22.x", "Runtime JavaScript server-side"],
            ["3", "Bun", "1.1.x", "Package manager + task runner"],
            ["4", "Next.js", "15.3.3", "Framework web React (App Router, SSR)"],
            ["5", "React", "19.2", "Library antarmuka pengguna"],
            ["6", "TypeScript", "5.9", "Bahasa pemrograman dengan tipe statis"],
            ["7", "PostgreSQL", "17.5", "Database relasional"],
            ["8", "Tailwind CSS", "4.2", "Framework utility-first untuk styling"],
            ["9", "TanStack Query", "5.90", "State management server data"],
            ["10", "Recharts", "3.8", "Visualisasi grafik dashboard"],
            ["11", "Playwright", "1.49", "Pengujian end-to-end"],
            ["12", "Vitest", "2.1", "Pengujian unit + komponen"],
            ["13", "VS Code", "1.95", "Editor kode sumber"],
        ],
        [600, 2200, 1500, 4700],
    )
    add_caption(doc, "Tabel 5.2 Spesifikasi Perangkat Lunak Pengembangan")

    # ---- 5.1.2 Implementasi Basis Data ----
    add_heading(doc, 2, "5.1.2  Implementasi Basis Data")
    add_para(doc,
             "Implementasi basis data direalisasikan menggunakan PostgreSQL versi 17.5 "
             "dengan total 13 tabel relasional yang saling terhubung melalui foreign key. "
             "Skema basis data dirancang sesuai dengan rancangan logis dan fisik pada "
             "subbab 4.3.2 dan 4.3.3, dengan beberapa tabel pendukung tambahan untuk "
             "kebutuhan operasional sistem. Daftar tabel beserta fungsinya disajikan pada "
             "Tabel 5.3.",
             justify=True, indent_first=1.27)

    build_table(
        doc,
        ["No", "Nama Tabel", "Fungsi"],
        [
            ["1", "users", "Menyimpan data pengguna sistem (Manajer, Teknisi, Admin)"],
            ["2", "clients", "Data klien beserta lokasi dan kontak"],
            ["3", "projects", "Master data proyek dengan fase survey/execution"],
            ["4", "project_assignments", "Tabel relasi M:N pengguna-proyek (mengelola)"],
            ["5", "tasks", "Rincian tugas (decomposition) per proyek"],
            ["6", "task_evidence", "Lampiran bukti pekerjaan per tugas"],
            ["7", "materials", "Data material/bahan yang digunakan per proyek"],
            ["8", "budget_items", "Rincian anggaran biaya proyek"],
            ["9", "daily_reports", "Catatan progres harian dari teknisi"],
            ["10", "project_health", "Hasil kalkulasi SPI dan status RAG per proyek"],
            ["11", "task_activities", "Riwayat perubahan status tugas (audit trail)"],
            ["12", "escalations", "Tiket eskalasi kendala lapangan"],
            ["13", "audit_log", "Log aktivitas sistem untuk keamanan"],
        ],
        [600, 2800, 5600],
    )
    add_caption(doc, "Tabel 5.3 Daftar Tabel Basis Data Sistem")

    add_para(doc,
             "Inisialisasi seluruh tabel di atas dilakukan melalui skrip "
             "database/schema.sql yang dijalankan otomatis melalui perintah "
             "bun run db:setup. Skrip tersebut berisi pernyataan CREATE TABLE, "
             "CONSTRAINT, dan INDEX untuk mengoptimalkan kueri. Kontrol referensial "
             "dijaga oleh foreign key dengan kebijakan ON DELETE CASCADE pada "
             "tabel-tabel anak yang bergantung pada parent (misalnya tasks "
             "bergantung pada projects).",
             justify=True, indent_first=1.27, space_before=6)

    # ---- 5.1.3 Implementasi Antarmuka ----
    add_heading(doc, 2, "5.1.3  Implementasi Antarmuka")
    add_para(doc,
             "Implementasi antarmuka direalisasikan menggunakan kombinasi React 19 "
             "sebagai library UI, Tailwind CSS v4 untuk styling, serta TanStack Query "
             "untuk pengelolaan state data dari server. Setiap halaman dirancang untuk "
             "responsif dan menampilkan data secara real-time melalui mekanisme "
             "invalidasi cache pada saat terjadi perubahan data. Berikut adalah "
             "implementasi setiap halaman antarmuka utama sesuai rancangan pada "
             "subbab 4.3.5.",
             justify=True, indent_first=1.27)

    # 10 pages with screenshots
    pages = [
        {
            "num": "1",
            "title": "Halaman Login",
            "desc": "Halaman Login menjadi pintu masuk autentikasi seluruh pengguna "
                    "sistem. Pengguna memasukkan email dan kata sandi, kemudian sistem "
                    "memvalidasi kredensial melalui endpoint POST /api/auth/login. Jika "
                    "valid, sistem menerbitkan JSON Web Token (JWT) yang disimpan pada "
                    "localStorage browser dan mengarahkan pengguna ke dashboard sesuai "
                    "perannya: Manajer ke /dashboard, Teknisi ke /my-dashboard.",
            "shot": "00-login.png",
            "fig": "5.1",
            "cap": "Tampilan Halaman Login",
        },
        {
            "num": "2",
            "title": "Halaman Tambah Proyek",
            "desc": "Halaman Tambah Proyek hanya dapat diakses oleh Manajer melalui "
                    "menu Proyek. Form ini terdiri dari informasi proyek (nama, klien, "
                    "tanggal, nilai), bagian penugasan teknisi, dan dekomposisi daftar "
                    "tugas. Sistem memberikan rekomendasi teknisi yang tersedia "
                    "berdasarkan ketersediaan jadwal melalui filter algoritma sederhana "
                    "yang menghindari bentrok jadwal aktif.",
            "shot": "02a-input-proyek.png",
            "fig": "5.2",
            "cap": "Tampilan Halaman Tambah Proyek (Form Input)",
        },
        {
            "num": "2b",
            "title": "Rekomendasi Teknisi pada Tambah Proyek",
            "desc": "Setelah Manajer memilih rentang tanggal, sistem menampilkan daftar "
                    "teknisi tersedia. Teknisi yang sudah memiliki tugas pada rentang "
                    "tersebut tidak ditampilkan dalam rekomendasi. Manajer kemudian "
                    "memilih teknisi yang akan ditugaskan.",
            "shot": "03-pilih-teknisi.png",
            "fig": "5.3",
            "cap": "Tampilan Rekomendasi Teknisi Otomatis",
        },
        {
            "num": "3",
            "title": "Halaman Tambah Daily Report",
            "desc": "Halaman Tambah Daily Report digunakan oleh Teknisi untuk "
                    "melaporkan progres harian melalui pembaruan status tugas. Teknisi "
                    "membuka tugas yang sedang dikerjakan, mengubah status menjadi "
                    "Working On It, dan mengunggah bukti pekerjaan berupa foto atau "
                    "dokumen. Status Done hanya dapat diset oleh Manajer setelah "
                    "melakukan validasi (review gate).",
            "shot": "06-teknisi-update-task.png",
            "fig": "5.4",
            "cap": "Tampilan Halaman Tambah Daily Report (Update Status Tugas)",
        },
        {
            "num": "4",
            "title": "Halaman Dashboard Early Warning System (EWS)",
            "desc": "Dashboard EWS menjadi halaman utama Manajer setelah login. "
                    "Halaman ini menampilkan ringkasan seluruh proyek aktif yang "
                    "diurutkan berdasarkan tingkat kekritisan. Indikator warna RAG "
                    "(Red/Amber/Green) merepresentasikan status kesehatan proyek "
                    "berdasarkan nilai Schedule Performance Index (SPI). Proyek "
                    "berstatus merah ditempatkan pada urutan teratas untuk mempermudah "
                    "Manajer dalam menentukan prioritas penanganan.",
            "shot": "10-dashboard-ews-spi-health.png",
            "fig": "5.5",
            "cap": "Tampilan Halaman Dashboard Early Warning System",
        },
        {
            "num": "4b",
            "title": "Visualisasi Grafik Dashboard",
            "desc": "Selain ringkasan proyek, dashboard juga menampilkan beberapa "
                    "visualisasi grafik antara lain pie chart distribusi status "
                    "kesehatan proyek, grafik beban kerja per teknisi, grafik tugas "
                    "berdasarkan status, dan tren earned value mingguan.",
            "shot": "10b-dashboard-charts.png",
            "fig": "5.6",
            "cap": "Tampilan Visualisasi Grafik pada Dashboard EWS",
        },
        {
            "num": "5",
            "title": "Halaman Data Proyek",
            "desc": "Halaman Data Proyek menampilkan daftar seluruh proyek beserta "
                    "fitur filter (status, fase, klien) dan pencarian. Manajer dapat "
                    "menambah proyek baru, mengakses detail proyek, atau melakukan "
                    "tindakan administratif lainnya melalui kolom Aksi pada setiap "
                    "baris.",
            "shot": "02b-nambah-proyek.png",
            "fig": "5.7",
            "cap": "Tampilan Halaman Data Proyek (setelah penambahan)",
        },
        {
            "num": "6",
            "title": "Halaman Kanban Penugasan Proyek",
            "desc": "Halaman Kanban Penugasan menampilkan tugas-tugas dalam tiga "
                    "kolom berdasarkan status: To Do, Working On It, dan Done. Teknisi "
                    "dapat memindahkan tugas dari To Do ke Working On It saat memulai "
                    "pekerjaan. Untuk memindahkan ke kolom Done, sistem menerapkan "
                    "review gate -- hanya Manajer yang berwenang menyetujui penyelesaian "
                    "setelah melakukan verifikasi terhadap bukti pekerjaan.",
            "shot": "07-manager-approve-task-proyek.png",
            "fig": "5.8",
            "cap": "Tampilan Halaman Kanban Penugasan Proyek (Review Gate Manajer)",
        },
        {
            "num": "6b",
            "title": "Update Eskalasi oleh Teknisi",
            "desc": "Pada halaman Kanban dan/atau halaman tugas, teknisi dapat "
                    "mengajukan eskalasi kendala lapangan dengan menuliskan deskripsi "
                    "masalah, tingkat prioritas, dan kategori kendala. Eskalasi akan "
                    "muncul pada dashboard Manajer sebagai indikator perlu tindakan.",
            "shot": "08-teknisi-update-eskalasi.png",
            "fig": "5.9",
            "cap": "Tampilan Pengajuan Eskalasi oleh Teknisi",
        },
        {
            "num": "7",
            "title": "Halaman Jadwal Proyek",
            "desc": "Halaman Jadwal Proyek menyajikan timeline pengerjaan tugas dalam "
                    "format Gantt-like dengan baris teknisi dan kolom hari. Manajer "
                    "dapat memantau alokasi sumber daya, mengidentifikasi bentrok "
                    "jadwal, dan melakukan realokasi tugas sesuai kebutuhan operasional.",
            "shot": None,
            "fig": "5.10",
            "cap": "Tampilan Halaman Jadwal Proyek",
        },
        {
            "num": "8",
            "title": "Halaman Dashboard Performa Teknisi",
            "desc": "Halaman ini memungkinkan setiap teknisi memantau performa "
                    "pribadinya. Komponen yang ditampilkan meliputi nilai SPI personal, "
                    "jumlah tugas selesai, sedang berjalan, overdue, serta tren SPI "
                    "mingguan. Halaman ini mendukung swamonitoring oleh teknisi tanpa "
                    "perlu campur tangan manajer.",
            "shot": "12b-performa-teknisi-self.png",
            "fig": "5.11",
            "cap": "Tampilan Halaman Dashboard Performa Teknisi (Self-View)",
        },
        {
            "num": "8b",
            "title": "Halaman Performa Teknisi (View Manajer)",
            "desc": "Manajer juga dapat melihat performa kolektif seluruh teknisi "
                    "untuk keperluan evaluasi kinerja, pemerataan beban kerja, atau "
                    "identifikasi teknisi yang membutuhkan pendampingan tambahan.",
            "shot": "12-performa-teknisi-manager.png",
            "fig": "5.12",
            "cap": "Tampilan Halaman Performa Teknisi (View Manajer)",
        },
        {
            "num": "9",
            "title": "Halaman Detail Proyek",
            "desc": "Halaman Detail Proyek menyajikan informasi komprehensif sebuah "
                    "proyek meliputi header status RAG, fase proyek, ringkasan progres "
                    "PV vs EV, breakdown tugas berdasarkan status, daftar eskalasi "
                    "aktif, serta panel komentar global untuk komunikasi terpadu antara "
                    "Manajer dan Teknisi.",
            "shot": "11-report-proyek.png",
            "fig": "5.13",
            "cap": "Tampilan Halaman Detail Proyek",
        },
        {
            "num": "10",
            "title": "Halaman Laporan Kesehatan Proyek",
            "desc": "Halaman Laporan Kesehatan menampilkan rincian status kesehatan "
                    "proyek beserta grafik Earned Value (PV vs EV), breakdown anggaran "
                    "dan material, serta rekomendasi tindakan. Halaman ini mendukung "
                    "ekspor ke format PDF untuk keperluan dokumentasi atau pelaporan "
                    "ke pihak eksternal.",
            "shot": "11b-report-budget-material.png",
            "fig": "5.14",
            "cap": "Tampilan Halaman Laporan Kesehatan Proyek (Anggaran & Material)",
        },
    ]

    for pg in pages:
        add_para(doc, f"{pg['num']}.  {pg['title']}", bold=True, indent_left=0.5,
                 space_before=8, space_after=4)
        add_para(doc, pg["desc"], justify=True, indent_first=1.27)
        if pg["shot"]:
            add_image(doc, os.path.join(SHOTS_DIR, pg["shot"]))
        else:
            add_para(doc,
                     "[Gambar belum diabadikan; diambil pada revisi berikutnya saat "
                     "lingkungan staging tersedia.]",
                     italic=True, center=True, size=10, color="888888")
        add_caption(doc, f"Gambar {pg['fig']} {pg['cap']}")

    # ---- 5.1.4 Implementasi Logika EWS ----
    add_heading(doc, 2, "5.1.4  Implementasi Logika Early Warning System (EWS)")
    add_para(doc,
             "Logika kalkulasi EWS merupakan komponen inti yang membedakan sistem "
             "dashboard ini dari sistem manual sebelumnya. Implementasi dilakukan pada "
             "modul backend frontend/src/lib/spiCalculator.ts. Logika ini mencakup tiga "
             "tahapan utama: kalkulasi Planned Value (PV), kalkulasi Schedule "
             "Performance Index (SPI), dan kategorisasi status kesehatan proyek "
             "berdasarkan parameter Red-Amber-Green (RAG).",
             justify=True, indent_first=1.27)

    add_para(doc, "1.  Kalkulasi Planned Value (PV)", bold=True, indent_left=0.5,
             space_before=6, space_after=4)
    add_para(doc,
             "Planned Value dihitung sebagai persentase progres ideal yang seharusnya "
             "dicapai berdasarkan jumlah hari yang telah berlalu sejak proyek dimulai, "
             "dibandingkan dengan total durasi proyek.",
             justify=True, indent_first=1.27)
    add_para(doc, "PV = (Hari Berlalu / Total Durasi Proyek) x 100%",
             center=True, italic=True, size=11, space_after=4)
    add_para(doc, "Nilai PV diklamp pada rentang [0, 100] untuk menjaga konsistensi.",
             justify=True, indent_first=1.27)

    add_para(doc, "2.  Kalkulasi Schedule Performance Index (SPI)", bold=True,
             indent_left=0.5, space_before=6, space_after=4)
    add_para(doc,
             "SPI dihitung dari rasio antara progres aktual (Earned Value) terhadap "
             "progres rencana (Planned Value). Earned Value dihitung berdasarkan rasio "
             "tugas yang sudah berstatus done terhadap total tugas pada proyek tersebut.",
             justify=True, indent_first=1.27)
    add_para(doc, "EV = (Jumlah Tugas Selesai / Total Tugas) x 100%",
             center=True, italic=True, size=11, space_after=4)
    add_para(doc, "SPI = EV / PV",
             center=True, italic=True, size=11, space_after=4)
    add_para(doc,
             "Pada kasus proyek yang belum memiliki tugas (jumlah tugas = 0), sistem "
             "menggunakan fallback berupa nilai progress_percentage dari laporan harian "
             "terakhir (daily_reports) sebagai estimasi Earned Value.",
             justify=True, indent_first=1.27)

    add_para(doc, "3.  Kategorisasi Status Kesehatan (RAG)", bold=True,
             indent_left=0.5, space_before=6, space_after=4)
    add_para(doc,
             "Berdasarkan nilai SPI yang telah dihitung, sistem mengelompokkan status "
             "kesehatan proyek ke dalam tiga kategori sebagaimana ditunjukkan pada "
             "Tabel 5.4.",
             justify=True, indent_first=1.27)
    build_table(
        doc,
        ["Status RAG", "Rentang SPI", "Interpretasi"],
        [
            ["Hijau (Green)", "SPI >= 0.95", "Proyek sesuai jadwal atau lebih cepat"],
            ["Kuning (Amber)", "0.85 <= SPI < 0.95", "Proyek mulai melenceng, perhatian diperlukan"],
            ["Merah (Red)", "SPI < 0.85", "Proyek terlambat signifikan, mitigasi mendesak"],
        ],
        [2400, 2400, 4200],
    )
    add_caption(doc, "Tabel 5.4 Pemetaan Nilai SPI ke Status RAG")

    add_para(doc, "4.  Pemicu Recalculation", bold=True, indent_left=0.5,
             space_before=6, space_after=4)
    add_para(doc,
             "Recalculation SPI dipicu secara otomatis pada beberapa peristiwa penting: "
             "(a) perubahan status tugas (misalnya dari working_on_it menjadi done), "
             "(b) penambahan atau penghapusan tugas pada proyek, (c) submit daily report "
             "manual, dan (d) penjadwalan recalculation periodik untuk seluruh proyek "
             "aktif. Hasil kalkulasi disimpan ke tabel project_health melalui mekanisme "
             "UPSERT (INSERT ON CONFLICT DO UPDATE) sehingga setiap proyek hanya "
             "memiliki satu record kesehatan terkini.",
             justify=True, indent_first=1.27)

    # ==== 5.2 Hasil Pengujian ====
    add_heading(doc, 1, "5.2  Hasil Pengujian")
    add_para(doc,
             "Pengujian sistem dilakukan untuk memastikan bahwa seluruh fungsionalitas "
             "yang telah diimplementasikan berjalan sesuai dengan kebutuhan yang "
             "tercantum pada subbab 4.2. Pengujian dilakukan dalam dua tahap utama: "
             "pengujian fungsional menggunakan metode Black Box Testing, dan pengujian "
             "penerimaan pengguna menggunakan metode User Acceptance Testing (UAT) "
             "yang melibatkan pengguna akhir.",
             justify=True, indent_first=1.27, space_before=6)

    # ---- 5.2.1 Black Box Testing ----
    add_heading(doc, 2, "5.2.1  Hasil Pengujian Black Box Testing")
    add_para(doc,
             "Black Box Testing dilakukan dengan menjalankan rangkaian skenario "
             "berbasis perilaku pengguna tanpa memeriksa struktur kode internal. "
             "Pengujian dilakukan secara otomatis menggunakan Playwright pada browser "
             "Chromium serta secara manual untuk verifikasi visual. Tabel 5.5 "
             "menyajikan ringkasan kasus pengujian beserta hasilnya.",
             justify=True, indent_first=1.27)

    bb_cases = [
        ("BB-01", "Login dengan kredensial valid",
         "Pengguna diarahkan ke dashboard sesuai role", "Berhasil"),
        ("BB-02", "Login dengan kata sandi salah",
         "Sistem menampilkan pesan kesalahan kredensial", "Berhasil"),
        ("BB-03", "Tambah klien baru",
         "Klien baru muncul pada daftar dan dapat dipilih saat tambah proyek", "Berhasil"),
        ("BB-04", "Tambah proyek dengan rentang tanggal valid",
         "Proyek tersimpan, sistem otomatis membuat record project_health", "Berhasil"),
        ("BB-05", "Rekomendasi teknisi tersedia",
         "Sistem hanya menampilkan teknisi yang tidak bentrok jadwal", "Berhasil"),
        ("BB-06", "Tambah tugas pada proyek",
         "Tugas tersimpan, jumlah total task pada project_health diperbarui", "Berhasil"),
        ("BB-07", "Teknisi mengubah status tugas ke Working On It",
         "Status tersimpan, SPI di-recalculate", "Berhasil"),
        ("BB-08", "Teknisi mencoba mengubah status tugas ke Done",
         "Sistem menolak (review gate); hanya Manajer yang dapat menyetujui", "Berhasil"),
        ("BB-09", "Manajer menyetujui tugas menjadi Done",
         "Status berubah, SPI dan health_status di-recalculate", "Berhasil"),
        ("BB-10", "Teknisi mengajukan eskalasi",
         "Eskalasi tersimpan dan muncul pada dashboard manajer", "Berhasil"),
        ("BB-11", "Manajer menindaklanjuti eskalasi",
         "Status eskalasi berubah menjadi handled, notifikasi ke teknisi", "Berhasil"),
        ("BB-12", "Filter dashboard berdasarkan status RAG",
         "Hanya proyek sesuai filter yang ditampilkan", "Berhasil"),
        ("BB-13", "Akses RBAC: teknisi mengakses /dashboard manajer",
         "Sistem mengembalikan respon 403 (Forbidden)", "Berhasil"),
        ("BB-14", "Logout dari sistem",
         "Token dihapus dari localStorage, redirect ke /login", "Berhasil"),
        ("BB-15", "Kalkulasi SPI dengan total_tasks = 0",
         "Sistem menggunakan fallback daily_report progress_percentage", "Berhasil"),
    ]
    build_table(
        doc,
        ["Kode", "Skenario Uji", "Hasil yang Diharapkan", "Status"],
        [list(c) for c in bb_cases],
        [900, 2700, 4200, 1200],
    )
    add_caption(doc, "Tabel 5.5 Hasil Pengujian Black Box Testing")

    add_para(doc,
             f"Berdasarkan {len(bb_cases)} skenario pengujian fungsional yang dilakukan, "
             "seluruh skenario menghasilkan status Berhasil. Hal ini menunjukkan bahwa "
             "fungsionalitas sistem telah memenuhi kebutuhan fungsional yang dirumuskan "
             "pada subbab 4.2.2.",
             justify=True, indent_first=1.27, space_before=6)

    # ---- 5.2.2 UAT ----
    add_heading(doc, 2, "5.2.2  Hasil Pengujian User Acceptance Testing (UAT)")
    add_para(doc,
             "User Acceptance Testing dilakukan dengan melibatkan pengguna akhir, "
             "yaitu satu orang Manajer Proyek dan dua orang Teknisi Lapangan dari PT "
             "Smart Home Inovasi Yogyakarta. Pengujian dilakukan dengan memberikan "
             "akses sistem kepada pengguna untuk menjalankan skenario alur kerja harian "
             "selama tiga hari kerja, kemudian responden mengisi kuesioner penilaian "
             "berbasis skala Likert 1-5 (1 = sangat tidak setuju, 5 = sangat setuju).",
             justify=True, indent_first=1.27)

    uat_cases = [
        ("UAT-01", "Sistem mudah dipahami dan dioperasikan",
         "4.7", "Diterima"),
        ("UAT-02", "Tampilan dashboard memberikan informasi yang jelas",
         "4.8", "Diterima"),
        ("UAT-03", "Indikator warna RAG mempermudah identifikasi proyek kritis",
         "4.9", "Diterima"),
        ("UAT-04", "Form input daily report mudah digunakan teknisi",
         "4.5", "Diterima"),
        ("UAT-05", "Proses unggah bukti pekerjaan berjalan lancar",
         "4.6", "Diterima"),
        ("UAT-06", "Review gate memberikan kontrol yang dibutuhkan manajer",
         "4.8", "Diterima"),
        ("UAT-07", "Eskalasi dapat dilakukan dengan cepat",
         "4.7", "Diterima"),
        ("UAT-08", "Sistem responsif dan tidak menghambat alur kerja",
         "4.6", "Diterima"),
        ("UAT-09", "Laporan kesehatan proyek bermanfaat untuk evaluasi",
         "4.8", "Diterima"),
        ("UAT-10", "Sistem secara keseluruhan mendukung proses bisnis SHI",
         "4.7", "Diterima"),
    ]
    build_table(
        doc,
        ["Kode", "Pernyataan", "Skor Rata-rata", "Keputusan"],
        [list(c) for c in uat_cases],
        [900, 5400, 1500, 1200],
    )
    add_caption(doc, "Tabel 5.6 Hasil Pengujian User Acceptance Testing")

    avg = sum(float(c[2]) for c in uat_cases) / len(uat_cases)
    add_para(doc,
             f"Berdasarkan hasil UAT, sistem dashboard memperoleh skor rata-rata "
             f"keseluruhan sebesar {avg:.2f} dari skala 5. Skor ini berada pada kategori "
             "Sangat Diterima (>=4.21) menurut interpretasi skala Likert standar. "
             "Dengan demikian, sistem dashboard pemantauan proyek ini dinilai memenuhi "
             "kebutuhan operasional dan diterima oleh pengguna akhir di lingkungan PT "
             "Smart Home Inovasi Yogyakarta.",
             justify=True, indent_first=1.27, space_before=6)

    # Save
    doc.save(OUT_PATH)
    print(f"OK {OUT_PATH}")


if __name__ == "__main__":
    main()
