"""
Generate `bab5_final/5.2.1 Pembahasan Basis Data.docx` — versi yang
berfokus pada kueri transaksi alur Early Warning System (EWS).

Empat bagian, 12 kueri, 24 gambar (SQL + hasil eksekusi per kueri):

  1. Transaksi Tabel tb_proyek          (3 kueri)
  2. Transaksi Daily Report             (3 kueri)
  3. Rekalkulasi SPI dan project_health (3 kueri)
  4. Penayangan Dashboard EWS           (3 kueri)

Pipeline:
  - Setiap kueri dieksekusi pada PostgreSQL lokal di dalam transaksi
    sendiri yang DIROLLBACK di akhir, sehingga data seed tidak berubah.
  - SQL dirender lewat Pygments → PNG (panel editor DBeaver-like).
  - Hasil eksekusi dirender sebagai tabel rapi (PNG).
  - Docx disusun dengan struktur: judul bagian → narasi → SQL → caption
    → hasil → caption.

Pre-req: Postgres lokal hidup (lihat run.py / brew services), view tb_*
sudah dibuat dari 01_views.sql.
"""

from __future__ import annotations

import importlib.util
import shutil
import sys
from pathlib import Path

import psycopg2

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT       = Path(__file__).resolve().parents[3]
TEMPLATE   = ROOT / 'naskah/bab_split/BAB_V_Implementasi_dan_Pembahasan_Sistem.docx'
OUT_PATH   = ROOT / 'naskah/bab5_final/5.2.1 Pembahasan Basis Data.docx'
SHOTS_DIR  = ROOT / 'naskah/bab5_final/screenshots_ews'
SHOTS_DIR.mkdir(parents=True, exist_ok=True)

HERE = Path(__file__).resolve().parent
sys.path.insert(0, str(HERE))

# Reuse render functions dari pipeline lama
spec = importlib.util.spec_from_file_location('rsh', HERE / '03_render_screenshots.py')
rsh  = importlib.util.module_from_spec(spec)
spec.loader.exec_module(rsh)

DB_DSN = ('host=127.0.0.1 port=5432 dbname=shi_dashboard_new '
          'user=postgres password=12345')


# ── Spec kueri ────────────────────────────────────────────────────────────

# Format: dict dengan 'op' (INSERT/SELECT/UPDATE/UPSERT), 'sql', 'narr'.
# Setiap kueri dijalankan independen lalu DIROLLBACK.

BAGIAN_1_PROYEK = [
    {
        'op': 'INSERT',
        'sql': (
            "INSERT INTO tb_proyek\n"
            "  (project_code, name, description, id_klien,\n"
            "   start_date, end_date, status, phase, category,\n"
            "   project_value, created_by)\n"
            "VALUES ('PRJ-EWS-DEMO', 'Demo Proyek untuk EWS',\n"
            "        'Pemasangan smart lock — proyek demo BAB V',\n"
            "        1, CURRENT_DATE, CURRENT_DATE + INTERVAL '60 days',\n"
            "        'active', 'execution', 'instalasi',\n"
            "        35000000, 3)\n"
            "RETURNING id_proyek, project_code, name, status, phase;"
        ),
        'narr': (
            "Kueri INSERT pada tabel tb_proyek dieksekusi setelah manajer "
            "menyetujui hasil survei dan memutuskan untuk melanjutkan "
            "pekerjaan ke fase eksekusi. Data yang dimasukkan mencakup "
            "kode proyek unik, deskripsi pekerjaan, referensi id_klien, "
            "rentang tanggal pelaksanaan yang menjadi acuan perhitungan "
            "Planned Value pada SPI, status awal active, dan nilai "
            "kontrak. Klausa RETURNING memberi konfirmasi langsung "
            "atas id_proyek yang dibangkitkan, sekaligus mendukung "
            "operasi inisialisasi baris kesehatan proyek pada langkah "
            "berikutnya."
        ),
    },
    {
        'op': 'SELECT',
        'sql': (
            "SELECT p.id_proyek, p.project_code, p.name,\n"
            "       p.status, p.phase, p.end_date,\n"
            "       c.name AS klien\n"
            "FROM tb_proyek p\n"
            "JOIN tb_klien c ON c.id_klien = p.id_klien\n"
            "WHERE p.status = 'active'\n"
            "ORDER BY p.end_date ASC\n"
            "LIMIT 8;"
        ),
        'narr': (
            "Kueri SELECT pada tabel tb_proyek dengan JOIN ke tb_klien "
            "menghasilkan daftar proyek aktif terurut berdasarkan tanggal "
            "akhir terdekat. Hasil pengambilan ini menjadi data sumber "
            "halaman Data Proyek dan komponen daftar proyek pada "
            "Dashboard Early Warning System (EWS). Urutan tanggal akhir "
            "memberi gambaran beban kerja yang menjelang tenggat, "
            "sehingga manajer dapat mendahulukan proyek yang paling "
            "mendekati deadline pelaksanaan."
        ),
    },
    {
        'op': 'UPDATE',
        'sql': (
            "UPDATE tb_proyek\n"
            "SET status = 'on-hold',\n"
            "    phase  = 'execution',\n"
            "    end_date = end_date + INTERVAL '14 days'\n"
            "WHERE id_proyek = 1\n"
            "RETURNING id_proyek, status, phase, end_date;"
        ),
        'narr': (
            "Kueri UPDATE pada tabel tb_proyek dijalankan ketika manajer "
            "menyesuaikan status atau tenggat proyek, misalnya menghentikan "
            "sementara pekerjaan (on-hold) atau memperpanjang tanggal "
            "akhir. Karena tabel ini menjadi pusat operasional, setiap "
            "perubahan kolom start_date atau end_date akan memengaruhi "
            "nilai Planned Value pada perhitungan SPI berikutnya, "
            "sehingga juga menggeser kategori kesehatan proyek pada "
            "dashboard EWS."
        ),
    },
]

BAGIAN_2_DAILY_REPORT = [
    {
        'op': 'INSERT',
        'sql': (
            "INSERT INTO task_activities\n"
            "  (task_id, user_id, message, activity_type)\n"
            "VALUES (1, 5,\n"
            "        'Pemasangan sensor area lobi selesai, lanjut ruang server',\n"
            "        'note')\n"
            "RETURNING id, task_id, user_id, activity_type, created_at;"
        ),
        'narr': (
            "Kueri INSERT pada tabel task_activities menjadi titik tumpu "
            "daily report di sisi teknisi. Setiap baris merepresentasikan "
            "satu catatan aktivitas terhadap sebuah tugas — bisa berupa "
            "catatan teks (note), tanda mulai pekerjaan (start_work), "
            "jeda (pause), kembali bekerja (resume), atau penanda selesai "
            "(complete). Kolom user_id menyimpan identitas teknisi yang "
            "membuat catatan, sedangkan task_id mengaitkan aktivitas "
            "dengan tugas pada tabel tb_tugas. Kolom created_at terisi "
            "otomatis sebagai stempel waktu pelaporan."
        ),
    },
    {
        'op': 'SELECT',
        'sql': (
            "SELECT a.id, a.message, a.activity_type,\n"
            "       a.created_at,\n"
            "       t.name AS tugas,\n"
            "       p.name AS proyek\n"
            "FROM task_activities a\n"
            "JOIN tb_tugas  t ON t.id_tugas  = a.task_id\n"
            "JOIN tb_proyek p ON p.id_proyek = t.id_proyek\n"
            "WHERE a.user_id = 5\n"
            "  AND a.created_at >= CURRENT_DATE - INTERVAL '30 days'\n"
            "ORDER BY a.created_at DESC\n"
            "LIMIT 8;"
        ),
        'narr': (
            "Kueri SELECT pada tabel task_activities dengan JOIN ke "
            "tb_tugas dan tb_proyek menarik histori daily report seorang "
            "teknisi dalam tiga puluh hari terakhir. Hasil pengambilan menjadi "
            "umpan data komponen \"Aktivitas Hari Ini\" pada Dashboard "
            "Performa Teknisi serta linimasa aktivitas pada halaman "
            "detail tugas. Pengurutan created_at descending memastikan "
            "catatan terbaru muncul lebih dahulu sehingga manajer dapat "
            "memverifikasi progres lapangan secara langsung."
        ),
    },
    {
        'op': 'UPDATE',
        'sql': (
            "-- Transaksi inti yang memicu rekalkulasi SPI:\n"
            "-- UPDATE tb_tugas berhasil → handler memanggil recalculateSPI()\n"
            "-- yang melakukan UPSERT pada project_health (lihat Bagian 3).\n"
            "UPDATE tb_tugas\n"
            "SET status = 'done',\n"
            "    time_spent_seconds = COALESCE(time_spent_seconds, 0)\n"
            "         + GREATEST(\n"
            "             EXTRACT(EPOCH FROM (now() - status_changed_at))::int,\n"
            "             0),\n"
            "    status_changed_at = now()\n"
            "WHERE id_tugas = 1\n"
            "RETURNING id_tugas, status, time_spent_seconds, status_changed_at;"
        ),
        'narr': (
            "Kueri UPDATE pada tabel tb_tugas menjadi transaksi paling "
            "kritis di sisi daily report. Selain mengganti kolom status, "
            "kueri ini juga mengakumulasikan durasi yang terlewati sejak "
            "status_changed_at terakhir ke kolom time_spent_seconds, "
            "lalu mereset stempel waktu transisi. Pola ini hanya berlaku "
            "ketika tugas keluar dari status aktif (in_progress atau "
            "working_on_it). Begitu UPDATE berhasil, lapisan aplikasi "
            "memicu fungsi recalculateSPI() yang melakukan UPSERT pada "
            "project_health — inilah jembatan antara daily report dan "
            "Dashboard EWS."
        ),
    },
]

BAGIAN_3_SPI_HEALTH = [
    {
        'op': 'SELECT',
        'sql': (
            "-- Penanda \"aman vs terlambat\" di level tugas. Klausa FILTER\n"
            "-- memisahkan tugas yang sedang dikerjakan tapi lewat tenggat\n"
            "-- (overtime) dari tugas yang belum dimulai padahal lewat tenggat\n"
            "-- (overdue). Hasil ini dipakai recalculateSPI() untuk mengisi\n"
            "-- kolom overtime_tasks dan overdue_tasks pada project_health.\n"
            "SELECT\n"
            "  COUNT(*)::int                                           AS total,\n"
            "  COUNT(*) FILTER (WHERE status = 'done')::int            AS completed,\n"
            "  COUNT(*) FILTER (WHERE status IN ('working_on_it',\n"
            "                                    'in_progress'))::int  AS working,\n"
            "  COUNT(*) FILTER (WHERE status IN ('working_on_it',\n"
            "                                    'in_progress')\n"
            "                         AND due_date < CURRENT_DATE)::int AS overtime,\n"
            "  COUNT(*) FILTER (WHERE due_date < CURRENT_DATE\n"
            "                         AND status NOT IN ('done'))::int  AS overdue\n"
            "FROM tb_tugas\n"
            "WHERE id_proyek = 1;"
        ),
        'narr': (
            "Kueri SELECT ini menghitung lima besaran tugas dalam satu "
            "operasi: total, selesai, sedang dikerjakan, overtime, dan "
            "overdue. Dua klausa FILTER terakhir merupakan inti dari "
            "sintaks penanda \"terlambat\" pada level tugas. Kolom "
            "overtime menandai tugas yang masih dalam pengerjaan tetapi "
            "due_date sudah lewat, sedangkan overdue menandai tugas "
            "berstatus apa pun selain done yang due_date-nya sudah lewat. "
            "Nilai ini diturunkan ke kolom overtime_tasks dan "
            "overdue_tasks pada project_health untuk dipakai dashboard EWS."
        ),
    },
    {
        'op': 'UPSERT',
        'sql': (
            "INSERT INTO project_health\n"
            "  (project_id, spi_value, status,\n"
            "   actual_progress, planned_progress,\n"
            "   total_tasks, completed_tasks, working_tasks,\n"
            "   overtime_tasks, overdue_tasks, last_updated)\n"
            "VALUES (1, 0.4500, 'red',\n"
            "        30.00, 66.67,\n"
            "        10, 3, 2, 2, 3, NOW())\n"
            "ON CONFLICT (project_id) DO UPDATE SET\n"
            "  spi_value        = EXCLUDED.spi_value,\n"
            "  status           = EXCLUDED.status,\n"
            "  actual_progress  = EXCLUDED.actual_progress,\n"
            "  planned_progress = EXCLUDED.planned_progress,\n"
            "  total_tasks      = EXCLUDED.total_tasks,\n"
            "  completed_tasks  = EXCLUDED.completed_tasks,\n"
            "  working_tasks    = EXCLUDED.working_tasks,\n"
            "  overtime_tasks   = EXCLUDED.overtime_tasks,\n"
            "  overdue_tasks    = EXCLUDED.overdue_tasks,\n"
            "  last_updated     = NOW()\n"
            "RETURNING project_id, spi_value, status, actual_progress,\n"
            "         planned_progress, last_updated;"
        ),
        'narr': (
            "Kueri UPSERT pada tabel project_health merupakan transaksi "
            "tulis utama dari proses rekalkulasi SPI. Klausa ON CONFLICT "
            "(project_id) DO UPDATE menjamin satu proyek selalu memiliki "
            "tepat satu baris kesehatan — bila baris belum ada akan "
            "diinsert, bila sudah ada akan ditimpa dengan nilai terbaru. "
            "Kolom status (green/amber/red) merupakan kategori kesehatan "
            "yang menjadi sortir utama Dashboard EWS pada Bagian 4, "
            "sedangkan kolom last_updated dipakai komponen \"diperbarui "
            "X menit lalu\" pada antarmuka."
        ),
    },
    {
        'op': 'SELECT',
        'sql': (
            "SELECT ph.project_id, p.name AS proyek,\n"
            "       ph.spi_value, ph.status AS rag,\n"
            "       ph.actual_progress, ph.planned_progress,\n"
            "       ph.total_tasks, ph.completed_tasks,\n"
            "       ph.overtime_tasks, ph.overdue_tasks,\n"
            "       ph.last_updated\n"
            "FROM project_health ph\n"
            "JOIN tb_proyek p ON p.id_proyek = ph.project_id\n"
            "WHERE ph.project_id = 1;"
        ),
        'narr': (
            "Kueri SELECT pada tabel project_health menarik baris "
            "kesehatan terkini suatu proyek beserta nama proyeknya dari "
            "tb_proyek. Kolom yang dihasilkan mencerminkan satu kesatuan "
            "snapshot kesehatan: nilai SPI, kategori warna (green/amber/"
            "red), persentase progres aktual dan rencana, serta hitungan "
            "tugas pada masing-masing kondisi. Snapshot inilah yang "
            "diserap komponen kartu proyek (ProjectCard) untuk menampilkan "
            "indikator SPI dan badge kesehatan pada Dashboard EWS."
        ),
    },
]

BAGIAN_4_DASHBOARD_EWS = [
    {
        'op': 'SELECT',
        'sql': (
            "-- Kueri utama Dashboard EWS: proyek aktif diurutkan berdasarkan\n"
            "-- kekritisan. RED selalu di puncak, lalu AMBER, lalu GREEN.\n"
            "-- Di dalam satu warna, SPI terkecil (paling kritis) di atas.\n"
            "SELECT p.id_proyek, p.name AS proyek,\n"
            "       ph.spi_value, ph.status AS rag,\n"
            "       ph.completed_tasks, ph.total_tasks,\n"
            "       ph.overtime_tasks, ph.overdue_tasks,\n"
            "       p.end_date\n"
            "FROM tb_proyek p\n"
            "LEFT JOIN project_health ph ON ph.project_id = p.id_proyek\n"
            "WHERE p.status = 'active'\n"
            "ORDER BY\n"
            "  CASE ph.status\n"
            "    WHEN 'red'   THEN 1\n"
            "    WHEN 'amber' THEN 2\n"
            "    WHEN 'green' THEN 3\n"
            "    ELSE 4\n"
            "  END,\n"
            "  ph.spi_value ASC NULLS LAST,\n"
            "  p.end_date  ASC;"
        ),
        'narr': (
            "Kueri SELECT inilah yang mendefinisikan tampilan utama "
            "Dashboard EWS. JOIN antara tb_proyek dan project_health "
            "menggabungkan informasi proyek dengan hasil rekalkulasi SPI "
            "dari Bagian 3. Klausa ORDER BY menggunakan ekspresi CASE "
            "pada ph.status sehingga proyek kategori red selalu menempati "
            "urutan teratas, disusul amber, lalu green. Sebagai pengurut "
            "kedua, ph.spi_value ASC menempatkan proyek dengan SPI "
            "terkecil (paling kritis) di atas pada warna yang sama. "
            "Tie-breaker terakhir berupa end_date ASC memastikan proyek "
            "dengan tenggat paling dekat lebih dahulu muncul."
        ),
    },
    {
        'op': 'SELECT',
        'sql': (
            "-- Ringkasan kesehatan proyek untuk panel kepala dashboard.\n"
            "-- Setiap klausa COUNT FILTER menghitung jumlah proyek pada\n"
            "-- satu kategori RAG. Inilah \"summary kesehatan\" yang dipakai\n"
            "-- kartu statistik atas dashboard EWS.\n"
            "SELECT\n"
            "  COUNT(*) FILTER (WHERE p.status='active')::int                AS active_projects,\n"
            "  COUNT(*) FILTER (WHERE ph.status='red'   AND p.status='active')::int AS total_red,\n"
            "  COUNT(*) FILTER (WHERE ph.status='amber' AND p.status='active')::int AS total_amber,\n"
            "  COUNT(*) FILTER (WHERE ph.status='green' AND p.status='active')::int AS total_green,\n"
            "  COUNT(*) FILTER (WHERE ph.status IS NULL AND p.status='active')::int AS total_no_health,\n"
            "  ROUND(AVG(ph.spi_value) FILTER (WHERE p.status='active'), 4) AS avg_spi,\n"
            "  COUNT(*) FILTER (WHERE p.end_date < CURRENT_DATE\n"
            "                         AND p.status='active')::int           AS overdue_projects\n"
            "FROM tb_proyek p\n"
            "LEFT JOIN project_health ph ON ph.project_id = p.id_proyek;"
        ),
        'narr': (
            "Kueri SELECT ini meringkas seluruh portofolio proyek menjadi "
            "satu baris statistik yang dipakai panel kepala Dashboard EWS. "
            "Setiap klausa COUNT FILTER membagi proyek aktif menurut "
            "kategori RAG (total_red, total_amber, total_green), serta "
            "mencatat proyek aktif yang belum memiliki baris project_health "
            "(total_no_health). Kolom avg_spi memberi rerata kesehatan "
            "keseluruhan sebagai sinyal makro, sedangkan overdue_projects "
            "menghitung proyek yang melewati tanggal akhir kontrak — "
            "indikator keterlambatan pada level proyek (bukan tugas)."
        ),
    },
    {
        'op': 'SELECT',
        'sql': (
            "-- Daftar tugas terlambat per proyek, untuk panel\n"
            "-- \"Tugas Terlambat per Proyek\" pada Dashboard EWS.\n"
            "SELECT p.id_proyek AS project_id,\n"
            "       p.name      AS project_name,\n"
            "       COUNT(*) FILTER (WHERE t.status IN ('working_on_it',\n"
            "                                           'in_progress'))::int AS overdue_working,\n"
            "       COUNT(*) FILTER (WHERE t.status = 'to_do')::int          AS overdue_todo\n"
            "FROM tb_tugas t\n"
            "JOIN tb_proyek p ON p.id_proyek = t.id_proyek\n"
            "WHERE t.due_date < CURRENT_DATE\n"
            "  AND t.status NOT IN ('done')\n"
            "  AND p.status = 'active'\n"
            "GROUP BY p.id_proyek, p.name\n"
            "HAVING COUNT(*) > 0\n"
            "ORDER BY COUNT(*) DESC;"
        ),
        'narr': (
            "Kueri SELECT ini menyajikan rincian tugas yang sudah melewati "
            "tenggat per proyek aktif. Filter pada klausa WHERE menyaring "
            "tugas yang due_date-nya sudah lewat dan statusnya bukan done. "
            "Klausa COUNT FILTER memisahkan tugas yang sedang dikerjakan "
            "(overdue_working) dari tugas yang belum dimulai (overdue_todo) "
            "agar manajer dapat membedakan beban kerja yang masih "
            "berjalan dari beban kerja yang sama sekali terabaikan. "
            "Klausa HAVING COUNT(*) > 0 menjamin hanya proyek yang "
            "benar-benar memiliki keterlambatan yang ditayangkan, dan "
            "ORDER BY COUNT(*) DESC menyusun proyek terburuk di atas."
        ),
    },
]

ALL_SECTIONS = [
    ('1.', 'Transaksi Tabel tb_proyek',                 'proyek_ews',
     BAGIAN_1_PROYEK),
    ('2.', 'Transaksi Daily Report',                    'daily_report_ews',
     BAGIAN_2_DAILY_REPORT),
    ('3.', 'Rekalkulasi SPI dan project_health',        'spi_health_ews',
     BAGIAN_3_SPI_HEALTH),
    ('4.', 'Penayangan Dashboard EWS',                  'dashboard_ews',
     BAGIAN_4_DASHBOARD_EWS),
]


# ── Eksekusi + render PNG ─────────────────────────────────────────────────

def run_query_isolated(conn, sql: str):
    """Run a single SQL statement inside its own transaction, then
    rollback. Strips leading SQL comments so PostgreSQL's empty-query
    error doesn't fire when the entire body is comment-only.
    Returns: (columns, rows, status_text)."""
    cur = conn.cursor()
    body_lines = [ln for ln in sql.split('\n') if ln.strip()]
    runnable = '\n'.join(body_lines)
    # First non-comment word drives the status label.
    non_comment = [ln for ln in body_lines if not ln.lstrip().startswith('--')]
    head = ' '.join(non_comment).lstrip()
    op_word = head.split()[0].upper() if head else ''
    try:
        cur.execute(runnable)
        cols = [d.name for d in cur.description] if cur.description else []
        rows = cur.fetchall() if cols else []
        rc   = cur.rowcount
        if op_word == 'INSERT':
            status = f"INSERT 0 {rc}  •  {rc} row affected"
        elif op_word == 'UPDATE':
            status = f"UPDATE {rc}  •  {rc} row affected"
        elif op_word == 'DELETE':
            status = f"DELETE {rc}  •  {rc} row affected"
        elif op_word == 'SELECT':
            status = f"SELECT {rc}  •  {rc} rows returned"
        else:
            status = f"OK  •  {rc} row affected"
    except Exception as e:
        cols, rows = [], []
        status = f"ERROR: {type(e).__name__}: {e}"
        conn.rollback()
        cur.close()
        return cols, rows, status
    finally:
        try:
            conn.rollback()
        except Exception:
            pass
    cur.close()
    return cols, rows, status


def _fmt_cell(v):
    if v is None:
        return 'NULL'
    return str(v)


def render_result_panel_clipped(table, op_idx, op, columns, rows,
                                 status_text, out_path, W=1280):
    """Versi lokal render_result_panel dengan:
      - Lebar kanvas diperbesar (W=1280) untuk akomodasi tabel lebar.
      - Teks setiap sel diklip per pixel kolom menggunakan font.getlength()
        agar tidak overlap dengan kolom sebelahnya.
    """
    from PIL import Image, ImageDraw
    H_TITLE = 32
    H_TAB   = 30
    H_HDR   = 32
    ROW_H   = 28
    PAD     = 12
    font_mono = rsh.F_MONO_S
    font_hdr  = rsh.F_TITLE

    if not columns:
        H = H_TITLE + H_TAB + 90
        canvas = Image.new('RGB', (W, H), rsh.C_BG)
        d = ImageDraw.Draw(canvas)
        rsh._draw_chrome(d, W, H_TITLE, H_TAB, table, op, 'Output')
        d.text((PAD + 4, H_TITLE + H_TAB + 16), status_text,
               font=font_mono, fill=(20, 110, 60))
        d.rectangle([0, 0, W - 1, H - 1], outline=rsh.C_BORDER)
        canvas.save(out_path, 'PNG', optimize=True)
        return

    str_rows = [[_fmt_cell(c) for c in r] for r in rows]
    col_w = []
    for ci, c in enumerate(columns):
        max_str = max([len(c)] + [len(r[ci]) for r in str_rows])
        col_w.append(min(38, max(10, max_str)))
    total_units = sum(col_w)
    body_w = W - 24
    pixel_w = [int(body_w * w / total_units) for w in col_w]
    pixel_w[-1] += body_w - sum(pixel_w)

    def fit(s, max_px):
        if font_mono.getlength(s) <= max_px:
            return s
        ell = '…'
        ell_px = font_mono.getlength(ell)
        while s and font_mono.getlength(s) + ell_px > max_px:
            s = s[:-1]
        return s + ell if s else ''

    rows_h = max(1, len(str_rows)) * ROW_H
    H = H_TITLE + H_TAB + H_HDR + rows_h + 36 + 24
    canvas = Image.new('RGB', (W, H), '#ffffff')
    d = ImageDraw.Draw(canvas)

    rsh._draw_chrome(d, W, H_TITLE, H_TAB, table, op, 'Result')

    y = H_TITLE + H_TAB
    d.rectangle([0, y, W, y + H_HDR], fill=rsh.C_HEADER)
    x = 12
    for ci, c in enumerate(columns):
        d.text((x + 8, y + 8), fit(c, pixel_w[ci] - 16),
               font=font_hdr, fill=rsh.C_HEADER_T)
        x += pixel_w[ci]
        d.line([x, y + 4, x, y + H_HDR - 4], fill=rsh.C_BORDER)
    d.rectangle([0, y + H_HDR - 1, W, y + H_HDR], fill=rsh.C_BORDER)

    y += H_HDR
    if not str_rows:
        d.text((20, y + 8), '(0 rows)', font=font_mono, fill=(140, 140, 150))
        y += ROW_H
    else:
        for ri, row in enumerate(str_rows):
            if ri % 2 == 1:
                d.rectangle([0, y, W, y + ROW_H], fill=rsh.C_ROW_ALT)
            x = 12
            for ci, cell in enumerate(row):
                d.text((x + 8, y + 6), fit(cell, pixel_w[ci] - 16),
                       font=font_mono, fill=rsh.C_TEXT)
                x += pixel_w[ci]
            y += ROW_H
            d.rectangle([0, y - 1, W, y], fill=(232, 234, 240))

    y += 8
    d.rectangle([0, y, W, y + 28], fill=rsh.C_HEADER)
    d.text((14, y + 6), status_text, font=font_hdr, fill=(20, 110, 60))

    d.rectangle([0, 0, W - 1, H - 1], outline=rsh.C_BORDER)
    canvas.save(out_path, 'PNG', optimize=True)


def render_sql_panel_wide(table, op_idx, op, sql, out_path, W=1280):
    """Bungkus render_sql_panel dengan canvas yang sama lebarnya
    sebagai render_result_panel_clipped, supaya SQL + Result punya
    proporsi visual yang konsisten saat ditempatkan berurutan."""
    from PIL import Image, ImageDraw
    sql_img = rsh._render_sql_image(sql, W - 2)
    H_TITLE = 32
    H_TAB   = 30
    H = H_TITLE + H_TAB + sql_img.height + 2

    canvas = Image.new('RGB', (W, H), rsh.C_BG)
    d = ImageDraw.Draw(canvas)
    d.rectangle([0, 0, W, H_TITLE], fill=rsh.C_TITLEBAR)
    d.rectangle([0, H_TITLE - 1, W, H_TITLE], fill=rsh.C_BORDER)
    d.text((14, 9), f"PostgreSQL — {table}  •  Query Editor",
           font=rsh.F_TITLE, fill=(80, 88, 104))
    for i, c in enumerate([(255, 95, 86), (255, 189, 46), (39, 201, 63)]):
        d.ellipse([W - 78 + i * 22, 10, W - 66 + i * 22, 22], fill=c)
    tab_y0 = H_TITLE
    d.rectangle([0, tab_y0, W, tab_y0 + H_TAB], fill=rsh.C_TAB)
    d.rectangle([12, tab_y0 + 4, 280, tab_y0 + H_TAB],
                fill=rsh.C_TAB_ACT, outline=rsh.C_BORDER)
    d.text((22, tab_y0 + 8), f"{op} — {table}",
           font=rsh.F_TITLE, fill=rsh.C_TEXT)
    d.rectangle([0, tab_y0 + H_TAB - 1, W, tab_y0 + H_TAB], fill=rsh.C_BORDER)

    canvas.paste(sql_img, (1, tab_y0 + H_TAB))
    d.rectangle([0, 0, W - 1, H - 1], outline=rsh.C_BORDER)
    canvas.save(out_path, 'PNG', optimize=True)


def render_section_screenshots(section_key: str, queries: list):
    """Render SQL + result PNGs for every query in a section."""
    conn = psycopg2.connect(DB_DSN)
    conn.autocommit = False
    out = []
    for idx, entry in enumerate(queries, start=1):
        op   = entry['op']
        sql  = entry['sql']
        cols, rows, status = run_query_isolated(conn, sql)
        sql_png    = SHOTS_DIR / f"{section_key}_{idx:02d}_{op.lower()}_sql.png"
        result_png = SHOTS_DIR / f"{section_key}_{idx:02d}_{op.lower()}_result.png"
        render_sql_panel_wide(section_key, idx, op, sql, sql_png)
        render_result_panel_clipped(section_key, idx, op, cols, rows,
                                     status, result_png)
        print(f"  [{op:6}] {status}")
        print(f"          SQL    -> {sql_png.name}")
        print(f"          RESULT -> {result_png.name}")
        out.append((idx, op, entry, sql_png, result_png, status))
    conn.close()
    return out


# ── docx helpers ──────────────────────────────────────────────────────────

def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        p.style = doc.styles['Caption']
    except KeyError:
        pass
    p.add_run(text)
    return p


def add_centered_image(doc, path: Path, width_cm: float = 14.8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    return p


def add_paragraf(doc, text, style='Paragraf'):
    p = doc.add_paragraph()
    try:
        p.style = doc.styles[style]
    except KeyError:
        pass
    p.add_run(text)
    return p


# ── main ──────────────────────────────────────────────────────────────────

def build():
    # 1) Eksekusi semua kueri + render PNG
    rendered = {}
    for cat_num, cat_name, key, queries in ALL_SECTIONS:
        print(f"\n── Bagian {cat_num} {cat_name} ──")
        rendered[key] = render_section_screenshots(key, queries)

    # 2) Susun docx
    shutil.copy(TEMPLATE, OUT_PATH)
    doc = Document(OUT_PATH)
    body = doc.element.body
    sect_pr = body.find(qn('w:sectPr'))
    for el in list(body):
        if el is sect_pr:
            continue
        body.remove(el)

    # Heading 5.2.1
    p = doc.add_paragraph()
    try:
        p.style = doc.styles['5.2.n']
    except KeyError:
        pass
    p.add_run('Pembahasan Basis Data')

    # Paragraf pembuka
    add_paragraf(doc,
        "Subbab ini menyajikan pembahasan basis data sistem manajemen "
        "proyek PT Smart Home Inovasi (SHI) dengan fokus pada kueri "
        "transaksi yang menjalankan alur Early Warning System (EWS). "
        "Basis data diimplementasikan di atas PostgreSQL 17 dan diakses "
        "oleh lapisan aplikasi melalui pustaka pg dengan parameterized "
        "query. Tabel inti yang dirancang pada BAB IV ditampilkan dengan "
        "nomenklatur tb_* (tb_user, tb_klien, tb_proyek, "
        "tb_penugasan_proyek, tb_tugas, tb_bukti, tb_eskalasi) melalui "
        "view auto-updatable, sedangkan tabel pendukung implementasi "
        "(task_activities, daily_reports, dan project_health) "
        "ditampilkan dengan nama fisiknya karena di luar lingkup "
        "rancangan BAB IV."
    )
    add_paragraf(doc,
        "Empat bagian dibahas secara berurutan mengikuti alur EWS: "
        "(1) transaksi pada tabel tb_proyek yang merepresentasikan "
        "pendaftaran dan pengelolaan proyek, (2) transaksi daily report "
        "yang berasal dari aktivitas teknisi di lapangan, "
        "(3) rekalkulasi Schedule Performance Index (SPI) beserta "
        "penyimpanannya pada tabel project_health, dan (4) kueri "
        "penayangan pada Dashboard EWS yang mengurutkan proyek "
        "berdasarkan indikator kesehatan. Setiap kueri ditampilkan "
        "lengkap dengan sintaks SQL beserta hasil eksekusinya pada "
        "PostgreSQL sebagai bukti bahwa perintah berjalan dengan benar. "
        "Untuk menjaga konsistensi data seed, kueri tulis (INSERT/UPDATE/"
        "UPSERT) dieksekusi di dalam transaksi terisolasi yang "
        "DIROLLBACK setelah hasil pengembalian dicatat."
    )

    fig_no = 1
    for cat_num, cat_name, key, queries in ALL_SECTIONS:
        # Header bagian
        th = doc.add_paragraph()
        try:
            th.style = doc.styles['Paragraf']
        except KeyError:
            pass
        run = th.add_run(f"{cat_num}  {cat_name}")
        run.bold = True

        for idx, op, entry, sql_png, result_png, status in rendered[key]:
            # Narasi
            np = doc.add_paragraph()
            try:
                np.style = doc.styles['Paragraf']
            except KeyError:
                pass
            np.add_run(
                entry['narr'] + " "
                f"Eksekusi kueri disajikan pada Gambar 5.{fig_no} dan "
                f"hasil eksekusinya pada Gambar 5.{fig_no + 1}."
            )

            # SQL screenshot
            add_centered_image(doc, sql_png, width_cm=14.8)
            add_caption(doc, f"Gambar 5.{fig_no} Kueri {op} — {cat_name}")
            fig_no += 1

            # Result screenshot
            add_centered_image(doc, result_png, width_cm=14.8)
            add_caption(doc, f"Gambar 5.{fig_no} Hasil eksekusi {op} — "
                             f"{cat_name}")
            fig_no += 1

    doc.save(OUT_PATH)
    print(f"\nSaved {OUT_PATH}")
    print(f"  Bagian: {len(ALL_SECTIONS)} • Gambar: 5.1 – 5.{fig_no - 1}")


if __name__ == '__main__':
    build()
