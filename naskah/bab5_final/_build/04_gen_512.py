"""
Generate `bab5_final/5.1.2 Implementasi Sistem.docx`.

For each module from BAB IV §4.3.5 Antarmuka:
  - heading (numbered "a./b./c.")
  - 1 narrative paragraph
  - 1 code block (monospaced text) with the corresponding backend handler
  - caption "Gambar 5.x Implementasi <Modul>"

Code is taken from frontend/src/lib/handlers/*.ts (TypeScript).
"""

import re
import json
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT       = Path(__file__).resolve().parents[3]
HANDLERS   = ROOT / 'frontend/src/lib/handlers'
TEMPLATE   = ROOT / 'naskah/bab_split/BAB_V_Implementasi_dan_Pembahasan_Sistem.docx'
OUT_PATH   = ROOT / 'naskah/bab5_final/5.1.2 Implementasi Sistem.docx'
HERE       = Path(__file__).resolve().parent
CUT_JSON   = HERE / 'gemini_outputs/cut_code.json'

sys.path.insert(0, str(HERE))
from code_transform import transform_code

# ── BAB IV §4.3.5 modules → handler function ───────────────────────────────
# Format: (category, letter, module name, handler file, function name, narrative)
NARR_LOGIN = (
    "Halaman Login merupakan gerbang autentikasi utama sistem. Implementasi "
    "ditangani oleh fungsi login() pada modul handler auth. Fungsi ini "
    "menerima payload email dan kata sandi, memvalidasi format input, "
    "mencari pengguna pada tabel tb_user berdasarkan email, kemudian "
    "membandingkan kata sandi dengan kolom password_hash menggunakan "
    "bcrypt. Jika kredensial valid, sistem membangkitkan JSON Web Token "
    "(JWT) yang ditempelkan pada cookie httpOnly sebagai mekanisme sesi."
)
NARR_TAMBAH_PROYEK = (
    "Halaman Tambah Proyek diakses oleh pengguna dengan peran manajer "
    "untuk merekam proyek baru. Fungsi createProject() melakukan validasi "
    "skema payload (nama, klien, tanggal mulai, tanggal akhir, kategori), "
    "menghasilkan kode proyek unik, melakukan operasi INSERT pada tabel "
    "tb_proyek, lalu menginisialisasi baris pada tabel kesehatan proyek "
    "dengan SPI awal sebesar 1.0. Operasi dijalankan dalam satu transaksi "
    "agar konsistensi tetap terjaga."
)
NARR_DAILY_REPORT = (
    "Halaman Tambah Daily Report digunakan oleh teknisi untuk mencatat "
    "aktivitas harian terhadap suatu tugas. Fungsi createActivity() "
    "memvalidasi keberadaan id_tugas, memastikan tugas tersebut "
    "ditugaskan kepada teknisi yang sedang aktif, mencatat aktivitas "
    "ke dalam tabel task_activities, kemudian memperbarui kolom "
    "time_spent_seconds pada tb_tugas. Sistem secara otomatis menarik "
    "rekap aktivitas harian via fungsi getMyTodayActivities()."
)
NARR_DASHBOARD_EWS = (
    "Dashboard Early Warning System (EWS) menjadi luaran inti sistem. "
    "Fungsi getDashboard() melakukan agregasi data dari tabel "
    "tb_proyek dan tabel kesehatan proyek untuk menghitung distribusi "
    "status proyek (RAG: Green/Amber/Red), mengurutkan proyek "
    "berdasarkan SPI menurun, serta menyusun ringkasan metrik kunci "
    "seperti jumlah proyek aktif, proyek dengan SPI < 0.85, dan beban "
    "kerja teknisi. Hasil agregasi dikemas dalam satu respons JSON."
)
NARR_DATA_PROYEK = (
    "Halaman Data Proyek menampilkan tabel interaktif berisi seluruh "
    "proyek aktif. Fungsi listProjects() membaca tabel tb_proyek dengan "
    "JOIN ke tb_klien untuk menarik nama pelanggan, lalu menerapkan "
    "filter berdasarkan parameter query (status, fase, kategori). "
    "Pengurutan default berdasarkan kolom created_at descending sehingga "
    "proyek terbaru muncul di urutan teratas."
)
NARR_KANBAN = (
    "Halaman Kanban Penugasan Proyek memungkinkan pengguna mengubah "
    "status tugas dengan menggeser kartu antar kolom (to_do, in_progress, "
    "review, done). Fungsi changeTaskStatus() menangani perubahan status, "
    "memvalidasi otorisasi (teknisi tidak dapat menandai status done), "
    "mencatat status_changed_at, kemudian memicu rekalkulasi SPI pada "
    "tabel kesehatan proyek karena perubahan status tugas berdampak "
    "langsung pada nilai earned value."
)
NARR_JADWAL = (
    "Halaman Jadwal Proyek menampilkan linimasa tugas pada kalender. "
    "Fungsi getScheduleTasks() menarik kumpulan tugas dengan rentang "
    "timeline_start dan timeline_end yang berpotongan dengan jendela "
    "waktu yang diminta klien. Hasil pengambilan dikelompokkan berdasarkan "
    "teknisi sehingga dapat divisualisasikan sebagai swimlane pada "
    "antarmuka kalender."
)
NARR_PERFORMA_TEKNISI = (
    "Dashboard Performa Teknisi ditujukan bagi pengguna dengan peran "
    "teknisi. Fungsi getTechnicianDashboard() melakukan agregasi tugas "
    "yang ditugaskan kepada teknisi yang sedang aktif, menghitung "
    "distribusi tugas berdasarkan status, jumlah tugas overdue, total "
    "waktu kerja harian, serta ringkasan eskalasi yang dilaporkan. "
    "Data ini digunakan untuk memotivasi teknisi melalui visualisasi "
    "produktivitas pribadi."
)
NARR_DETAIL_PROYEK = (
    "Halaman Detail Proyek menjadi pusat pengawasan untuk satu proyek. "
    "Fungsi getProject() menarik baris dari tabel tb_proyek beserta JOIN "
    "ke tb_klien dan tabel kesehatan proyek, kemudian melampirkan "
    "rincian tim teknisi (tb_penugasan_proyek) serta jumlah tugas per "
    "status. Respons yang dihasilkan menjadi sumber tunggal data bagi "
    "seluruh komponen tab pada halaman detail."
)
NARR_LAPORAN_KESEHATAN = (
    "Halaman Laporan Kesehatan Proyek menyediakan analitik komprehensif "
    "tentang performa proyek dari segi jadwal. Fungsi chartEarnedValue() "
    "menghitung kurva Planned Value (PV), Earned Value (EV), dan Actual "
    "Cost (AC) selama rentang proyek berdasarkan persentase penyelesaian "
    "tugas. Hasil perhitungan dipakai untuk menggambar kurva S yang "
    "menjadi indikator visual kesehatan jadwal proyek."
)

MODULES = [
    ('1.', 'Antarmuka Data Input', [
        ('a.', 'Halaman Login',            'auth.ts',       'login',                 NARR_LOGIN),
        ('b.', 'Halaman Tambah Proyek',    'projects.ts',   'createProject',         NARR_TAMBAH_PROYEK),
        ('c.', 'Halaman Tambah Daily Report','activities.ts','createActivity',       NARR_DAILY_REPORT),
    ]),
    ('2.', 'Antarmuka Data Proses', [
        ('a.', 'Halaman Dashboard Early Warning System (EWS)', 'dashboard.ts', 'getDashboard', NARR_DASHBOARD_EWS),
        ('b.', 'Halaman Data Proyek',            'projects.ts',   'listProjects',          NARR_DATA_PROYEK),
        ('c.', 'Halaman Kanban Penugasan Proyek','tasks.ts',      'changeTaskStatus',      NARR_KANBAN),
        ('d.', 'Halaman Jadwal Proyek',          'tasks.ts',      'getScheduleTasks',      NARR_JADWAL),
    ]),
    ('3.', 'Antarmuka Data Output', [
        ('a.', 'Halaman Dashboard Performa Teknisi', 'dashboard.ts', 'getTechnicianDashboard', NARR_PERFORMA_TEKNISI),
        ('b.', 'Halaman Detail Proyek',              'projects.ts',  'getProject',             NARR_DETAIL_PROYEK),
        ('c.', 'Halaman Laporan Kesehatan Proyek',   'dashboard.ts', 'chartEarnedValue',       NARR_LAPORAN_KESEHATAN),
    ]),
]


def get_code(filename: str, fn_name: str) -> str:
    """Return cut+tb_*-transformed code from gemini_outputs/cut_code.json.
    Falls back to raw extraction + transform if no Gemini cut available."""
    if CUT_JSON.exists():
        data = json.load(open(CUT_JSON))
        key = f"{filename}::{fn_name}"
        if key in data:
            return data[key]['code']
    # Fallback path
    text = (HANDLERS / filename).read_text()
    m = re.search(rf'^export (async )?function {fn_name}\b.*?\{{', text, re.M | re.S)
    if not m:
        raise RuntimeError(f"Function not found: {filename}:{fn_name}")
    start = m.start(); depth = 0; i = m.end() - 1
    while i < len(text):
        c = text[i]
        if c == '{': depth += 1
        elif c == '}':
            depth -= 1
            if depth == 0:
                return transform_code(text[start:i+1].rstrip())
        i += 1
    raise RuntimeError("brace mismatch")


# ── docx helpers ────────────────────────────────────────────────────────────
def _set_para_border(p, color="A0A4B0", size="4", space="6"):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:space'), space)
        b.set(qn('w:color'), color)
        pBdr.append(b)
    pPr.append(pBdr)


def _set_para_shading(p, fill="F4F5F8"):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def _set_para_keep(p):
    pPr = p._p.get_or_add_pPr()
    for tag in ('w:keepNext', 'w:keepLines'):
        el = OxmlElement(tag)
        pPr.append(el)


def add_code_block(doc, code: str):
    """Add a monospaced code block. Each non-empty line -> own paragraph.
    Border + shading wraps the whole block by applying to each para."""
    lines = code.split('\n')
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.4)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Consolas'
        # Ensure rFonts gets East Asian fallback too
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), 'Consolas')
        rFonts.set(qn('w:hAnsi'), 'Consolas')
        rFonts.set(qn('w:cs'), 'Consolas')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x20, 0x24, 0x30)
        _set_para_shading(p)
        # Borders: only top on first, bottom on last, sides on all
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        sides_to_draw = ['left', 'right']
        if i == 0: sides_to_draw.insert(0, 'top')
        if i == len(lines) - 1: sides_to_draw.append('bottom')
        for side in sides_to_draw:
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '4')
            b.set(qn('w:space'), '4')
            b.set(qn('w:color'), 'A0A4B0')
            pBdr.append(b)
        pPr.append(pBdr)
        _set_para_keep(p)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        p.style = doc.styles['Caption']
    except KeyError:
        pass
    run = p.add_run(text)
    run.italic = False
    run.bold = False
    return p


# ── main ────────────────────────────────────────────────────────────────────
def build():
    # Start from template to inherit styles
    shutil.copy(TEMPLATE, OUT_PATH)
    doc = Document(OUT_PATH)
    body = doc.element.body
    sect_pr = body.find(qn('w:sectPr'))
    # Wipe all body children except sectPr
    for el in list(body):
        if el is sect_pr:
            continue
        body.remove(el)

    # Title heading
    p = doc.add_paragraph()
    try:
        p.style = doc.styles['5.1.n']
    except KeyError:
        pass
    p.add_run('Implementasi Sistem')

    # Intro paragraph
    intro = doc.add_paragraph()
    try:
        intro.style = doc.styles['Paragraf New']
    except KeyError:
        pass
    intro.add_run(
        "Implementasi sistem merupakan tahap perwujudan rancangan menjadi "
        "modul-modul kode yang dapat dieksekusi. Sistem dibangun dengan "
        "kerangka kerja Next.js (App Router) yang menjalankan rute API "
        "lewat satu dispatcher terpusat. Setiap endpoint dipetakan ke "
        "fungsi handler pada lib/handlers, dengan akses basis data "
        "PostgreSQL menggunakan pustaka pg dan parameterized query. "
        "Subbab ini menyajikan implementasi sintaks dari fungsi-fungsi "
        "handler kunci untuk setiap modul antarmuka yang dirancang pada "
        "BAB IV §4.3.5."
    )

    bridge = doc.add_paragraph()
    try:
        bridge.style = doc.styles['Paragraf New']
    except KeyError:
        pass
    bridge.add_run(
        "Catatan pemetaan nama: listing kode pada subbab ini menggunakan "
        "nomenklatur BAB IV (tb_user, tb_klien, tb_proyek, "
        "tb_penugasan_proyek, tb_tugas, tb_bukti, dan tb_eskalasi) untuk "
        "menjaga konsistensi dengan rancangan basis data. Pada penyimpanan "
        "fisik PostgreSQL, tabel-tabel tersebut diimplementasikan dengan "
        "nama singkat (users, clients, projects, dan seterusnya) dan "
        "dijembatani oleh view auto-updatable, sehingga seluruh perintah "
        "SQL yang ditampilkan tetap valid saat dieksekusi langsung pada "
        "basis data. Nomenklatur kolom (id_user, id_klien, dan "
        "seterusnya) telah diselaraskan dengan rancangan BAB IV melalui "
        "alias pada view tersebut."
    )

    fig_no = 1
    for cat_num, cat_name, modules in MODULES:
        # Category header (e.g., "1. Antarmuka Data Input")
        ph = doc.add_paragraph()
        try:
            ph.style = doc.styles['Paragraf New']
        except KeyError:
            pass
        rcat = ph.add_run(f"{cat_num}  {cat_name}")
        rcat.bold = True

        for letter, modname, fname, fn_name, narrative in modules:
            # Module entry header (e.g., "a. Halaman Login")
            hp = doc.add_paragraph()
            try:
                hp.style = doc.styles['Paragraf New']
            except KeyError:
                pass
            hp.paragraph_format.left_indent = Cm(0.6)
            hp.add_run(f"{letter}  {modname}").bold = True

            # Narrative
            np = doc.add_paragraph()
            try:
                np.style = doc.styles['Paragraf New']
            except KeyError:
                pass
            np.paragraph_format.left_indent = Cm(0.6)
            np.add_run(narrative + " Kode implementasinya disajikan pada "
                       f"Gambar 5.{fig_no}.")

            # Code block
            code = get_code(fname, fn_name)
            add_code_block(doc, code)

            # Caption
            cap = add_caption(doc, f"Gambar 5.{fig_no} Implementasi {modname}")
            fig_no += 1

    doc.save(OUT_PATH)
    print(f"Saved {OUT_PATH}  •  figures: {fig_no - 1}")


if __name__ == '__main__':
    build()
