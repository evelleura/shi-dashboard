"""
Generate `bab5_final/5.1.2 Implementasi Sistem.docx` — versi terfokus
pada alur Early Warning System (EWS).

Alur yang dibahas (3 bagian, 9 gambar):

  1. Input Daily Report (Transaksi Tulis)
       1a. createActivity()      — INSERT task_activities + otorisasi teknisi
       1b. changeTaskStatus()    — UPDATE tasks + memicu rekalkulasi SPI
       1c. screenshot            — antarmuka teknisi mengubah status tugas

  2. Proses Rekalkulasi & Penandaan Aman/Terlambat (Transaksi Proses)
       2a. recalculateSPI()      — hitung SPI, UPSERT project_health
       2b. categorizeHealth()    — threshold green/amber/red
       2c. getTaskCounts() SQL   — FILTER overtime + overdue

  3. Penayangan Dashboard EWS (Transaksi Baca)
       3a. getDashboard() SQL    — JOIN + ORDER BY RAG → SPI ASC
       3b. chartOverdueTasks()   — daftar tugas terlambat per proyek
       3c. screenshot            — dashboard EWS dengan badge RAG

Semua cuplikan kode dikutip langsung dari file source agar konsisten
dengan implementasi nyata. Tabel `task_activities`, `project_health`,
dan `daily_reports` tampil dengan nama fisiknya (tidak masuk ke
nomenklatur BAB IV).
"""

import shutil
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT     = Path(__file__).resolve().parents[3]
TEMPLATE = ROOT / 'naskah/bab_split/BAB_V_Implementasi_dan_Pembahasan_Sistem.docx'
OUT_PATH = ROOT / 'naskah/bab5_final/5.1.2 Implementasi Sistem.docx'
SHOT_DIR = ROOT / 'alur-sistem'
HERE     = Path(__file__).resolve().parent

sys.path.insert(0, str(HERE))
from code_transform import transform_code   # type: ignore


# ── Code excerpts (dikutip langsung dari source; lihat komentar HEAD) ─────

# HEAD: frontend/src/lib/handlers/activities.ts  (createActivity, L23-86)
CODE_1A_CREATE_ACTIVITY = """\
export async function createActivity(request: NextRequest) {
  const auth = authenticateRequest(request);
  if (!auth.user) return auth.errorResponse;

  const formData = await request.formData();
  const file          = formData.get('file') as File | null;
  const task_id       = formData.get('task_id') as string;
  const message       = formData.get('message') as string;
  const activity_type = formData.get('activity_type') as string | null;

  // Validasi payload
  if (!task_id)
    return NextResponse.json({ success: false, error: 'task_id is required' }, { status: 400 });
  if (!message || message.trim().length === 0)
    return NextResponse.json({ success: false, error: 'message is required' }, { status: 400 });

  const finalType = activity_type && VALID_ACTIVITY_TYPES.includes(activity_type)
    ? activity_type : 'note';

  // Otorisasi: teknisi hanya boleh mencatat aktivitas untuk tugas
  // yang ditugaskan kepadanya.
  const taskCheck = await query(
    'SELECT id, project_id, assigned_to FROM tasks WHERE id = $1', [task_id]);
  if (taskCheck.rowCount === 0)
    return NextResponse.json({ success: false, error: 'Task not found' }, { status: 404 });
  const task = taskCheck.rows[0];
  if (auth.user.role === 'technician' && task.assigned_to !== auth.user.userId) {
    return NextResponse.json(
      { success: false, error: 'You can only add activities to tasks assigned to you' },
      { status: 403 });
  }

  // Transaksi tulis utama: catat aktivitas harian ke task_activities.
  const result = await query(
    `INSERT INTO task_activities
       (task_id, user_id, message, activity_type,
        file_path, file_name, file_type, file_size)
     VALUES ($1, $2, $3, $4, $5, $6, $7, $8) RETURNING *`,
    [task_id, auth.user.userId, message.trim(), finalType,
     filePath, fileName, fileType, fileSize]);

  return NextResponse.json({ success: true, data: result.rows[0] }, { status: 201 });
}
"""

# HEAD: frontend/src/lib/handlers/tasks.ts  (changeTaskStatus, kutipan
# bagian inti L516-540 — UPDATE bersyarat + pemicu rekalkulasi SPI)
CODE_1B_CHANGE_STATUS = """\
// Bagian inti changeTaskStatus(): UPDATE bersyarat lalu memicu
// rekalkulasi SPI. Jika tugas keluar dari status aktif (in_progress /
// working_on_it), durasi pengerjaan diakumulasikan ke time_spent_seconds.
const wasActive = ACTIVE_STATUSES.has(currentStatus);
const result = await query(
  wasActive
    ? `UPDATE tasks
         SET status = $1,
             time_spent_seconds = COALESCE(time_spent_seconds, 0)
                                + GREATEST(EXTRACT(EPOCH FROM
                                            (NOW() - status_changed_at))::int, 0),
             status_changed_at = NOW(),
             updated_at = NOW()
       WHERE id = $2 RETURNING *`
    : `UPDATE tasks
         SET status = $1,
             status_changed_at = NOW(),
             updated_at = NOW()
       WHERE id = $2 RETURNING *`,
  [status, taskId]);

// Pemicu utama EWS: setiap perubahan status memanggil recalculateSPI()
// sehingga baris project_health untuk proyek terkait diperbarui.
await recalculateSPI(task.project_id as number);

// Catat aktivitas otomatis (start_work / pause / resume / complete)
// agar daily report konsisten dengan transisi status tugas.
const activityType = activityTypeForTransition(currentStatus, status);
"""

# HEAD: frontend/src/lib/spiCalculator.ts  (recalculateSPI, L111-177)
CODE_2A_RECALC_SPI = """\
export async function recalculateSPI(projectId: number)
  : Promise<ProjectHealth | null> {
  // Ambil rentang waktu proyek
  const projectResult = await query<{ start_date: Date; end_date: Date; status: string }>(
    'SELECT start_date, end_date, status FROM projects WHERE id = $1', [projectId]);
  if (projectResult.rowCount === 0) return null;
  const project = projectResult.rows[0];

  // Rekap jumlah tugas: total, selesai, sedang dikerjakan,
  // overtime, dan overdue. Dipakai untuk menghitung rasio progres
  // aktual dan untuk pelaporan EWS.
  const taskCounts = await getTaskCounts(projectId);

  // Planned Value (PV) = (hari berjalan / total durasi) × 100
  const today        = new Date();
  const plannedValue = calculatePlannedValue(
    project.start_date, project.end_date, today);

  // Actual Progress = (tugas selesai / total tugas) × 100.
  // Bila proyek belum memiliki tugas, gunakan progress_percentage
  // dari laporan harian terakhir sebagai fallback.
  let actualProgress: number;
  if (taskCounts.total === 0) {
    const reportProgress = await getLatestReportProgress(projectId);
    actualProgress = reportProgress !== null ? reportProgress : 0;
  } else {
    actualProgress = (taskCounts.completed / taskCounts.total) * 100;
  }

  // SPI = Actual Progress / Planned Value
  const spiValue = plannedValue > 0
    ? Math.round((actualProgress / plannedValue) * 10000) / 10000
    : 1;

  const status = categorizeHealth(spiValue);   // → 'green' | 'amber' | 'red'

  // UPSERT ke project_health: satu baris per proyek, last_updated diperbarui.
  const upsertResult = await query<ProjectHealth>(
    `INSERT INTO project_health
       (project_id, spi_value, status, actual_progress, planned_progress,
        total_tasks, completed_tasks, working_tasks,
        overtime_tasks, overdue_tasks, last_updated)
     VALUES ($1, $2, $3, $4, $5, $6, $7, $8, $9, $10, NOW())
     ON CONFLICT (project_id) DO UPDATE SET
       spi_value        = EXCLUDED.spi_value,
       status           = EXCLUDED.status,
       actual_progress  = EXCLUDED.actual_progress,
       planned_progress = EXCLUDED.planned_progress,
       total_tasks      = EXCLUDED.total_tasks,
       completed_tasks  = EXCLUDED.completed_tasks,
       working_tasks    = EXCLUDED.working_tasks,
       overtime_tasks   = EXCLUDED.overtime_tasks,
       overdue_tasks    = EXCLUDED.overdue_tasks,
       last_updated     = NOW()
     RETURNING *`,
    [projectId, spiValue, status,
     Math.round(actualProgress * 100) / 100, plannedValue,
     taskCounts.total, taskCounts.completed, taskCounts.working,
     taskCounts.overtime, taskCounts.overdue]);

  return upsertResult.rows[0];
}
"""

# HEAD: frontend/src/lib/spiCalculator.ts  (categorizeHealth, L47-51)
CODE_2B_CATEGORIZE = """\
/**
 * Memetakan nilai SPI ke kategori kesehatan proyek.
 *   Aman / On Track    : SPI ≥ 0.95   → 'green'
 *   Peringatan / Warn  : 0.85 ≤ SPI < 0.95   → 'amber'
 *   Terlambat / Critical : SPI < 0.85   → 'red'
 */
export function categorizeHealth(spiValue: number): HealthStatus {
  if (spiValue >= 0.95) return 'green';
  if (spiValue >= 0.85) return 'amber';
  return 'red';
}
"""

# HEAD: frontend/src/lib/spiCalculator.ts  (getTaskCounts, L56-89)
CODE_2C_GET_TASK_COUNTS = """\
// Penghitungan dua kategori "terlambat" pada level tugas:
//   - overtime : sedang dikerjakan TAPI due_date sudah lewat
//   - overdue  : status apa pun selain 'done' DAN due_date sudah lewat
// Hasilnya dipakai recalculateSPI() untuk mengisi kolom
// overtime_tasks dan overdue_tasks pada project_health.
const result = await query<{
  total: string; completed: string; working: string;
  overtime: string; overdue: string;
}>(
  `SELECT
     COUNT(*)::int                                            AS total,
     COUNT(*) FILTER (WHERE status = 'done')::int             AS completed,
     COUNT(*) FILTER (WHERE status IN ('working_on_it',
                                       'in_progress'))::int   AS working,
     COUNT(*) FILTER (WHERE status IN ('working_on_it',
                                       'in_progress')
                            AND due_date < CURRENT_DATE)::int AS overtime,
     COUNT(*) FILTER (WHERE due_date < CURRENT_DATE
                            AND status NOT IN ('done'))::int  AS overdue
   FROM tasks
   WHERE project_id = $1`, [projectId]);
"""

# HEAD: frontend/src/lib/handlers/dashboard.ts  (getDashboard, L24-46)
CODE_3A_GET_DASHBOARD = """\
// Query utama EWS: gabungkan proyek aktif dengan baris kesehatannya,
// lalu urutkan berdasarkan kekritisan. Proyek MERAH selalu di atas,
// disusul KUNING, lalu HIJAU; di dalam satu warna, SPI terkecil
// (paling kritis) didahulukan, dan tanggal akhir terdekat berikutnya.
const projectsResult = await query(
  `SELECT p.id, p.name, p.description, p.client_id, p.start_date, p.end_date,
     p.duration, p.status, p.phase, p.project_value,
     c.name AS client_name,
     ph.spi_value, ph.status AS health_status,
     ph.actual_progress, ph.planned_progress,
     ph.last_updated AS health_last_updated,
     ph.total_tasks, ph.completed_tasks, ph.working_tasks,
     ph.overtime_tasks, ph.overdue_tasks,
     dr.constraints           AS latest_constraints,
     dr.report_date           AS last_report_date,
     dr.progress_percentage   AS last_reported_progress
   FROM projects p
   LEFT JOIN clients         c  ON c.id  = p.client_id
   LEFT JOIN project_health  ph ON ph.project_id = p.id
   LEFT JOIN LATERAL (
     SELECT constraints, report_date, progress_percentage
     FROM daily_reports
     WHERE project_id = p.id
     ORDER BY report_date DESC, created_at DESC LIMIT 1
   ) dr ON true
   WHERE p.status = 'active'
   ORDER BY
     CASE ph.status
       WHEN 'red'   THEN 1
       WHEN 'amber' THEN 2
       WHEN 'green' THEN 3
       ELSE 4
     END,
     ph.spi_value ASC NULLS LAST,
     p.end_date  ASC`);
"""

# HEAD: frontend/src/lib/handlers/dashboard.ts  (statsResult, L57-69)
CODE_3B_STATS = """\
// Statistik ringkas untuk panel kepala dashboard. Setiap kolom
// menghitung proyek yang masuk dalam satu kategori RAG ─ inilah
// "penanda aman vs terlambat" pada level agregat.
const statsResult = await query(
  `SELECT COUNT(*)::int                                              AS total_projects,
     COUNT(*) FILTER (WHERE p.status = 'active')::int                AS active_projects,
     COUNT(*) FILTER (WHERE ph.status = 'red'   AND p.status='active')::int AS total_red,
     COUNT(*) FILTER (WHERE ph.status = 'amber' AND p.status='active')::int AS total_amber,
     COUNT(*) FILTER (WHERE ph.status = 'green' AND p.status='active')::int AS total_green,
     COUNT(*) FILTER (WHERE ph.status IS NULL   AND p.status='active')::int AS total_no_health,
     ROUND(AVG(ph.spi_value) FILTER (WHERE p.status='active'), 4)    AS avg_spi,
     COUNT(*) FILTER (WHERE p.end_date < CURRENT_DATE
                            AND p.status='active')::int              AS overdue_projects
   FROM projects p
   LEFT JOIN project_health ph ON ph.project_id = p.id
   WHERE 1=1`);
"""

# HEAD: frontend/src/lib/handlers/dashboard.ts  (chartOverdueTasks, L207-238)
CODE_3C_OVERDUE = """\
export async function chartOverdueTasks(request: NextRequest) {
  const auth = authenticateRequest(request);
  if (!auth.user) return auth.errorResponse;
  const roleCheck = authorizeRoles(auth.user, ['manager', 'admin']);
  if (roleCheck) return roleCheck;

  // Daftar proyek beserta jumlah tugas terlambatnya, dipakai
  // panel "Tugas Terlambat per Proyek" pada dashboard EWS.
  const result = await query(
    `SELECT p.id   AS project_id,
       p.name AS project_name,
       COUNT(*) FILTER (WHERE t.status IN ('working_on_it',
                                           'in_progress'))::int AS overdue_working,
       COUNT(*) FILTER (WHERE t.status = 'to_do')::int          AS overdue_todo
     FROM tasks t
     JOIN projects p ON p.id = t.project_id
     WHERE t.due_date < CURRENT_DATE
       AND t.status NOT IN ('done')
       AND p.status = 'active'
     GROUP BY p.id, p.name
     HAVING COUNT(*) > 0
     ORDER BY COUNT(*) DESC`);

  return NextResponse.json({ success: true, data: result.rows });
}
"""


# ── Narasi per bagian ────────────────────────────────────────────────────

NARR_INTRO = (
    "Subbab ini menyajikan implementasi alur Early Warning System (EWS) "
    "secara terfokus, yaitu rangkaian transaksi yang dimulai dari "
    "pencatatan progres harian oleh teknisi (daily report) sampai dengan "
    "penayangan kondisi proyek pada dashboard EWS. Tiga tahap aliran "
    "dibahas berurutan: (1) transaksi tulis pada saat teknisi memasukkan "
    "laporan harian, (2) transaksi proses yang menghitung ulang nilai "
    "Schedule Performance Index (SPI) dan memutakhirkan penanda "
    "kesehatan proyek, serta (3) transaksi baca yang menyusun tampilan "
    "EWS bagi manajer."
)
NARR_BRIDGE = (
    "Pada listing kode di subbab ini, tabel inti BAB IV tetap "
    "ditampilkan dengan nomenklatur tb_user, tb_klien, tb_proyek, "
    "tb_penugasan_proyek, tb_tugas, tb_bukti, dan tb_eskalasi. "
    "Tabel pendukung implementasi (task_activities untuk catatan "
    "aktivitas harian, daily_reports untuk arsip laporan persen "
    "manual, dan project_health untuk turunan SPI) ditampilkan "
    "dengan nama fisiknya karena bukan bagian rancangan basis "
    "data BAB IV."
)

NARR_BAGIAN_1 = (
    "Bagian ini memuat dua transaksi tulis yang berfungsi sebagai daily "
    "report di sisi teknisi. Transaksi pertama mencatat aktivitas pada "
    "sebuah tugas (foto, catatan, atau penanda waktu), sedangkan "
    "transaksi kedua mengubah status tugas dan secara tidak langsung "
    "menjadi pemicu utama pemutakhiran dashboard EWS."
)
NARR_1A = (
    "Fungsi createActivity() menangani penambahan catatan aktivitas "
    "harian terhadap satu tugas. Validasi payload dilakukan pada awal, "
    "dilanjutkan dengan otorisasi sehingga teknisi hanya dapat "
    "menambahkan aktivitas pada tugas yang ditugaskan kepada dirinya "
    "sendiri (perbandingan task.assigned_to terhadap auth.user.userId). "
    "Transaksi inti berupa INSERT pada tabel task_activities dengan "
    "kolom message, activity_type, dan opsional file_path bila "
    "diunggah bukti foto/dokumen. Sintaks implementasi disajikan "
    "pada Gambar 5.1."
)
NARR_1B = (
    "Fungsi changeTaskStatus() pada modul tasks menjadi titik tumpu "
    "alur EWS karena setiap perubahan status oleh teknisi maupun manajer "
    "akan memicu rekalkulasi SPI. Operasi UPDATE pada tabel tugas "
    "memperbarui kolom status, status_changed_at, dan updated_at; "
    "khusus pada transisi yang meninggalkan status aktif (in_progress "
    "atau working_on_it), durasi pengerjaan diakumulasikan ke kolom "
    "time_spent_seconds. Setelah UPDATE berhasil, dilakukan pemanggilan "
    "recalculateSPI(task.project_id) yang menjadi pemicu utama "
    "pemutakhiran penanda terlambat pada dashboard EWS. Aktivitas "
    "transisi (start_work, pause, resume, complete) juga dicatat "
    "otomatis ke task_activities sehingga konsisten dengan catatan "
    "manual pada bagian a. Sintaks implementasi disajikan pada "
    "Gambar 5.2, sedangkan tampilan teknisi mengubah status disajikan "
    "pada Gambar 5.3."
)

NARR_BAGIAN_2 = (
    "Bagian ini memuat tiga komponen kunci dari proses rekalkulasi "
    "yang dipicu oleh transaksi pada Bagian 1. Komponen pertama "
    "menghitung nilai SPI dan menyimpannya pada tabel project_health. "
    "Komponen kedua mengklasifikasikan SPI menjadi tiga kategori "
    "kesehatan proyek (aman, peringatan, terlambat). Komponen ketiga "
    "berisi sintaks SQL yang membedakan tugas terlambat dari tugas "
    "yang masih aman."
)
NARR_2A = (
    "Fungsi recalculateSPI() merupakan jantung perhitungan EWS. Fungsi "
    "ini mengambil rentang waktu proyek dari tabel proyek, kemudian "
    "memperoleh hitungan tugas (total, selesai, sedang dikerjakan, "
    "overtime, dan overdue) melalui getTaskCounts(). Planned Value "
    "(PV) dihitung sebagai persentase hari yang sudah berjalan terhadap "
    "total durasi proyek, sedangkan Actual Progress dihitung sebagai "
    "rasio jumlah tugas selesai terhadap jumlah tugas total. Bila "
    "proyek belum memiliki tugas, sistem mengambil progress_percentage "
    "dari laporan harian terakhir pada tabel daily_reports sebagai "
    "fallback. Hasil pembagian Actual Progress dengan Planned Value "
    "menghasilkan nilai SPI; nilai inilah yang menentukan apakah "
    "proyek tergolong aman, peringatan, atau terlambat. Operasi "
    "UPSERT pada tabel project_health dilakukan dengan klausa ON "
    "CONFLICT (project_id) sehingga satu proyek selalu memiliki satu "
    "baris kesehatan yang konsisten. Sintaks implementasi disajikan "
    "pada Gambar 5.4."
)
NARR_2B = (
    "Fungsi categorizeHealth() memetakan nilai SPI ke tiga label warna "
    "yang dipakai dashboard. Proyek dengan SPI lebih dari atau sama "
    "dengan 0,95 dikategorikan green (label On Track / aman), proyek "
    "dengan SPI antara 0,85 dan 0,95 dikategorikan amber (Warning / "
    "peringatan), dan proyek dengan SPI di bawah 0,85 dikategorikan "
    "red (Critical / terlambat). Threshold ini menjadi acuan tunggal "
    "yang dipakai baik di sisi basis data (kolom status pada "
    "project_health) maupun di sisi antarmuka (komponen StatusBadge "
    "yang merender lingkaran hijau, kuning, atau merah). Sintaks "
    "implementasi disajikan pada Gambar 5.5."
)
NARR_2C = (
    "Selain pada level proyek, sistem juga membedakan tugas yang "
    "terlambat dari yang aman menggunakan dua klausa FILTER pada SQL. "
    "Kolom overtime menghitung tugas yang sedang dikerjakan namun "
    "tenggatnya sudah terlewati (status IN ('working_on_it', "
    "'in_progress') AND due_date < CURRENT_DATE), sedangkan kolom "
    "overdue menghitung tugas pada status apa pun selain done yang "
    "tenggatnya sudah terlewati. Cuplikan SQL ini berjalan di dalam "
    "getTaskCounts() yang dipanggil oleh recalculateSPI() pada Bagian "
    "2a, sehingga setiap perubahan status tugas akan langsung "
    "memperbarui nilai overtime_tasks dan overdue_tasks pada "
    "project_health. Sintaks implementasi disajikan pada Gambar 5.6."
)

NARR_BAGIAN_3 = (
    "Bagian ini memuat transaksi baca yang menampilkan hasil "
    "rekalkulasi dari Bagian 2 kepada manajer. Query utama dashboard "
    "EWS mengurutkan proyek berdasarkan tingkat kekritisan, panel "
    "statistik meringkas jumlah proyek per kategori RAG, dan grafik "
    "tugas terlambat menampilkan rincian per proyek."
)
NARR_3A = (
    "Fungsi getDashboard() menyusun tampilan EWS dengan dua tahap. "
    "Pertama, melakukan JOIN antara tabel proyek dengan baris "
    "project_health dan laporan harian terakhir (subquery LATERAL "
    "pada daily_reports). Kedua, mengurutkan hasil JOIN dengan klausa "
    "ORDER BY CASE ph.status sehingga proyek berkategori red menempati "
    "urutan teratas, disusul amber, lalu green. Di dalam satu kategori "
    "warna, proyek diurutkan menurut nilai SPI menaik supaya proyek "
    "yang paling kritis (SPI terkecil) muncul lebih dahulu; sebagai "
    "tie-breaker terakhir, tanggal akhir terdekat didahulukan. Pola "
    "pengurutan inilah yang membuat manajer langsung melihat proyek "
    "terlambat di pucuk dashboard tanpa perlu memilah daftar secara "
    "manual. Sintaks implementasi disajikan pada Gambar 5.7."
)
NARR_3B = (
    "Selain daftar proyek, dashboard juga menampilkan panel statistik "
    "ringkas berupa jumlah proyek pada masing-masing kategori RAG. "
    "Query menggunakan beberapa klausa COUNT FILTER untuk menghitung "
    "total_red, total_amber, total_green, dan total_no_health "
    "(proyek tanpa baris project_health), serta avg_spi sebagai "
    "rerata kesehatan secara keseluruhan. Kolom overdue_projects "
    "menambah dimensi keterlambatan pada level proyek dengan "
    "membandingkan end_date dengan tanggal sistem. Sintaks "
    "implementasi disajikan pada Gambar 5.8."
)
NARR_3C = (
    "Fungsi chartOverdueTasks() menyajikan rincian tugas terlambat "
    "untuk panel \"Tugas Terlambat per Proyek\" pada dashboard EWS. "
    "Query mem-filter tugas yang due_date-nya sudah terlewati dan "
    "statusnya bukan done, kemudian mengelompokkan per proyek dengan "
    "memisahkan jumlah tugas yang sedang dikerjakan (overdue_working) "
    "dari yang belum dimulai (overdue_todo). Klausa HAVING COUNT(*) > "
    "0 memastikan hanya proyek yang benar-benar memiliki tugas "
    "terlambat yang muncul, sedangkan ORDER BY COUNT(*) DESC menyusun "
    "proyek dengan keterlambatan terbanyak di atas. Sintaks "
    "implementasi disajikan pada Gambar 5.9, sedangkan tampilan "
    "dashboard EWS dengan pengurutan RAG nyata disajikan pada "
    "Gambar 5.10."
)


# ── Struktur dokumen ──────────────────────────────────────────────────────

# Setiap item: (letter, judul subitem, narasi, code|None, screenshot|None,
#               caption pendek setelah "Gambar 5.x")
BAGIAN_1 = ('1.', 'Transaksi Input Daily Report', NARR_BAGIAN_1, [
    ('a.', 'Pencatatan Aktivitas Harian Tugas',
        NARR_1A, CODE_1A_CREATE_ACTIVITY, None,
        'Implementasi createActivity() — INSERT INTO task_activities'),
    ('b.', 'Perubahan Status Tugas sebagai Pemicu Rekalkulasi',
        NARR_1B, CODE_1B_CHANGE_STATUS, None,
        'Implementasi changeTaskStatus() — UPDATE tasks dan pemicu '
        'recalculateSPI()'),
    ('c.', 'Antarmuka Pengubahan Status oleh Teknisi',
        ('Antarmuka teknisi memperlihatkan kartu tugas pada papan Kanban. '
         'Perubahan kolom kartu memanggil endpoint changeTaskStatus() yang '
         'memuat dua transaksi pada Bagian 1a dan 1b di atas. Tampilan '
         'antarmukanya disajikan pada Gambar 5.3.'),
        None, '06-teknisi-update-task.png',
        'Antarmuka teknisi mengubah status tugas'),
])

BAGIAN_2 = ('2.', 'Transaksi Proses Rekalkulasi dan Penanda Aman / Terlambat',
            NARR_BAGIAN_2, [
    ('a.', 'Perhitungan SPI dan UPSERT ke project_health',
        NARR_2A, CODE_2A_RECALC_SPI, None,
        'Implementasi recalculateSPI() — komputasi SPI dan UPSERT '
        'project_health'),
    ('b.', 'Klasifikasi Aman, Peringatan, dan Terlambat',
        NARR_2B, CODE_2B_CATEGORIZE, None,
        'Implementasi categorizeHealth() — threshold green/amber/red'),
    ('c.', 'Sintaks SQL Penanda Tugas Terlambat',
        NARR_2C, CODE_2C_GET_TASK_COUNTS, None,
        'Klausa FILTER untuk overtime dan overdue pada getTaskCounts()'),
])

BAGIAN_3 = ('3.', 'Transaksi Baca pada Dashboard EWS', NARR_BAGIAN_3, [
    ('a.', 'Query Utama dengan Pengurutan Berdasarkan Kekritisan',
        NARR_3A, CODE_3A_GET_DASHBOARD, None,
        'Implementasi getDashboard() — JOIN proyek dan project_health '
        'dengan ORDER BY RAG'),
    ('b.', 'Statistik Ringkasan Proyek per Kategori RAG',
        NARR_3B, CODE_3B_STATS, None,
        'Implementasi statistik panel kepala EWS — COUNT FILTER per '
        'kategori RAG'),
    ('c.', 'Daftar Tugas Terlambat per Proyek dan Tampilan EWS',
        NARR_3C, CODE_3C_OVERDUE, None,
        'Implementasi chartOverdueTasks() — daftar tugas terlambat per proyek'),
    ('d.', 'Tampilan Dashboard EWS dengan Sortir RAG',
        ('Tampilan dashboard EWS pada Gambar 5.10 menunjukkan hasil '
         'pengurutan dari query getDashboard(). Kartu proyek berbingkai '
         'merah (kategori terlambat) menempati posisi teratas, disusul '
         'kartu berbingkai kuning (peringatan) dan hijau (aman). Setiap '
         'kartu menampilkan nilai SPI, rasio tugas selesai terhadap '
         'total, serta progres aktual vs rencana yang berasal langsung '
         'dari kolom-kolom project_health hasil rekalkulasi pada '
         'Bagian 2.'),
        None, '10-dashboard-ews-spi-health.png',
        'Tampilan Dashboard EWS dengan urutan proyek berdasarkan '
        'kekritisan'),
])

ALL_SECTIONS = [BAGIAN_1, BAGIAN_2, BAGIAN_3]


# ── docx helpers (sama gaya dengan 04_gen_512.py) ─────────────────────────

def _set_para_shading(p, fill='F4F5F8'):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def _set_para_keep(p):
    pPr = p._p.get_or_add_pPr()
    for tag in ('w:keepNext', 'w:keepLines'):
        el = OxmlElement(tag)
        pPr.append(el)


def add_code_block(doc, code: str):
    lines = code.rstrip('\n').split('\n')
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.4)
        run = p.add_run(line if line else ' ')
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Consolas')
        rFonts.set(qn('w:hAnsi'), 'Consolas')
        rFonts.set(qn('w:cs'),    'Consolas')
        rPr.insert(0, rFonts)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x20, 0x24, 0x30)
        _set_para_shading(p)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        sides = ['left', 'right']
        if i == 0:
            sides.insert(0, 'top')
        if i == len(lines) - 1:
            sides.append('bottom')
        for side in sides:
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '4')
            b.set(qn('w:space'), '4')
            b.set(qn('w:color'), 'A0A4B0')
            pBdr.append(b)
        pPr.append(pBdr)
        _set_para_keep(p)


def add_centered_image(doc, path: Path, width_cm: float = 14.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        p.style = doc.styles['Caption']
    except KeyError:
        pass
    run = p.add_run(text)
    run.italic = False
    run.bold = False
    return p


def add_paragraf(doc, text, style='Paragraf New', indent_cm=0):
    p = doc.add_paragraph()
    try:
        p.style = doc.styles[style]
    except KeyError:
        pass
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    p.add_run(text)
    return p


# ── main ──────────────────────────────────────────────────────────────────

def build():
    shutil.copy(TEMPLATE, OUT_PATH)
    doc = Document(OUT_PATH)
    body = doc.element.body
    sect_pr = body.find(qn('w:sectPr'))
    for el in list(body):
        if el is sect_pr:
            continue
        body.remove(el)

    # Heading 5.1.2 Implementasi Sistem
    p = doc.add_paragraph()
    try:
        p.style = doc.styles['5.1.n']
    except KeyError:
        pass
    p.add_run('Implementasi Sistem')

    add_paragraf(doc, NARR_INTRO)
    add_paragraf(doc, NARR_BRIDGE)

    fig_no = 1
    for cat_num, cat_name, cat_narr, items in ALL_SECTIONS:
        # Header bagian
        ph = doc.add_paragraph()
        try:
            ph.style = doc.styles['Paragraf New']
        except KeyError:
            pass
        rcat = ph.add_run(f'{cat_num}  {cat_name}')
        rcat.bold = True

        # Narasi pengantar bagian
        add_paragraf(doc, cat_narr)

        for letter, sub_title, narr, code, screenshot, cap_short in items:
            # Sub-judul (a. / b. / c.)
            hp = doc.add_paragraph()
            try:
                hp.style = doc.styles['Paragraf New']
            except KeyError:
                pass
            hp.paragraph_format.left_indent = Cm(0.6)
            hp.add_run(f'{letter}  {sub_title}').bold = True

            # Narasi
            add_paragraf(doc, narr, indent_cm=0.6)

            # Gambar = code listing
            if code is not None:
                add_code_block(doc, transform_code(code))
                add_caption(doc, f'Gambar 5.{fig_no} {cap_short}')
                fig_no += 1

            # Gambar = screenshot real dari alur-sistem/
            if screenshot is not None:
                img_path = SHOT_DIR / screenshot
                if not img_path.is_file():
                    print(f'[WARN] Missing screenshot: {img_path}')
                else:
                    add_centered_image(doc, img_path, width_cm=14.5)
                add_caption(doc, f'Gambar 5.{fig_no} {cap_short}')
                fig_no += 1

    doc.save(OUT_PATH)
    print(f'Saved {OUT_PATH}')
    print(f'  Bagian: {len(ALL_SECTIONS)} • Gambar: 5.1 – 5.{fig_no - 1}')


if __name__ == '__main__':
    build()
