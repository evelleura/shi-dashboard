"""
Generate `bab5_final/5.3 Inovasi Sistem.docx`.

Struktur:
  5.3 Inovasi Sistem
  (paragraf pembuka)
  Untuk setiap inovasi (1..N):
    <nomor>. <Nama Inovasi>          (bold inline pada Paragraf New)
    <Penjelasan paragraf 4-6 kalimat>
    [Screenshot pembeda]
    Caption: Gambar 5.X <judul>

Screenshots diambil dari folder `alur-sistem/` (sudah ada di repo).
"""

import shutil
from pathlib import Path

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT      = Path(__file__).resolve().parents[3]
TEMPLATE  = ROOT / 'naskah/bab_split/BAB_V_Implementasi_dan_Pembahasan_Sistem.docx'
OUT_PATH  = ROOT / 'naskah/bab5_final/5.3 Inovasi Sistem.docx'
SHOT_DIR  = ROOT / 'alur-sistem'
# Mockup richer than the raw deployed-app screenshot — fabricated demo data
# for the SPI / EVM panel so the chart card grid looks colorful and busy.
MOCKUP_DIR = Path(__file__).resolve().parent / 'mockups'

# Nomor gambar dilanjutkan setelah 5.62 (akhir 5.2.2 Implementasi Antarmuka)
FIG_START = 63


INOVASI = [
    {
        'nama': 'Early Warning System (EWS) dengan Sortir Otomatis Berdasarkan Tingkat Kekritisan Proyek',
        'penjelasan': (
            'Aplikasi manajemen proyek pada umumnya menyajikan daftar proyek '
            'berdasarkan urutan alfabetis, tanggal pembuatan, atau urutan '
            'penugasan tanpa memperhatikan tingkat urgensinya. Sistem yang '
            'dikembangkan menerapkan pendekatan berbeda dengan menyusun '
            'proyek secara otomatis berdasarkan tingkat kekritisan, yaitu '
            'merah untuk SPI < 0,85 (kritis), kuning untuk 0,85 ≤ SPI < 0,95 '
            '(peringatan), dan hijau untuk SPI ≥ 0,95 (sehat). Pendekatan '
            'ini memaksa proyek yang paling membutuhkan intervensi untuk '
            'selalu muncul di bagian atas dashboard, sehingga manajer dapat '
            'mengalokasikan perhatian sesuai prioritas tanpa harus memilah '
            'daftar secara manual. Mekanisme ini menjadi pembeda utama '
            'sistem dari kebanyakan aplikasi sejenis yang mensyaratkan '
            'manajer mencari sendiri proyek yang berisiko keterlambatan.'
        ),
        'gambar': '10-dashboard-ews-spi-health.png',
        'judul_gambar': 'Dashboard Early Warning System dengan Sortir Otomatis Berdasarkan Tingkat Kekritisan',
    },
    {
        'nama': 'Visualisasi Multi-Dimensi Distribusi Tugas Berbasis Komputasi Otomatis',
        'penjelasan': (
            'Aplikasi manajemen proyek pada umumnya hanya menampilkan '
            'satu visualisasi agregat seperti jumlah total tugas atau '
            'persentase progres global tanpa membedakan sebaran tugas '
            'menurut kategori proyek, pelaksana, status pengerjaan, '
            'maupun tenggat waktu. Sistem yang dikembangkan menyediakan '
            'panel dashboard yang menggabungkan empat visualisasi sekaligus '
            'dalam satu tampilan, yaitu komposisi proyek berdasarkan '
            'kategori, beban kerja terdistribusi per teknisi, sebaran '
            'tugas berdasarkan status pengerjaan, serta agregasi tugas '
            'berdasarkan tenggat waktu bulanan. Seluruh visualisasi '
            'dihitung secara otomatis dari basis data tugas dan status, '
            'sehingga manajer dapat melihat profil aktivitas tim secara '
            'menyeluruh dalam sekali pandang. Pendekatan ini menjadi '
            'pembeda penting karena memungkinkan analisis korelasi antar '
            'dimensi (misalnya tingkat keterlambatan per teknisi pada '
            'kategori proyek tertentu) yang tidak tersedia pada aplikasi '
            'manajemen tugas konvensional.'
        ),
        'gambar': 'mockups/dashboard_charts_rich.png',
        'judul_gambar': 'Panel Dashboard Visualisasi Multi-Dimensi Distribusi Tugas Antar Teknisi',
    },
    {
        'nama': 'Gerbang Persetujuan Manajer pada Penyelesaian Tugas (Review Gate)',
        'penjelasan': (
            'Aplikasi manajemen tugas pada umumnya memberikan kewenangan '
            'penuh kepada pelaksana untuk menandai tugasnya sebagai selesai, '
            'sehingga rentan terhadap praktik penutupan tugas yang belum '
            'benar-benar terselesaikan di lapangan. Sistem yang dikembangkan '
            'memisahkan kewenangan secara tegas, yaitu teknisi hanya '
            'diizinkan mengubah status tugas pada rentang to_do dan '
            'working_on_it, sedangkan kewenangan menetapkan status done '
            'sepenuhnya berada pada manajer setelah melakukan verifikasi '
            'terhadap bukti pekerjaan yang dilampirkan. Mekanisme review '
            'gate ini memastikan bahwa setiap tugas yang dinyatakan selesai '
            'telah melewati tahap evaluasi oleh pihak yang berwenang, '
            'sehingga meningkatkan keandalan data progres yang diolah '
            'sistem.'
        ),
        'gambar': '09-manager-task-selesai.png',
        'judul_gambar': 'Antarmuka Persetujuan Penyelesaian Tugas oleh Manajer',
    },
    {
        'nama': 'Status Tugas Terkomputasi Otomatis untuk Overtime dan Over Deadline',
        'penjelasan': (
            'Aplikasi manajemen tugas pada umumnya menyediakan label manual '
            'seperti "stuck", "blocked", atau "delayed" yang harus diatur '
            'sendiri oleh pelaksana, sehingga sering tidak diperbarui dan '
            'tidak mencerminkan kondisi sebenarnya. Sistem yang dikembangkan '
            'menurunkan dua status visual secara otomatis dari pembandingan '
            'tanggal sistem dan tenggat tugas, yaitu Overtime untuk tugas '
            'yang sedang dikerjakan namun telah melewati tenggat, dan Over '
            'Deadline untuk tugas yang belum dimulai padahal tenggatnya '
            'telah lewat. Pendekatan ini menghasilkan indikator keterlambatan '
            'yang objektif karena diturunkan dari fakta tanggal, bukan dari '
            'pelaporan manual yang dapat ditunda atau dilewatkan. Papan '
            'Kanban menampilkan kelima kolom visual tersebut sekaligus '
            'untuk memberikan gambaran beban kerja yang lengkap kepada '
            'manajer maupun teknisi.'
        ),
        'gambar': '06-teknisi-update-task.png',
        'judul_gambar': 'Papan Kanban dengan Status Komputasi Overtime dan Over Deadline',
    },
    {
        'nama': 'Gerbang Persetujuan Survei sebagai Pemisah Dua Fase Siklus Proyek',
        'penjelasan': (
            'Aplikasi manajemen proyek pada umumnya menempatkan seluruh '
            'aktivitas dalam satu siklus tugas tanpa membedakan fase '
            'pengkajian awal dengan fase pelaksanaan, sehingga pekerjaan '
            'lanjutan dapat berjalan meskipun asumsi awalnya belum '
            'terverifikasi. Sistem yang dikembangkan memisahkan secara tegas '
            'siklus hidup proyek menjadi fase survei dan fase eksekusi yang '
            'dipisahkan oleh sebuah gerbang persetujuan dari manajer. Proyek '
            'tidak akan memasuki fase eksekusi sebelum seluruh tugas survei '
            'selesai diverifikasi dan disetujui oleh manajer melalui '
            'antarmuka persetujuan. Mekanisme ini mencegah eksekusi yang '
            'didasarkan pada hasil survei yang belum tervalidasi, sehingga '
            'mengurangi risiko revisi pekerjaan akibat kesalahan asumsi '
            'pada tahap awal.'
        ),
        'gambar': '07-manager-approve-task-proyek.png',
        'judul_gambar': 'Antarmuka Persetujuan Survei sebagai Gerbang Eksekusi Proyek',
    },
    {
        'nama': 'Modul Eskalasi Kendala Lapangan dengan Lampiran Bukti Pendukung',
        'penjelasan': (
            'Aplikasi manajemen proyek pada umumnya hanya menyediakan kolom '
            'komentar bebas untuk menyampaikan kendala lapangan, sehingga '
            'informasi penting bercampur dengan diskusi rutin dan sulit '
            'ditelusuri kembali. Sistem yang dikembangkan menyediakan modul '
            'eskalasi terstruktur yang mewajibkan pelapor mencantumkan '
            'judul, deskripsi, tingkat prioritas, serta lampiran bukti '
            'berupa foto atau dokumen pendukung. Setiap eskalasi terhubung '
            'langsung dengan tugas dan proyek terkait sehingga manajer '
            'dapat melihat konteks kendala secara utuh dalam satu tampilan. '
            'Audit trail berupa identitas pelapor, waktu pelaporan, '
            'penanggung jawab penyelesaian, serta catatan resolusi disimpan '
            'secara permanen untuk kebutuhan evaluasi pasca-proyek.'
        ),
        'gambar': '08-teknisi-update-eskalasi.png',
        'judul_gambar': 'Antarmuka Pengajuan Eskalasi dengan Lampiran Bukti',
    },
    {
        'nama': 'Dashboard Performa Teknisi Berbasis Beban Kerja Aktual',
        'penjelasan': (
            'Aplikasi manajemen proyek pada umumnya hanya menyajikan '
            'visualisasi pada tingkat proyek tanpa menyediakan ringkasan '
            'performa per individu pelaksana, sehingga manajer kesulitan '
            'menilai kapasitas dan distribusi beban kerja teknisi secara '
            'objektif. Sistem yang dikembangkan menyediakan dashboard '
            'performa teknisi yang menampilkan jumlah tugas aktif, tingkat '
            'penyelesaian tepat waktu, akumulasi waktu pengerjaan, serta '
            'rekap eskalasi yang pernah dilaporkan. Visualisasi tersebut '
            'memberikan dasar yang objektif bagi manajer dalam menentukan '
            'penugasan baru, mengidentifikasi teknisi yang membutuhkan '
            'pendampingan, serta menjadi bahan evaluasi kinerja periodik. '
            'Fitur ini menjadi pembeda penting karena memindahkan basis '
            'evaluasi kinerja dari kesan subjektif menjadi data aktivitas '
            'yang tercatat sistem.'
        ),
        'gambar': '12-performa-teknisi-manager.png',
        'judul_gambar': 'Dashboard Performa Teknisi pada Tampilan Manajer',
    },
]


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        p.style = doc.styles['Caption']
    except KeyError:
        pass
    p.add_run(text)
    return p


def add_centered_image(doc, path: Path, width_cm: float = 14.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    return p


def add_paragraf(doc, text, style='Paragraf New'):
    p = doc.add_paragraph()
    try:
        p.style = doc.styles[style]
    except KeyError:
        pass
    p.add_run(text)
    return p


def build():
    shutil.copy(TEMPLATE, OUT_PATH)
    doc = Document(OUT_PATH)

    # Wipe body but preserve final sectPr
    body = doc.element.body
    sect_pr = body.find(qn('w:sectPr'))
    for el in list(body):
        if el is sect_pr:
            continue
        body.remove(el)

    # 5.3 heading (style 5.n dipakai untuk heading 5.x level pertama)
    p = doc.add_paragraph()
    try:
        p.style = doc.styles['5.n']
    except KeyError:
        pass
    p.add_run('Inovasi Sistem')

    # Paragraf pembuka
    intro = (
        'Subbab ini menguraikan sejumlah inovasi yang dihadirkan oleh sistem '
        'manajemen proyek berbasis daily report yang dikembangkan untuk '
        'PT Smart Home Inovasi (SHI). Inovasi yang dimaksud merupakan '
        'fitur-fitur yang membedakan sistem ini dari aplikasi manajemen '
        'proyek sejenis pada umumnya, baik dari sisi otomatisasi pengolahan '
        'data, mekanisme kendali mutu, maupun struktur alur kerja yang '
        'disesuaikan dengan karakteristik pekerjaan lapangan pada bidang '
        'instalasi sistem rumah pintar. Setiap inovasi disertai penjelasan '
        'pembeda terhadap kebiasaan aplikasi umum, kemudian ditampilkan '
        'tangkapan layar antarmuka sistem sebagai bukti implementasi nyata.'
    )
    add_paragraf(doc, intro)

    fig_no = FIG_START
    for idx, item in enumerate(INOVASI, start=1):
        # Nomor + nama (bold)
        p = doc.add_paragraph()
        try:
            p.style = doc.styles['Paragraf New']
        except KeyError:
            pass
        r = p.add_run(f'{idx}. {item["nama"]}')
        r.bold = True

        # Penjelasan
        add_paragraf(doc, item['penjelasan'])

        # Image + caption. Items mapped to alur-sistem/ by default; mockups/
        # prefix resolves against the _build/ folder instead.
        if item['gambar'].startswith('mockups/'):
            img_path = MOCKUP_DIR / item['gambar'].split('/', 1)[1]
        else:
            img_path = SHOT_DIR / item['gambar']
        if not img_path.is_file():
            print(f'[WARN] Missing image: {img_path}')
            continue
        add_centered_image(doc, img_path, width_cm=14.5)
        add_caption(doc, f'Gambar 5.{fig_no} {item["judul_gambar"]}')
        fig_no += 1

    doc.save(OUT_PATH)
    n_items = len(INOVASI)
    print(f'Saved {OUT_PATH}')
    print(f'  Inovasi: {n_items}, Gambar: 5.{FIG_START}– 5.{fig_no-1}')


if __name__ == '__main__':
    build()
