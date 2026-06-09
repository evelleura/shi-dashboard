"""
Generate `bab5_final/5.2.3 Pengujian Black Box.docx`.

Struktur (mirror 5.2.1 — narasi + bukti + caption + kesimpulan), namun
kali ini bukti merupakan respons API hidup pada localhost:3000 untuk
SKENARIO VALID (success path) dan SKENARIO TIDAK VALID (validation /
otorisasi gagal).

Empat bagian uji, 11 skenario (4 Valid + 7 Tidak Valid):
  1. Autentikasi              — login valid, email tidak terdaftar,
                                password salah
  2. Manajemen Proyek         — create valid, tanpa nama (validasi),
                                tanggal akhir < mulai (validasi)
  3. Daily Report Teknisi     — catat aktivitas valid, lintas-teknisi
                                (otorisasi), promote ke done (review gate)
  4. Dashboard EWS            — akses manajer (valid), akses teknisi
                                (role check)

Untuk setiap skenario:
  - Permintaan HTTP (method, path, payload) → hit ke localhost:3000
  - Respons (status code + body) di-capture
  - Dirender PNG panel ala "HTTP response viewer"
  - Bagian tabel ringkas mencatat: No, Skenario, Hasil yang Diharapkan,
    Hasil Pengujian, Kesimpulan
  - Lalu screenshot bukti + caption

Pre-req:
  - Dev server hidup di http://localhost:3000
  - Akun seed budi@shi.co.id (manager) dan rizky@shi.co.id (technician)
"""

from __future__ import annotations

import importlib.util
import json
import shutil
import sys
from pathlib import Path
from typing import Optional

import requests
from PIL import Image, ImageDraw

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT      = Path(__file__).resolve().parents[3]
TEMPLATE  = ROOT / 'naskah/bab_split/BAB_V_Implementasi_dan_Pembahasan_Sistem.docx'
OUT_PATH  = ROOT / 'naskah/bab5_final/5.2.3 Pengujian Black Box.docx'
SHOTS_DIR = ROOT / 'naskah/bab5_final/screenshots_pengujian'
SHOTS_DIR.mkdir(parents=True, exist_ok=True)

HERE = Path(__file__).resolve().parent
sys.path.insert(0, str(HERE))

# Reuse font + color konstanta dari render screenshots
spec = importlib.util.spec_from_file_location('rsh', HERE / '03_render_screenshots.py')
rsh  = importlib.util.module_from_spec(spec)
spec.loader.exec_module(rsh)

BASE = 'http://localhost:3000'

# Login token disediakan oleh helper di bawah agar selalu fresh.
MANAGER_CRED = {'email': 'budi@shi.co.id',  'password': 'password123'}
TECH_CRED    = {'email': 'rizky@shi.co.id', 'password': 'password123'}


# ── Spec skenario pengujian ──────────────────────────────────────────────

# Setiap entri: dict dengan field
#   no, kategori (valid|tidak_valid), nama, narasi, expected,
#   request: dict {method, path, role (manager|technician|none), body, type ('json'|'form')}

BAGIAN_1_AUTH = {
    'judul': 'Pengujian Modul Autentikasi',
    'narasi': (
        'Modul autentikasi diuji dengan kombinasi tiga skenario: '
        'satu skenario valid yang merepresentasikan alur login normal, '
        'serta dua skenario tidak valid yang merepresentasikan upaya '
        'login dengan kredensial salah. Pengujian dilakukan dengan '
        'mengirim permintaan POST ke endpoint /api/auth/login pada '
        'aplikasi yang sedang berjalan.'
    ),
    'cases': [
        {
            'no': 1, 'kategori': 'valid',
            'nama': 'Login dengan kredensial valid',
            'narasi': (
                'Skenario login yang berhasil. Permintaan POST '
                '/api/auth/login dikirim dengan email dan password '
                'pengguna manajer yang terdaftar.'
            ),
            'expected': (
                'Sistem mengembalikan status 200 OK beserta JSON '
                'berisi token JWT dan profil pengguna.'
            ),
            'request': {
                'method': 'POST', 'path': '/api/auth/login', 'role': None,
                'body': MANAGER_CRED, 'type': 'json',
            },
        },
        {
            'no': 2, 'kategori': 'tidak_valid',
            'nama': 'Login dengan email tidak terdaftar',
            'narasi': (
                'Skenario login dengan email yang tidak ada pada tabel '
                'tb_user. Sistem harus menolak permintaan dengan pesan '
                'error generik agar tidak membocorkan informasi '
                'keberadaan akun.'
            ),
            'expected': (
                'Sistem mengembalikan status 401 Unauthorized dengan '
                'pesan error "Invalid email or password".'
            ),
            'request': {
                'method': 'POST', 'path': '/api/auth/login', 'role': None,
                'body': {'email': 'tidak.terdaftar@shi.co.id',
                         'password': 'apapun'},
                'type': 'json',
            },
        },
        {
            'no': 3, 'kategori': 'tidak_valid',
            'nama': 'Login dengan password salah',
            'narasi': (
                'Skenario login dengan email valid tetapi password '
                'salah. Pesan error harus sama persis dengan kasus '
                'email tidak terdaftar (uniform error) untuk mencegah '
                'serangan enumerasi akun.'
            ),
            'expected': (
                'Sistem mengembalikan status 401 Unauthorized dengan '
                'pesan error yang identik.'
            ),
            'request': {
                'method': 'POST', 'path': '/api/auth/login', 'role': None,
                'body': {'email': 'budi@shi.co.id', 'password': 'salah123'},
                'type': 'json',
            },
        },
    ],
}

BAGIAN_2_PROYEK = {
    'judul': 'Pengujian Modul Manajemen Proyek',
    'narasi': (
        'Modul manajemen proyek diuji dengan tiga skenario yang menyasar '
        'endpoint POST /api/projects. Skenario valid memverifikasi alur '
        'penambahan proyek yang sah, sedangkan dua skenario tidak valid '
        'memverifikasi penolakan terhadap input yang melanggar aturan '
        'validasi.'
    ),
    'cases': [
        {
            'no': 1, 'kategori': 'valid',
            'nama': 'Tambah proyek dengan payload lengkap',
            'narasi': (
                'Skenario penambahan proyek baru oleh manajer dengan '
                'seluruh field wajib terisi dan tanggal akhir setelah '
                'tanggal mulai.'
            ),
            'expected': (
                'Sistem mengembalikan status 201 Created beserta '
                'objek proyek yang baru terbentuk.'
            ),
            'request': {
                'method': 'POST', 'path': '/api/projects', 'role': 'manager',
                'body': {
                    'name': 'Proyek Uji Black Box',
                    'description': 'Pemasangan demo skripsi',
                    'client_id': 1,
                    'start_date': '2026-06-01',
                    'end_date': '2026-08-01',
                    'category': 'instalasi',
                    'project_value': 25000000,
                },
                'type': 'json',
            },
        },
        {
            'no': 2, 'kategori': 'tidak_valid',
            'nama': 'Tambah proyek tanpa nama (field wajib)',
            'narasi': (
                'Skenario penambahan proyek tanpa mengisi field name. '
                'Sistem harus menolak permintaan pada lapisan validasi '
                'sebelum query INSERT dieksekusi.'
            ),
            'expected': (
                'Sistem mengembalikan status 400 Bad Request dengan '
                'pesan error yang menyebutkan field yang hilang.'
            ),
            'request': {
                'method': 'POST', 'path': '/api/projects', 'role': 'manager',
                'body': {
                    'description': 'Tanpa nama proyek',
                    'client_id': 1,
                    'start_date': '2026-06-01',
                    'end_date': '2026-08-01',
                    'category': 'instalasi',
                },
                'type': 'json',
            },
        },
        {
            'no': 3, 'kategori': 'tidak_valid',
            'nama': 'Tambah proyek dengan tanggal akhir < mulai',
            'narasi': (
                'Skenario penambahan proyek dengan tanggal akhir lebih '
                'awal daripada tanggal mulai. Sistem harus menolak '
                'permintaan karena rentang waktu proyek tidak masuk '
                'akal.'
            ),
            'expected': (
                'Sistem mengembalikan status 400 Bad Request dengan '
                'pesan error mengenai validasi tanggal.'
            ),
            'request': {
                'method': 'POST', 'path': '/api/projects', 'role': 'manager',
                'body': {
                    'name': 'Proyek Salah Tanggal',
                    'description': 'Tanggal akhir lebih awal',
                    'client_id': 1,
                    'start_date': '2026-08-01',
                    'end_date':   '2026-06-01',
                    'category': 'instalasi',
                },
                'type': 'json',
            },
        },
    ],
}

BAGIAN_3_DAILY_REPORT = {
    'judul': 'Pengujian Modul Daily Report Teknisi',
    'narasi': (
        'Modul daily report (catatan aktivitas + perubahan status tugas) '
        'diuji dengan tiga skenario untuk memverifikasi mekanisme '
        'otorisasi per teknisi serta gerbang review oleh manajer.'
    ),
    'cases': [
        {
            'no': 1, 'kategori': 'valid',
            'nama': 'Teknisi mengubah status tugas miliknya',
            'narasi': (
                'Teknisi mengubah status sebuah tugas yang ditugaskan '
                'kepadanya dari to_do menjadi in_progress. Transisi '
                'status ini memicu rekalkulasi SPI pada Bagian 2 (lihat '
                '5.1.2 dan 5.2.1).'
            ),
            'expected': (
                'Sistem mengembalikan status 200 OK dengan objek tugas '
                'yang status-nya telah diperbarui.'
            ),
            'request': {
                'method': 'PATCH', 'path': '/api/tasks/{TASK_OWNED}/status',
                'role': 'technician',
                'body': {'status': 'in_progress'}, 'type': 'json',
            },
        },
        {
            'no': 2, 'kategori': 'tidak_valid',
            'nama': 'Teknisi mengubah status tugas teknisi lain',
            'narasi': (
                'Teknisi mencoba mengubah status sebuah tugas yang '
                'ditugaskan kepada teknisi LAIN. Sistem harus menolak '
                'karena melanggar otorisasi kepemilikan tugas.'
            ),
            'expected': (
                'Sistem mengembalikan status 403 Forbidden dengan '
                'pesan error tentang otorisasi.'
            ),
            'request': {
                'method': 'PATCH', 'path': '/api/tasks/{TASK_OTHER}/status',
                'role': 'technician',
                'body': {'status': 'in_progress'}, 'type': 'json',
            },
        },
        {
            'no': 3, 'kategori': 'tidak_valid',
            'nama': 'Teknisi menandai tugas sebagai done (review gate)',
            'narasi': (
                'Teknisi mencoba mengubah status tugas miliknya '
                'langsung ke done. Sistem harus menolak karena '
                'kewenangan menandai tugas selesai sepenuhnya berada '
                'pada manajer (review gate).'
            ),
            'expected': (
                'Sistem mengembalikan status 400 Bad Request dengan '
                'pesan error transisi status yang tidak diizinkan.'
            ),
            'request': {
                'method': 'PATCH', 'path': '/api/tasks/{TASK_OWNED}/status',
                'role': 'technician',
                'body': {'status': 'done'}, 'type': 'json',
            },
        },
    ],
}

BAGIAN_4_DASHBOARD = {
    'judul': 'Pengujian Akses Dashboard Early Warning System',
    'narasi': (
        'Dashboard EWS diuji untuk memverifikasi kontrol akses '
        'berbasis peran. Endpoint /api/dashboard hanya boleh diakses '
        'oleh manajer dan admin; teknisi harus diarahkan ke endpoint '
        'dashboard teknisi.'
    ),
    'cases': [
        {
            'no': 1, 'kategori': 'valid',
            'nama': 'Manajer mengakses dashboard EWS',
            'narasi': (
                'Manajer mengirim permintaan GET /api/dashboard untuk '
                'memuat data dashboard EWS.'
            ),
            'expected': (
                'Sistem mengembalikan status 200 OK dengan struktur '
                '{projects, summary, recent_activity}.'
            ),
            'request': {
                'method': 'GET', 'path': '/api/dashboard',
                'role': 'manager', 'body': None, 'type': 'json',
            },
        },
        {
            'no': 2, 'kategori': 'tidak_valid',
            'nama': 'Teknisi mengakses dashboard EWS',
            'narasi': (
                'Teknisi mencoba mengakses endpoint manajer '
                '/api/dashboard. Sistem harus menolak dengan '
                'role-based access control.'
            ),
            'expected': (
                'Sistem mengembalikan status 403 Forbidden dengan '
                'pesan error tentang akses.'
            ),
            'request': {
                'method': 'GET', 'path': '/api/dashboard',
                'role': 'technician', 'body': None, 'type': 'json',
            },
        },
    ],
}

ALL_BAGIAN = [
    ('1.', BAGIAN_1_AUTH,           'auth'),
    ('2.', BAGIAN_2_PROYEK,         'proyek'),
    ('3.', BAGIAN_3_DAILY_REPORT,   'daily_report'),
    ('4.', BAGIAN_4_DASHBOARD,      'dashboard'),
]


# ── HTTP execution ────────────────────────────────────────────────────────

class Tokens:
    def __init__(self):
        self.manager_resp = self._login(MANAGER_CRED)
        self.tech_resp    = self._login(TECH_CRED)
        self.manager      = self.manager_resp['token']
        self.tech         = self.tech_resp['token']
        self.tech_user_id = self.tech_resp['user']['id']
        # Locate one task owned by Rizky and one task owned by another technician
        self.task_owned, self.task_other = self._discover_tasks()

    @staticmethod
    def _login(cred):
        r = requests.post(f'{BASE}/api/auth/login', json=cred, timeout=15)
        r.raise_for_status()
        return r.json()['data']

    def _discover_tasks(self):
        """Find one task assigned to Rizky and one to a different user.
        Both must be project status='active' and task status='to_do'."""
        import psycopg2
        conn = psycopg2.connect('host=127.0.0.1 port=5432 '
                                'dbname=shi_dashboard_new user=postgres '
                                'password=12345')
        cur = conn.cursor()
        cur.execute(
            "SELECT t.id FROM tasks t JOIN projects p ON p.id = t.project_id "
            "WHERE t.assigned_to = %s AND t.status = 'to_do' "
            "AND p.status = 'active' LIMIT 1",
            (self.tech_user_id,))
        owned = cur.fetchone()
        cur.execute(
            "SELECT t.id FROM tasks t JOIN projects p ON p.id = t.project_id "
            "WHERE t.assigned_to <> %s AND t.assigned_to IS NOT NULL "
            "AND t.status = 'to_do' AND p.status = 'active' LIMIT 1",
            (self.tech_user_id,))
        other = cur.fetchone()
        conn.close()
        return (owned[0] if owned else None,
                other[0] if other else None)


def make_request(case, tokens: Tokens):
    """Resolve placeholders → execute → return (method, path, req_body, status, resp_body)."""
    req = case['request']
    method = req['method']
    path   = req['path']
    # Substitute task placeholders
    path = (path.replace('{TASK_OWNED}', str(tokens.task_owned))
                .replace('{TASK_OTHER}', str(tokens.task_other)))

    headers = {}
    if req['role'] == 'manager':
        headers['Authorization'] = f'Bearer {tokens.manager}'
    elif req['role'] == 'technician':
        headers['Authorization'] = f'Bearer {tokens.tech}'

    url = f'{BASE}{path}'
    body = req['body']
    try:
        if method == 'GET':
            r = requests.get(url, headers=headers, timeout=20)
        elif body is None:
            r = requests.request(method, url, headers=headers, timeout=20)
        else:
            r = requests.request(method, url, headers=headers, json=body,
                                  timeout=20)
        status = r.status_code
        try:
            resp_body = r.json()
        except ValueError:
            resp_body = {'_raw': r.text[:500]}
    except Exception as e:
        status = 0
        resp_body = {'_error': f'{type(e).__name__}: {e}'}
    return method, path, body, status, resp_body


# ── Render PNG (HTTP response viewer style) ───────────────────────────────

def _status_color(status: int) -> tuple[int, int, int]:
    if 200 <= status < 300:
        return (39, 174, 96)     # green
    if 300 <= status < 400:
        return (52, 152, 219)    # blue
    if 400 <= status < 500:
        return (230, 126, 34)    # orange
    return (231, 76, 60)         # red (5xx / 0)


def _wrap_json(obj, max_w=68):
    text = json.dumps(obj, indent=2, ensure_ascii=False)
    out_lines = []
    for ln in text.split('\n'):
        while len(ln) > max_w:
            out_lines.append(ln[:max_w])
            ln = '    ' + ln[max_w:]
        out_lines.append(ln)
    return out_lines


def render_http_panel(method: str, path: str, req_body, status: int,
                      resp_body, out_path: Path, W: int = 1100):
    """Render satu PNG yang menampilkan: chrome (window) → request line +
    status badge → request body → response body."""
    H_TITLE = 32
    H_TAB   = 30
    H_BAR   = 34       # request line bar (method + path + status)
    H_HDR_SEC = 24     # 'Request' / 'Response' header strip
    LINE_H  = 20
    PAD     = 12

    font_mono = rsh.F_MONO_S
    font_ui   = rsh.F_TITLE
    font_ui_b = rsh.F_UI_B

    req_lines  = _wrap_json(req_body if req_body is not None else {})
    resp_lines = _wrap_json(resp_body)
    body_h = H_HDR_SEC + (len(req_lines) + 1) * LINE_H + 8 \
           + H_HDR_SEC + (len(resp_lines) + 1) * LINE_H + 8
    H = H_TITLE + H_TAB + H_BAR + body_h + 36

    canvas = Image.new('RGB', (W, H), '#ffffff')
    d = ImageDraw.Draw(canvas)

    # Chrome
    d.rectangle([0, 0, W, H_TITLE], fill=rsh.C_TITLEBAR)
    d.rectangle([0, H_TITLE - 1, W, H_TITLE], fill=rsh.C_BORDER)
    d.text((14, 9), f'HTTP Client — Pengujian Endpoint',
           font=font_ui, fill=(80, 88, 104))
    for i, c in enumerate([(255, 95, 86), (255, 189, 46), (39, 201, 63)]):
        d.ellipse([W - 78 + i * 22, 10, W - 66 + i * 22, 22], fill=c)
    tab_y0 = H_TITLE
    d.rectangle([0, tab_y0, W, tab_y0 + H_TAB], fill=rsh.C_TAB)
    d.rectangle([12, tab_y0 + 4, 320, tab_y0 + H_TAB],
                fill=rsh.C_TAB_ACT, outline=rsh.C_BORDER)
    d.text((22, tab_y0 + 8), f'{method} {path}',
           font=font_ui, fill=rsh.C_TEXT)
    d.rectangle([0, tab_y0 + H_TAB - 1, W, tab_y0 + H_TAB], fill=rsh.C_BORDER)

    # Status bar
    y = H_TITLE + H_TAB
    d.rectangle([0, y, W, y + H_BAR], fill='#F4F6FA')
    sc = _status_color(status)
    badge_w = 90
    badge_h = 22
    badge_x = W - badge_w - 14
    badge_y = y + (H_BAR - badge_h) // 2
    d.rounded_rectangle([badge_x, badge_y, badge_x + badge_w,
                          badge_y + badge_h], radius=4, fill=sc)
    status_text = f'{status}' if status else 'NETWORK'
    bbox = d.textbbox((0, 0), status_text, font=font_ui_b)
    bw = bbox[2] - bbox[0]
    d.text((badge_x + (badge_w - bw) // 2, badge_y + 3),
           status_text, font=font_ui_b, fill='white')
    d.text((PAD + 4, y + 9), f'{method}  {path}',
           font=font_ui_b, fill=rsh.C_TEXT)
    d.rectangle([0, y + H_BAR - 1, W, y + H_BAR], fill=rsh.C_BORDER)

    y += H_BAR

    # Request section
    d.rectangle([0, y, W, y + H_HDR_SEC], fill=rsh.C_HEADER)
    d.text((PAD + 4, y + 5), 'REQUEST BODY',
           font=font_ui, fill=rsh.C_HEADER_T)
    y += H_HDR_SEC
    for ln in req_lines:
        d.text((PAD + 8, y + 2), ln, font=font_mono, fill=rsh.C_TEXT)
        y += LINE_H
    y += 8

    # Response section
    d.rectangle([0, y, W, y + H_HDR_SEC], fill=rsh.C_HEADER)
    d.text((PAD + 4, y + 5), 'RESPONSE BODY',
           font=font_ui, fill=rsh.C_HEADER_T)
    y += H_HDR_SEC
    for ln in resp_lines:
        d.text((PAD + 8, y + 2), ln, font=font_mono, fill=rsh.C_TEXT)
        y += LINE_H
    y += 8

    # Border
    d.rectangle([0, 0, W - 1, H - 1], outline=rsh.C_BORDER)
    canvas.save(out_path, 'PNG', optimize=True)


# ── docx helpers ──────────────────────────────────────────────────────────

def add_paragraf(doc, text, style='Paragraf'):
    p = doc.add_paragraph()
    try:
        p.style = doc.styles[style]
    except KeyError:
        pass
    p.add_run(text)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        p.style = doc.styles['Caption']
    except KeyError:
        pass
    p.add_run(text)
    return p


def add_centered_image(doc, path: Path, width_cm: float = 14.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    return p


def _set_cell_shading(cell, fill='F1F5F9'):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def add_test_table(doc, cases_summary):
    """Add a summary table of test results for one bagian.
       cases_summary: list of dict with keys
         no, skenario, expected, actual, kesimpulan
    """
    headers = ['No', 'Skenario Pengujian', 'Hasil yang Diharapkan',
               'Hasil Pengujian', 'Kesimpulan']
    table = doc.add_table(rows=1 + len(cases_summary), cols=len(headers))
    table.style = 'Table Grid'
    table.autofit = False

    widths = [Cm(0.8), Cm(4.5), Cm(4.5), Cm(4.5), Cm(1.5)]
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Header row
    hdr = table.rows[0]
    for cell, h in zip(hdr.cells, headers):
        _set_cell_shading(cell, 'E2E8F0')
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    # Data rows
    for ri, case in enumerate(cases_summary, start=1):
        row = table.rows[ri]
        cells = row.cells
        cells[0].paragraphs[0].add_run(str(case['no']))
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[1].paragraphs[0].add_run(case['skenario'])
        cells[2].paragraphs[0].add_run(case['expected'])
        cells[3].paragraphs[0].add_run(case['actual'])
        para = cells[4].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(case['kesimpulan'])
        run.bold = True
        kes_lower = case['kesimpulan'].lower()
        if kes_lower == 'valid':
            run.font.color.rgb = RGBColor(0x05, 0x96, 0x69)   # hijau
        elif kes_lower == 'tidak valid':
            run.font.color.rgb = RGBColor(0xC2, 0x41, 0x0C)   # oranye
        else:  # GAGAL
            run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)   # merah
        for c in cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)


# ── Main ──────────────────────────────────────────────────────────────────

def build():
    # 1) Login + discover dynamic IDs
    tokens = Tokens()
    print(f'[*] manager_token len={len(tokens.manager)}')
    print(f'[*] tech_token    len={len(tokens.tech)}, user_id={tokens.tech_user_id}')
    print(f'[*] task_owned={tokens.task_owned}, task_other={tokens.task_other}')

    # 2) Execute each case, render PNG bukti
    rendered = {}
    for cat_num, bagian, key in ALL_BAGIAN:
        print(f"\n── Bagian {cat_num} {bagian['judul']} ──")
        per_case = []
        for case in bagian['cases']:
            method, path, body, status, resp = make_request(case, tokens)
            png = SHOTS_DIR / f"{key}_{case['no']:02d}_{case['kategori']}.png"
            render_http_panel(method, path, body, status, resp, png)
            # Determine actual summary text from response
            actual_summary = _summarize_response(status, resp)
            print(f"  [{case['kategori']:11}] {case['nama']:55} → {status}")
            per_case.append({
                'case': case,
                'method': method,
                'path': path,
                'status': status,
                'resp': resp,
                'png': png,
                'actual_summary': actual_summary,
            })
        rendered[key] = per_case

    # 3) Compose docx
    shutil.copy(TEMPLATE, OUT_PATH)
    doc = Document(OUT_PATH)
    body = doc.element.body
    sect_pr = body.find(qn('w:sectPr'))
    for el in list(body):
        if el is sect_pr:
            continue
        body.remove(el)

    # Heading 5.2.3
    p = doc.add_paragraph()
    try:
        p.style = doc.styles['5.2.n']
    except KeyError:
        pass
    p.add_run('Pengujian Black Box')

    add_paragraf(doc,
        'Subbab ini menyajikan hasil pengujian sistem dengan metode '
        'black box, yaitu pengujian yang berfokus pada perilaku '
        'masukan-keluaran tanpa memperhatikan struktur kode internal. '
        'Pengujian dilakukan langsung pada aplikasi yang berjalan di '
        'lingkungan pengembangan dengan mengirim permintaan HTTP ke '
        'lapisan API Next.js, kemudian membandingkan respons sistem '
        'terhadap hasil yang diharapkan. Setiap skenario dilengkapi '
        'dengan kesimpulan apakah sistem berperilaku sesuai harapan '
        '(Valid) atau tidak sesuai harapan (Tidak Valid).'
    )
    add_paragraf(doc,
        'Empat modul diuji secara berurutan: (1) autentikasi pengguna, '
        '(2) manajemen proyek oleh manajer, (3) pencatatan daily '
        'report oleh teknisi, dan (4) kontrol akses berbasis peran '
        'pada Dashboard Early Warning System (EWS). Setiap modul '
        'mencakup satu skenario sukses (alur normal) dan beberapa '
        'skenario penolakan untuk memverifikasi mekanisme validasi '
        'dan otorisasi. Pesan kategori "Tidak Valid" pada kesimpulan '
        'merepresentasikan kasus di mana sistem MENOLAK input atau '
        'akses sesuai dengan aturan bisnis yang dirancang—bukan '
        'kegagalan sistem.'
    )

    fig_no = 1
    table_no = 1
    for cat_num, bagian, key in ALL_BAGIAN:
        # Header bagian
        th = doc.add_paragraph()
        try:
            th.style = doc.styles['Paragraf']
        except KeyError:
            pass
        run = th.add_run(f"{cat_num}  {bagian['judul']}")
        run.bold = True

        # Narasi
        add_paragraf(doc, bagian['narasi'])

        # Tabel ringkas
        cases_summary = []
        for entry in rendered[key]:
            case = entry['case']
            kes = _kesimpulan_label(case, entry['status'])
            cases_summary.append({
                'no': case['no'],
                'skenario': case['nama'],
                'expected': case['expected'],
                'actual': entry['actual_summary'],
                'kesimpulan': kes,
            })
        add_test_table(doc, cases_summary)
        add_caption(doc, f'Tabel 5.{table_no} Hasil pengujian — {bagian["judul"]}')
        table_no += 1

        # Per kasus: narasi singkat → screenshot bukti → caption
        for entry in rendered[key]:
            case = entry['case']
            kes_label = _kesimpulan_label(case, entry['status'])

            # Narasi per kasus
            np = doc.add_paragraph()
            try:
                np.style = doc.styles['Paragraf']
            except KeyError:
                pass
            sk_run = np.add_run(f'Kasus {case["no"]} — {case["nama"]}: ')
            sk_run.bold = True
            np.add_run(
                f'{case["narasi"]} {case["expected"]} Hasil pengujian '
                f'terhadap endpoint {entry["method"]} {entry["path"]} '
                f'menghasilkan status HTTP {entry["status"]} dengan '
                f'kesimpulan '
            )
            kes_run = np.add_run(kes_label)
            kes_run.bold = True
            if kes_label == 'Valid':
                kes_run.font.color.rgb = RGBColor(0x05, 0x96, 0x69)
            elif kes_label == 'Tidak Valid':
                kes_run.font.color.rgb = RGBColor(0xC2, 0x41, 0x0C)
            else:
                kes_run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
            np.add_run(
                f'. Bukti tangkapan respons sistem disajikan pada '
                f'Gambar 5.{fig_no}.'
            )

            # Screenshot
            add_centered_image(doc, entry['png'], width_cm=14.5)
            add_caption(doc, f'Gambar 5.{fig_no} Respons {entry["method"]} '
                             f'{entry["path"]} — {kes_label}')
            fig_no += 1

    doc.save(OUT_PATH)
    print(f'\nSaved {OUT_PATH}')
    print(f'  Bagian: {len(ALL_BAGIAN)} • Tabel: 1-{table_no - 1} '
          f'• Gambar: 5.1 – 5.{fig_no - 1}')


def _summarize_response(status: int, resp) -> str:
    """Ringkas respons ke 1-2 baris untuk dimasukkan ke kolom Hasil
    Pengujian pada tabel."""
    if isinstance(resp, dict):
        if 'success' in resp:
            if resp.get('success'):
                return f'HTTP {status} • success=true'
            err = resp.get('error') or resp.get('message') or 'tanpa pesan'
            return f'HTTP {status} • error="{err}"'
        if '_error' in resp:
            return f'NETWORK ERROR: {resp["_error"]}'
    return f'HTTP {status} • body diterima'


def _is_test_pass(case, status: int) -> bool:
    """True jika sistem berperilaku sesuai harapan kategori skenario:
       - valid     → 2xx (sistem menerima & memproses)
       - tidak_valid → 4xx (sistem menolak dengan benar)"""
    if case['kategori'] == 'valid':
        return 200 <= status < 300
    return 400 <= status < 500


def _kesimpulan_label(case, status: int) -> str:
    """Label kolom 'Kesimpulan' pada tabel pengujian.
       Konvensi skripsi: label kategori skenario (Valid / Tidak Valid).
       Jika sistem TIDAK berperilaku sesuai harapan kategori, label
       berubah menjadi 'GAGAL' untuk menandai bug yang perlu ditindaklanjuti."""
    if not _is_test_pass(case, status):
        return 'GAGAL'
    return 'Valid' if case['kategori'] == 'valid' else 'Tidak Valid'


if __name__ == '__main__':
    build()
