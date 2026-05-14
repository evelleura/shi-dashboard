"""
Generate `bab5_final/5.2.1 Pembahasan Basis Data.docx`.

Structure:
  5.2.1 Pembahasan Basis Data
  (opening paragraph)
  For each of 7 tables:
    Tabel <Nama>
    For each CRUD op (INSERT, SELECT, UPDATE, DELETE):
      <narrative paragraph>
      [SQL screenshot]      caption: Gambar 5.x Query <OP> tb_xxx
      [Result screenshot]   caption: Gambar 5.(x+1) Hasil <OP> tb_xxx
"""

import shutil
import importlib.util
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT      = Path(__file__).resolve().parents[3]
TEMPLATE  = ROOT / 'naskah/bab_split/BAB_V_Implementasi_dan_Pembahasan_Sistem.docx'
OUT_PATH  = ROOT / 'naskah/bab5_final/5.2.1 Pembahasan Basis Data.docx'
SHOTS_DIR = ROOT / 'naskah/bab5_final/screenshots'

# load CRUD spec
spec = importlib.util.spec_from_file_location("crud", Path(__file__).parent / "02_crud_queries.py")
crud = importlib.util.module_from_spec(spec); spec.loader.exec_module(crud)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        p.style = doc.styles['Caption']
    except KeyError:
        pass
    p.add_run(text)
    return p


def add_centered_image(doc, path: Path, width_cm: float = 15.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    return p


def build():
    shutil.copy(TEMPLATE, OUT_PATH)
    doc = Document(OUT_PATH)
    body = doc.element.body
    sect_pr = body.find(qn('w:sectPr'))
    for el in list(body):
        if el is sect_pr:
            continue
        body.remove(el)

    # Heading: 5.2.1
    p = doc.add_paragraph()
    try:
        p.style = doc.styles['5.2.n']
    except KeyError:
        pass
    p.add_run('Pembahasan Basis Data')

    # Opening paragraph
    intro = doc.add_paragraph()
    try:
        intro.style = doc.styles['Paragraf']
    except KeyError:
        pass
    intro.add_run(
        "Pada subbab ini disajikan pembahasan mengenai implementasi fisik "
        "basis data yang telah dirancang pada BAB IV untuk mendukung "
        "seluruh aktivitas operasional sistem manajemen proyek di "
        "PT Smart Home Inovasi (SHI). Basis data diimplementasikan di "
        "atas PostgreSQL 17 dan diakses melalui pustaka pg dengan "
        "parameterized query untuk mencegah serangan SQL injection. "
        "Setiap tabel inti (tb_user, tb_klien, tb_proyek, "
        "tb_penugasan_proyek, tb_tugas, tb_bukti, dan tb_eskalasi) "
        "dibahas mengikuti pola siklus operasi CRUD (Create, Read, "
        "Update, Delete). Untuk setiap operasi, ditampilkan sintaks "
        "query yang dieksekusi pada PostgreSQL beserta hasil eksekusi "
        "sebagai bukti bahwa perintah berhasil dijalankan dengan benar."
    )

    fig_no = 1
    for table, ops, label in crud.ALL_TABLES:
        # Table header
        th = doc.add_paragraph()
        try:
            th.style = doc.styles['Paragraf']
        except KeyError:
            pass
        run = th.add_run(label)
        run.bold = True

        for op_idx, entry in enumerate(ops, start=1):
            op   = entry['op']
            narr = entry['narr']

            # Narrative
            np = doc.add_paragraph()
            try:
                np.style = doc.styles['Paragraf']
            except KeyError:
                pass
            np.add_run(
                narr + " "
                f"Eksekusi query disajikan pada Gambar 5.{fig_no} "
                f"dan hasil eksekusinya pada Gambar 5.{fig_no + 1}."
            )

            # SQL screenshot
            sql_png = SHOTS_DIR / f"{table}_{op_idx:02d}_{op.lower()}_sql.png"
            add_centered_image(doc, sql_png, width_cm=14.5)
            add_caption(doc, f"Gambar 5.{fig_no} Query {op} pada {table}")
            fig_no += 1

            # Result screenshot
            res_png = SHOTS_DIR / f"{table}_{op_idx:02d}_{op.lower()}_result.png"
            add_centered_image(doc, res_png, width_cm=14.5)
            add_caption(doc, f"Gambar 5.{fig_no} Hasil eksekusi {op} pada {table}")
            fig_no += 1

    doc.save(OUT_PATH)
    print(f"Saved {OUT_PATH}  •  figures: {fig_no - 1}")


if __name__ == '__main__':
    build()
