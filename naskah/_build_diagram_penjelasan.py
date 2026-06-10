"""
Generator dua file penjelasan diagram:
  1) naskah/erd shi.docx           -- penjelasan diagram/erd_shi.drawio
  2) naskah/class diagram shi.docx -- penjelasan diagram/Class/class_shi.drawio

Style mengikuti naskah/class diagram.docx (4.3.1.n, Paragraf New, Daftar Tabel).
"""
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn

TEMPLATE = 'naskah/class diagram.docx'
OUT_ERD = 'naskah/erd shi.docx'
OUT_CLASS = 'naskah/class diagram shi.docx'

STYLE_HEADING = '4.3.1.n'
STYLE_PARA = 'Paragraf New'
STYLE_CAPTION = 'Daftar Tabel'


def reset(doc):
    body = doc.element.body
    for child in list(body):
        if child.tag == qn('w:sectPr'):
            continue
        body.remove(child)


def H(doc, txt):
    doc.add_paragraph(txt, style=STYLE_HEADING)


def P(doc, txt):
    doc.add_paragraph(txt, style=STYLE_PARA)


def CAP(doc, txt):
    doc.add_paragraph(txt, style=STYLE_CAPTION)


def tabel_fungsi(doc, rows):
    """rows = [(fungsi, deskripsi), ...] -> tabel 3 kolom: No. | Fungsi | Deskripsi"""
    t = doc.add_table(rows=1 + len(rows), cols=3)
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    hdr[0].text = 'No.'
    hdr[1].text = 'Fungsi'
    hdr[2].text = 'Deskripsi'
    for i, (f, d) in enumerate(rows, start=1):
        c = t.rows[i].cells
        c[0].text = f'{i}.'
        c[1].text = f
        c[2].text = d


def tabel_atribut(doc, rows):
    """rows = [(atribut, tipe, keterangan), ...]"""
    t = doc.add_table(rows=1 + len(rows), cols=4)
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    hdr[0].text = 'No.'
    hdr[1].text = 'Atribut'
    hdr[2].text = 'Tipe'
    hdr[3].text = 'Keterangan'
    for i, (a, ti, k) in enumerate(rows, start=1):
        c = t.rows[i].cells
        c[0].text = f'{i}.'
        c[1].text = a
        c[2].text = ti
        c[3].text = k


def tabel_relasi(doc, rows):
    """rows = [(entitas_a, kardinalitas, entitas_b, deskripsi), ...]"""
    t = doc.add_table(rows=1 + len(rows), cols=5)
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    hdr[0].text = 'No.'
    hdr[1].text = 'Entitas A'
    hdr[2].text = 'Kardinalitas'
    hdr[3].text = 'Entitas B'
    hdr[4].text = 'Keterangan'
    for i, (a, k, b, d) in enumerate(rows, start=1):
        c = t.rows[i].cells
        c[0].text = f'{i}.'
        c[1].text = a
        c[2].text = k
        c[3].text = b
        c[4].text = d


# ============================================================================
# ERD CONTENT
# ============================================================================

def build_erd():
    doc = Document(TEMPLATE)
    reset(doc)

    H(doc, 'Model Entity Relationship Diagram')
    P(doc,
      'Entity Relationship Diagram (ERD) merupakan model konseptual yang '
      'digunakan untuk menggambarkan struktur basis data sistem dalam bentuk '
      'entitas, atribut, dan relasi antar entitas. Pada penelitian ini, ERD '
      'dirancang untuk merepresentasikan kebutuhan penyimpanan data Sistem '
      'Dashboard Manajemen Proyek PT Smart Home Inovasi. Diagram ini berfungsi '
      'sebagai dasar perancangan basis data fisik (skema tabel) sekaligus '
      'menjadi acuan dalam memetakan alur data mulai dari pencatatan klien, '
      'pengelolaan proyek, penugasan teknisi, pelaporan tugas harian, hingga '
      'pengarsipan bukti pekerjaan dan pencatatan eskalasi kendala lapangan.')
    P(doc,
      'Berdasarkan rancangan tersebut, ERD sistem terdiri dari enam entitas '
      'utama, yaitu tb_klien, tb_proyek, tb_user, tb_tugas, tb_bukti, dan '
      'tb_eskalasi. Setiap entitas memiliki atribut serta keterhubungan dengan '
      'kardinalitas one-to-many (1:M) maupun many-to-many (M:N) yang mencerminkan '
      'aturan bisnis yang berlaku. Penjabaran lebih lanjut mengenai entitas, '
      'atribut, dan relasi yang membentuk ERD sistem dapat dilihat pada uraian '
      'berikut.')

    # ENTITAS
    H(doc, 'Entitas tb_klien')
    P(doc,
      'Entitas tb_klien merepresentasikan pihak pelanggan atau instansi yang '
      'menjadi pemilik proyek. Entitas ini menyimpan data identitas dan '
      'kontak klien yang dibutuhkan dalam pencatatan proyek baru maupun '
      'komunikasi selama proyek berlangsung. Atribut yang dimiliki entitas '
      'tb_klien diuraikan pada Tabel 4.1.')
    CAP(doc, 'Tabel 4.1 Atribut Entitas tb_klien')
    tabel_atribut(doc, [
        ('id_klien', 'integer', 'Primary key, identitas unik tiap klien.'),
        ('nama_klien', 'varchar', 'Nama klien atau instansi.'),
        ('alamat', 'text', 'Alamat lokasi klien atau lokasi proyek.'),
        ('no_telp', 'varchar', 'Nomor telepon kontak klien.'),
    ])

    H(doc, 'Entitas tb_proyek')
    P(doc,
      'Entitas tb_proyek menyimpan data utama setiap proyek yang dikerjakan '
      'oleh PT Smart Home Inovasi. Entitas ini mencatat informasi proyek '
      'mulai dari nama proyek, klien pemilik, status pengerjaan, fase '
      '(survei atau eksekusi), hingga nilai kontrak. Data pada entitas ini '
      'menjadi acuan dalam proses perhitungan Schedule Performance Index '
      '(SPI) pada dashboard Early Warning System (EWS). Atribut yang '
      'dimiliki entitas tb_proyek diuraikan pada Tabel 4.2.')
    CAP(doc, 'Tabel 4.2 Atribut Entitas tb_proyek')
    tabel_atribut(doc, [
        ('id_proyek', 'integer', 'Primary key, identitas unik tiap proyek.'),
        ('nama_proyek', 'varchar', 'Nama atau judul proyek.'),
        ('id_klien', 'integer', 'Foreign key ke entitas tb_klien.'),
        ('status', 'enum', 'Status proyek: active, completed, on-hold.'),
        ('phase', 'enum', 'Fase proyek: survey atau execution.'),
        ('project_value', 'decimal', 'Nilai kontrak proyek dalam Rupiah.'),
    ])

    H(doc, 'Entitas tb_user')
    P(doc,
      'Entitas tb_user merepresentasikan pengguna sistem, baik manajer, '
      'teknisi lapangan, maupun administrator. Entitas ini mengelola data '
      'autentikasi dan otorisasi pengguna serta menjadi dasar dalam '
      'pencatatan siapa yang mengerjakan tugas, mengelola proyek, atau '
      'melaporkan eskalasi. Atribut yang dimiliki entitas tb_user diuraikan '
      'pada Tabel 4.3.')
    CAP(doc, 'Tabel 4.3 Atribut Entitas tb_user')
    tabel_atribut(doc, [
        ('id_user', 'integer', 'Primary key, identitas unik tiap pengguna.'),
        ('name', 'varchar', 'Nama lengkap pengguna.'),
        ('email', 'varchar', 'Alamat email pengguna, digunakan untuk login.'),
        ('role', 'enum', 'Peran pengguna: technician, manager, atau admin.'),
    ])

    H(doc, 'Entitas tb_tugas')
    P(doc,
      'Entitas tb_tugas merepresentasikan unit pekerjaan yang dirinci dari '
      'sebuah proyek. Setiap tugas memiliki status pengerjaan, tenggat '
      'waktu, dan teknisi yang bertanggung jawab. Pencatatan tugas pada '
      'entitas ini menjadi sumber data utama dalam perhitungan SPI, '
      'penentuan status keterlambatan, serta penayangan dashboard EWS. '
      'Atribut yang dimiliki entitas tb_tugas diuraikan pada Tabel 4.4.')
    CAP(doc, 'Tabel 4.4 Atribut Entitas tb_tugas')
    tabel_atribut(doc, [
        ('id_tugas', 'integer', 'Primary key, identitas unik tiap tugas.'),
        ('id_proyek', 'integer', 'Foreign key ke entitas tb_proyek.'),
        ('status', 'enum', 'Status tugas: to_do, working_on_it, done.'),
        ('due_date', 'date', 'Tenggat penyelesaian tugas.'),
    ])

    H(doc, 'Entitas tb_bukti')
    P(doc,
      'Entitas tb_bukti digunakan untuk menyimpan dokumen atau lampiran '
      'bukti pekerjaan yang diunggah teknisi lapangan, seperti foto hasil '
      'instalasi atau dokumen pendukung. Entitas ini terkait langsung '
      'dengan tugas yang dilaporkan dan menjadi catatan otentik penyelesaian '
      'pekerjaan. Atribut yang dimiliki entitas tb_bukti diuraikan pada '
      'Tabel 4.5.')
    CAP(doc, 'Tabel 4.5 Atribut Entitas tb_bukti')
    tabel_atribut(doc, [
        ('id_bukti', 'integer', 'Primary key, identitas unik tiap bukti.'),
        ('id_tugas', 'integer', 'Foreign key ke entitas tb_tugas.'),
        ('file_path', 'varchar', 'Lokasi penyimpanan berkas bukti pada server.'),
    ])

    H(doc, 'Entitas tb_eskalasi')
    P(doc,
      'Entitas tb_eskalasi merepresentasikan pencatatan kendala atau '
      'permasalahan lapangan yang dilaporkan teknisi kepada manajer. '
      'Entitas ini berfungsi sebagai sarana komunikasi formal terkait '
      'hambatan pelaksanaan proyek, sekaligus menjadi basis pelacakan '
      'tindak lanjut hingga kendala dinyatakan selesai. Atribut yang '
      'dimiliki entitas tb_eskalasi diuraikan pada Tabel 4.6.')
    CAP(doc, 'Tabel 4.6 Atribut Entitas tb_eskalasi')
    tabel_atribut(doc, [
        ('id_eskalasi', 'integer', 'Primary key, identitas unik tiap eskalasi.'),
        ('id_proyek', 'integer', 'Foreign key ke entitas tb_proyek.'),
        ('priority', 'enum', 'Tingkat prioritas: low, medium, high.'),
        ('status', 'enum', 'Status eskalasi: open, handled, closed.'),
    ])

    # RELASI
    H(doc, 'Relasi Antar Entitas')
    P(doc,
      'Hubungan antar entitas dalam ERD sistem digambarkan dalam bentuk '
      'relasi dengan kardinalitas tertentu. Relasi tersebut mencerminkan '
      'aturan bisnis yang berlaku, misalnya satu klien dapat memiliki '
      'banyak proyek, satu proyek dapat berisi banyak tugas, dan satu '
      'tugas dapat memiliki banyak bukti pekerjaan. Penjelasan rinci '
      'mengenai relasi antar entitas diuraikan pada Tabel 4.7.')
    CAP(doc, 'Tabel 4.7 Relasi Antar Entitas')
    tabel_relasi(doc, [
        ('tb_klien', '1 : M', 'tb_proyek',
         'Satu klien dapat memiliki banyak proyek (memiliki).'),
        ('tb_user', 'M : M', 'tb_proyek',
         'Banyak pengguna dapat dialokasikan untuk mengelola banyak proyek (mengelola).'),
        ('tb_proyek', '1 : M', 'tb_tugas',
         'Satu proyek terdiri dari banyak tugas (memiliki).'),
        ('tb_user', '1 : M', 'tb_tugas',
         'Satu pengguna dapat mengerjakan banyak tugas (mengerjakan).'),
        ('tb_tugas', '1 : M', 'tb_bukti',
         'Satu tugas dapat memiliki banyak bukti pekerjaan (memiliki).'),
        ('tb_user', '1 : M', 'tb_eskalasi',
         'Satu pengguna dapat melaporkan banyak eskalasi kendala (melaporkan).'),
    ])

    doc.save(OUT_ERD)
    print(f'[OK] {OUT_ERD}')


# ============================================================================
# CLASS DIAGRAM CONTENT
# ============================================================================

def build_class():
    doc = Document(TEMPLATE)
    reset(doc)

    H(doc, 'Model Class Diagram')
    P(doc,
      'Class diagram merupakan salah satu instrumen pemodelan dalam UML yang '
      'berfungsi untuk mendeskripsikan struktur statis sebuah sistem dengan '
      'menjabarkan entitas kelas, atribut, operasi, serta keterhubungan '
      'antarkelas. Pada penelitian ini, class diagram dirancang untuk '
      'merepresentasikan rancangan logis Sistem Dashboard Manajemen Proyek '
      'PT Smart Home Inovasi sebelum tahap implementasi kode dilakukan. '
      'Diagram ini menjadi acuan dalam pembentukan model data, kontrak '
      'metode antar modul, serta pemetaan tanggung jawab masing-masing '
      'kelas pada arsitektur sistem.')
    P(doc,
      'Sebagai penjabaran lebih lanjut dari rancangan Class Diagram yang '
      'telah diuraikan sebelumnya, berikut adalah pemaparan mengenai fungsi '
      'spesifik dari masing-masing kelas yang membentuk sistem. Kelas-kelas '
      'tersebut dibagi berdasarkan tanggung jawab utamanya, mulai dari '
      'autentikasi pengguna, pengelolaan klien dan proyek, koordinasi '
      'penugasan, perhitungan kesehatan proyek (Early Warning System), '
      'pengelolaan tugas dan bukti, hingga pencatatan eskalasi kendala '
      'lapangan.')

    # ---- User ----
    H(doc, 'Kelas User')
    P(doc,
      'Kelas User (pengguna) berfungsi sebagai lapisan autentikasi dan '
      'otorisasi pada sistem. Kelas ini merepresentasikan entitas pengguna '
      'serta mengelola hak akses berdasarkan peran yang dimiliki, yaitu '
      'technician dan manager. Deskripsi fungsi pada kelas User dapat '
      'dilihat pada Tabel 4.1.')
    CAP(doc, 'Tabel 4.1 Kelas User')
    tabel_fungsi(doc, [
        ('login',
         'Memvalidasi kredensial berupa email dan kata sandi untuk '
         'memberikan akses masuk kepada pengguna.'),
        ('hasRole',
         'Memeriksa dan memastikan hak akses atau peran pengguna dalam '
         'sistem sebelum tindakan tertentu dijalankan.'),
    ])

    # ---- Klien ----
    H(doc, 'Kelas Klien')
    P(doc,
      'Kelas Klien merepresentasikan entitas pelanggan atau instansi '
      'pemilik proyek. Kelas ini digunakan untuk mengelola data klien yang '
      'terdaftar dalam sistem sehingga proyek selalu memiliki pihak '
      'penanggung jawab eksternal yang jelas. Deskripsi fungsi pada kelas '
      'Klien dapat dilihat pada Tabel 4.2.')
    CAP(doc, 'Tabel 4.2 Kelas Klien')
    tabel_fungsi(doc, [
        ('create',
         'Menambahkan dan menyimpan data profil klien baru ke dalam basis data.'),
        ('update',
         'Memperbarui informasi klien yang telah terdaftar dalam sistem.'),
        ('delete',
         'Menghapus data klien dari sistem.'),
    ])

    # ---- Proyek ----
    H(doc, 'Kelas Proyek')
    P(doc,
      'Kelas Proyek merupakan inti utama sistem karena berfungsi untuk '
      'mengelola koordinasi proyek, penjadwalan, fase pekerjaan, serta '
      'hubungan antar tugas. Kelas ini juga menjadi titik integrasi antara '
      'data klien, pengguna, tugas, hingga indikator kesehatan proyek. '
      'Deskripsi fungsi pada kelas Proyek dapat dilihat pada Tabel 4.3.')
    CAP(doc, 'Tabel 4.3 Kelas Proyek')
    tabel_fungsi(doc, [
        ('create',
         'Membuat catatan proyek baru beserta informasi awal seperti klien, '
         'jadwal, dan nilai kontrak.'),
        ('update',
         'Memperbarui informasi proyek, termasuk perubahan jadwal atau status.'),
        ('delete',
         'Menghapus proyek dari sistem.'),
        ('assignTechnician',
         'Menambahkan teknisi sebagai anggota tim proyek tertentu.'),
        ('approveSurvey',
         'Menyetujui hasil tahap survei agar proyek dapat berlanjut ke fase eksekusi.'),
        ('sortByUrgency',
         'Mengurutkan proyek berdasarkan tingkat urgensi dari status '
         'kesehatan (merah, kuning, hijau) pada dashboard EWS.'),
        ('getHealth',
         'Mengambil informasi indikator kesehatan proyek dari kelas KesehatanProyek.'),
    ])

    # ---- PenugasanProyek ----
    H(doc, 'Kelas PenugasanProyek')
    P(doc,
      'Kelas PenugasanProyek berfungsi sebagai kelas asosiasi yang '
      'menghubungkan pengguna dengan proyek tertentu secara terstruktur. '
      'Kelas ini mencatat siapa yang ditugaskan pada proyek tertentu dan '
      'kapan penugasan tersebut dilakukan. Deskripsi fungsi pada kelas '
      'PenugasanProyek dapat dilihat pada Tabel 4.4.')
    CAP(doc, 'Tabel 4.4 Kelas PenugasanProyek')
    tabel_fungsi(doc, [
        ('assign',
         'Mencatat penugasan seorang pengguna ke dalam sebuah proyek.'),
        ('unassign',
         'Menghapus catatan penugasan pengguna dari proyek tertentu.'),
        ('getByProject',
         'Mengambil daftar pengguna yang ditugaskan pada sebuah proyek.'),
    ])

    # ---- KesehatanProyek ----
    H(doc, 'Kelas KesehatanProyek')
    P(doc,
      'Kelas KesehatanProyek berfungsi untuk mengelola proses komputasi '
      'EWS serta menghitung indikator kesehatan proyek. Dengan adanya '
      'pemisahan ini, perhitungan SPI, Earned Value, dan Planned Value '
      'dijaga konsistensinya dan tidak bercampur dengan logika pengelolaan '
      'proyek itu sendiri. Deskripsi fungsi pada kelas KesehatanProyek '
      'dapat dilihat pada Tabel 4.5.')
    CAP(doc, 'Tabel 4.5 Kelas KesehatanProyek')
    tabel_fungsi(doc, [
        ('calculateSPI',
         'Menghitung Schedule Performance Index berdasarkan rasio '
         'penyelesaian tugas terhadap rentang waktu proyek.'),
        ('recalculate',
         'Memicu perhitungan ulang seluruh indikator kesehatan proyek setelah '
         'terjadi perubahan status tugas.'),
        ('getStatus',
         'Mengembalikan status kesehatan proyek dalam kategori green, amber, '
         'atau red sesuai ambang SPI.'),
        ('computeEarnedValue',
         'Menghitung nilai pekerjaan yang telah diselesaikan (Earned Value).'),
        ('computePlannedValue',
         'Menghitung nilai pekerjaan yang seharusnya tercapai pada saat '
         'pengukuran (Planned Value).'),
    ])

    # ---- Tugas ----
    H(doc, 'Kelas Tugas')
    P(doc,
      'Kelas Tugas merepresentasikan unit pekerjaan atau tugas lapangan '
      'yang harus diselesaikan dalam periode waktu tertentu. Kelas ini '
      'menjadi titik utama interaksi teknisi dengan sistem sekaligus sumber '
      'utama data perhitungan SPI. Deskripsi fungsi pada kelas Tugas dapat '
      'dilihat pada Tabel 4.6.')
    CAP(doc, 'Tabel 4.6 Kelas Tugas')
    tabel_fungsi(doc, [
        ('create',
         'Menambahkan tugas baru ke dalam sebuah proyek.'),
        ('update',
         'Memperbarui informasi tugas, seperti deskripsi, tenggat waktu, '
         'atau penanggung jawab.'),
        ('delete',
         'Menghapus tugas dari proyek.'),
        ('changeStatus',
         'Mengubah status tugas oleh teknisi (to_do atau working_on_it).'),
        ('markDone',
         'Menandai tugas sebagai selesai (done). Operasi ini hanya dapat '
         'dilakukan oleh manajer sebagai gerbang review.'),
        ('isOvertime',
         'Memeriksa apakah tugas berstatus working_on_it telah melewati '
         'tenggat waktu.'),
        ('isOverDeadline',
         'Memeriksa apakah tugas berstatus to_do telah melewati tenggat waktu.'),
    ])

    # ---- BuktiTugas ----
    H(doc, 'Kelas BuktiTugas')
    P(doc,
      'Kelas BuktiTugas digunakan untuk mengelola dokumen atau lampiran '
      'bukti pekerjaan yang dikirimkan oleh teknisi. Kelas ini menjamin '
      'bahwa setiap tugas yang dilaporkan dapat diverifikasi melalui '
      'dokumentasi pendukung. Deskripsi fungsi pada kelas BuktiTugas dapat '
      'dilihat pada Tabel 4.7.')
    CAP(doc, 'Tabel 4.7 Kelas BuktiTugas')
    tabel_fungsi(doc, [
        ('upload',
         'Mengunggah berkas bukti pekerjaan ke server dan menautkannya pada '
         'sebuah tugas.'),
        ('download',
         'Mengunduh berkas bukti yang telah tersimpan.'),
        ('delete',
         'Menghapus berkas bukti dari sistem.'),
    ])

    # ---- Eskalasi ----
    H(doc, 'Kelas Eskalasi')
    P(doc,
      'Kelas Eskalasi berfungsi untuk mencatat, mengelola prioritas, dan '
      'memantau proses komunikasi terkait kendala lapangan yang terjadi '
      'selama pelaksanaan proyek. Kelas ini memastikan setiap hambatan '
      'tercatat secara formal dan dapat ditindaklanjuti oleh manajer. '
      'Deskripsi fungsi pada kelas Eskalasi dapat dilihat pada Tabel 4.8.')
    CAP(doc, 'Tabel 4.8 Kelas Eskalasi')
    tabel_fungsi(doc, [
        ('create',
         'Membuat catatan eskalasi baru atas kendala yang dialami teknisi.'),
        ('update',
         'Memperbarui informasi eskalasi, termasuk prioritas atau deskripsi.'),
        ('resolve',
         'Menandai eskalasi sebagai selesai setelah kendala teratasi.'),
        ('sendInstruction',
         'Mengirimkan instruksi tindak lanjut dari manajer kepada teknisi '
         'terkait penanganan kendala.'),
    ])

    # ---- Relasi ----
    H(doc, 'Relasi Antar Kelas')
    P(doc,
      'Keterhubungan antar kelas dalam class diagram menggambarkan '
      'bagaimana objek-objek pada sistem saling berinteraksi pada saat '
      'aplikasi dijalankan. Relasi tersebut juga merefleksikan aturan '
      'bisnis yang berlaku, seperti satu klien dapat memiliki banyak '
      'proyek, dan satu proyek dapat memiliki banyak tugas yang dikerjakan '
      'oleh teknisi yang berbeda. Penjelasan rinci mengenai relasi antar '
      'kelas dapat dilihat pada Tabel 4.9.')
    CAP(doc, 'Tabel 4.9 Relasi Antar Kelas')
    tabel_relasi(doc, [
        ('User', '1 : M', 'PenugasanProyek',
         'Seorang pengguna dapat menerima banyak penugasan proyek.'),
        ('Klien', '1 : M', 'Proyek',
         'Satu klien dapat memiliki banyak proyek.'),
        ('Proyek', 'M : 1', 'PenugasanProyek',
         'Satu proyek dapat memiliki banyak penugasan pengguna (relasi melalui kelas asosiasi).'),
        ('Proyek', '1 : 1', 'KesehatanProyek',
         'Setiap proyek memiliki tepat satu catatan kesehatan proyek.'),
        ('Proyek', '1 : M', 'Tugas',
         'Satu proyek dapat memiliki banyak tugas.'),
        ('Tugas', '1 : M', 'BuktiTugas',
         'Satu tugas dapat memiliki banyak berkas bukti pekerjaan.'),
        ('Tugas', '1 : M', 'Eskalasi',
         'Satu tugas dapat memiliki banyak catatan eskalasi.'),
        ('User', '1 : M', 'Tugas',
         'Seorang pengguna (teknisi) dapat ditugaskan untuk mengerjakan banyak tugas (assignee).'),
    ])

    doc.save(OUT_CLASS)
    print(f'[OK] {OUT_CLASS}')


if __name__ == '__main__':
    build_erd()
    build_class()
