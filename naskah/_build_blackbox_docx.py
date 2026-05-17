"""Build STANDALONE docx for sub-bab 5.2.3 Pengujian (Black Box).

This is a separate deliverable, NOT a modification to Naskah TA 17-05-26.docx.
Content matches Dian's narrative style and uses TNR 12 / 1.5 spacing / justify
to align with typical Indonesian academic thesis formatting.

Output: naskah/5.2.3 Pengujian Black Box.docx
"""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "naskah/5.2.3 Pengujian Black Box.docx"

# Body content -----------------------------------------------------------------

HEADING = "5.2.3 Pengujian"

INTRO_P1 = (
    "Setelah seluruh antarmuka pengguna dan logika komputasi pada sistem "
    "diimplementasikan, tahapan selanjutnya adalah pelaksanaan pengujian sistem. "
    "Pengujian ini bertujuan untuk memvalidasi bahwa seluruh fungsionalitas yang "
    "telah dibangun berjalan sesuai dengan spesifikasi kebutuhan operasional yang "
    "didefinisikan pada Bab IV, serta memastikan sistem terbebas dari kesalahan "
    "pemrosesan (*bug*) sebelum diterapkan di lingkungan PT Smart Home Inovasi "
    "Yogyakarta."
)

INTRO_P2 = (
    "Metode yang digunakan pada tahap ini adalah pengujian *Black Box* (Kotak "
    "Hitam). Pendekatan ini dipilih karena evaluasi difokuskan pada fungsionalitas "
    "*input* dan *output* perangkat lunak—seperti akurasi perekaman *daily report*, "
    "ketepatan kalkulasi *Schedule Performance Index* (*SPI*), dan respons pemicu "
    "indikator *Early Warning System* (*EWS*) pada *dashboard*—tanpa perlu meninjau "
    "struktur internal dari kode program. Pengujian dilakukan berdasarkan skenario "
    "interaksi langsung antara aktor (Manajer dan Teknisi) dengan sistem."
)

CAPTION = "Tabel 5.1 Skenario dan Hasil Pengujian Black Box"

HEADERS = ["No", "ID Pengujian", "Deskripsi Skenario",
           "Hasil Diharapkan", "Hasil Aktual", "Status"]

ROWS = [
    ("TC-AUTH-01",
     "Manajer melakukan *login* ke dalam sistem menggunakan *kredensial* otorisasi yang valid.",
     "Sistem berhasil melakukan autentikasi *kredensial* dan merender *dashboard* utama dengan *Role-Based Access Control* (*RBAC*) tingkat Manajer secara presisi.",
     "Autentikasi berhasil dan *dashboard* Manajer dirender dengan tepat.",
     "Valid"),
    ("TC-AUTH-02",
     "Teknisi melakukan *login* ke dalam sistem operasional menggunakan *kredensial* yang terdaftar.",
     "Sistem memproses autentikasi secara valid dan mengarahkan pengguna ke halaman *kanban* khusus Teknisi sesuai konfigurasi *role*.",
     "Autentikasi Teknisi berhasil diproses dan halaman *kanban* ditampilkan.",
     "Valid"),
    ("TC-AUTH-03",
     "Teknisi menginisiasi upaya untuk mengakses halaman *dashboard* Manajer melalui manipulasi URL langsung.",
     "Sistem mengeksekusi validasi *Role-Based Access Control* (*RBAC*) dan memblokir akses tersebut sebagai bentuk mitigasi otorisasi.",
     "Validasi *RBAC* berhasil mencegah akses ilegal dari Teknisi.",
     "Valid"),
    ("TC-MN-01",
     "Manajer memasukkan entitas data proyek instalasi baru ke dalam matriks sistem perencanaan.",
     "Sistem melakukan perekaman data entitas secara presisi ke dalam basis data dan merender proyek baru pada daftar proyek.",
     "Perekaman entitas proyek baru pada basis data tervalidasi sukses.",
     "Valid"),
    ("TC-MN-02",
     "Manajer menugaskan Teknisi ke sebuah proyek baru dengan menginputkan rentang jadwal pelaksanaan.",
     "Sistem memvalidasi ketersediaan jadwal Teknisi terhadap matriks alokasi sebelumnya dan menyimpan penugasan jika tidak terjadi bentrok jadwal.",
     "Validasi jadwal berhasil dan penugasan Teknisi direkam ke sistem.",
     "Valid"),
    ("TC-MN-03",
     "Manajer mengeksekusi mekanisme *review gate* dengan meninjau *task evidence* dari Teknisi dan memperbarui *status* menjadi Done.",
     "Sistem merekam persetujuan Manajer secara presisi dan mengubah *status* tugas dari sedang ditinjau menjadi Done di basis data.",
     "Perekaman mekanisme *review gate* dan pembaruan *status* berhasil dieksekusi.",
     "Valid"),
    ("TC-MN-04",
     "Manajer menginputkan instruksi balasan sebagai mitigasi atas laporan eskalasi kendala yang dikirimkan oleh Teknisi.",
     "Sistem memproses instruksi balasan tersebut dan meneruskannya kembali ke halaman antarmuka Teknisi yang bersangkutan secara real-time.",
     "Pemrosesan balasan eskalasi Manajer berhasil diteruskan kepada Teknisi.",
     "Valid"),
    ("TC-TK-01",
     "Teknisi melakukan pembaruan *status* pada papan *kanban* operasional dengan memindahkan tugas dari to_do menuju working_on_it.",
     "Sistem merender perpindahan *status* secara dinamis pada antarmuka *kanban* dan merekam perubahan tersebut untuk pemantauan Manajer.",
     "Pembaruan *status* pada *kanban* terekam secara presisi.",
     "Valid"),
    ("TC-TK-02",
     "Teknisi mengunggah dokumentasi visual sebagai *task evidence* untuk tugas operasional lapangan yang telah diselesaikan.",
     "Sistem memvalidasi format dokumen dan menyimpan *task evidence* secara presisi agar dapat ditinjau melalui *review gate*.",
     "Penyimpanan *task evidence* dari Teknisi ke sistem tervalidasi sukses.",
     "Valid"),
    ("TC-TK-03",
     "Teknisi mengirimkan formulir laporan eskalasi terkait kendala teknis atau *bug* yang ditemukan di lapangan kerja.",
     "Sistem merekam agregasi kendala lapangan dan secara otomatis menampilkan indikator peringatan eskalasi pada *dashboard* Manajer.",
     "Perekaman laporan eskalasi dan notifikasi ke Manajer berhasil.",
     "Valid"),
    ("TC-TK-04",
     "Teknisi mengakses menu Self-Performance Dashboard sebagai bentuk swamonitoring atas capaian *Schedule Performance Index* (*SPI*) pribadinya.",
     "Sistem merender matriks kinerja individu dengan menyajikan kalkulasi *SPI* berdasarkan agregasi penugasan yang ditangani oleh Teknisi bersangkutan.",
     "Matriks swamonitoring kinerja Teknisi berhasil dirender secara presisi.",
     "Valid"),
    ("TC-SYS-01",
     "Sistem merespons aksi persetujuan *review gate* dari Manajer yang merubah *status* penugasan menjadi Done.",
     "Sistem mengomputasi nilai *Earned Value* secara otomatis berdasarkan bobot penugasan yang telah disetujui penyelesaiannya oleh Manajer.",
     "Komputasi *Earned Value* oleh sistem berjalan secara akurat.",
     "Valid"),
    ("TC-SYS-02",
     "Sistem mengeksekusi mekanisme *Early Warning System* (*EWS*) pada saat periode pelaporan harian (*daily report*) berakhir.",
     "Sistem mengomputasi *Schedule Performance Index* (*SPI*) dengan membandingkan *Earned Value* terhadap *Planned Value*, lalu merender indikator warna (Red-Amber-Green) secara presisi.",
     "Komputasi *SPI* dan penentuan indikator warna berfungsi tepat.",
     "Valid"),
    ("TC-SYS-03",
     "Sistem mengelola representasi daftar proyek pada halaman *dashboard* pemantauan Manajer.",
     "Sistem melakukan agregasi nilai *SPI* dan merender urutan proyek dengan mengutamakan entitas proyek berstatus kritis pada posisi teratas matriks.",
     "Agregasi prioritas proyek pada *dashboard* dirender dengan valid.",
     "Valid"),
    ("TC-SYS-04",
     "Sistem memproses input data alokasi penugasan Teknisi baru dari form entri yang diakses oleh Manajer.",
     "Sistem mengomputasi validasi jadwal untuk mencegah presensi jadwal yang bentrok antar proyek pada matriks penugasan Teknisi.",
     "Validasi sistem untuk mitigasi bentrok jadwal berfungsi sempurna.",
     "Valid"),
]

CLOSING = (
    "Berdasarkan hasil pengujian *Black Box* pada Tabel 5.1, keseluruhan 15 "
    "skenario yang terbagi ke dalam empat kategori utama—yakni otentikasi "
    "(AUTH), fungsionalitas manajer (MN), fungsionalitas teknisi (TK), dan "
    "pemrosesan sistem (SYS)—menunjukkan tingkat keberhasilan 100% dengan "
    "seluruh pengujian berstatus Valid. Pencapaian ini secara langsung "
    "memvalidasi bahwa seluruh kebutuhan fungsional yang didefinisikan pada "
    "Bab IV telah terpenuhi secara presisi di dalam satu ekosistem sistem yang "
    "terpadu. Lebih lanjut, hal ini membuktikan keandalan sistem dalam "
    "menerapkan otorisasi *Role-Based Access Control* (*RBAC*) serta "
    "mengomputasi kalkulasi *Schedule Performance Index* (*SPI*) dari agregasi "
    "input *daily report*. Proses komputasi tersebut juga telah tervalidasi "
    "mampu memicu indikator *Early Warning System* (*EWS*) pada *dashboard* "
    "secara otomatis sebagai langkah mitigasi risiko. Dengan demikian, "
    "perangkat lunak ini dinyatakan stabil dan siap untuk diimplementasikan di "
    "lingkungan operasional PT Smart Home Inovasi Yogyakarta. Setelah seluruh "
    "kapabilitas teknis dasar ini berhasil melewati *review gate* pengujian, "
    "pembahasan selanjutnya akan mengevaluasi keunggulan serta inovasi yang "
    "ditawarkan oleh sistem."
)


# Style helpers ---------------------------------------------------------------

def add_runs_with_italic(paragraph, text, font_name="Times New Roman", font_size_pt=12):
    parts = re.split(r"(\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        italic = part.startswith("*") and part.endswith("*") and len(part) > 2
        run = paragraph.add_run(part[1:-1] if italic else part)
        run.italic = italic
        run.font.name = font_name
        run.font.size = Pt(font_size_pt)


def set_paragraph_format(p, *, justify=True, line_spacing=1.5,
                         first_line_indent_cm=None, space_after=0,
                         space_before=0):
    pf = p.paragraph_format
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = line_spacing
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if first_line_indent_cm is not None:
        pf.first_line_indent = Cm(first_line_indent_cm)
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)


def add_heading(doc, text, level=3):
    """Sub-sub-heading (5.2.3). Use heading style + TNR 12 bold."""
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    set_paragraph_format(p, justify=False, line_spacing=1.5, space_before=12, space_after=6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_intro(doc, text):
    p = doc.add_paragraph()
    set_paragraph_format(p, justify=True, line_spacing=1.5,
                         first_line_indent_cm=1.27, space_after=0)
    add_runs_with_italic(p, text)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = 1.0
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    return p


def set_cell_runs(cell, text, *, bold=False, font_size_pt=10):
    """Replace cell text with italic-aware runs."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = 1.0
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    parts = re.split(r"(\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        italic = part.startswith("*") and part.endswith("*") and len(part) > 2
        run = p.add_run(part[1:-1] if italic else part)
        run.italic = italic
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(font_size_pt)


# Build the document ----------------------------------------------------------

def main():
    doc = Document()

    # Set default style to Times New Roman 12
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    add_heading(doc, HEADING, level=3)
    add_intro(doc, INTRO_P1)
    add_intro(doc, INTRO_P2)
    add_caption(doc, CAPTION)

    # Table: 1 header + 15 data rows, 6 columns
    table = doc.add_table(rows=1 + len(ROWS), cols=6)
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Header row
    for i, h in enumerate(HEADERS):
        set_cell_runs(table.rows[0].cells[i], h, bold=True, font_size_pt=10)
        # center align header
        table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for ri, row in enumerate(ROWS, start=1):
        cells = table.rows[ri].cells
        set_cell_runs(cells[0], str(ri), font_size_pt=10)
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_runs(cells[1], row[0], font_size_pt=10)
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_runs(cells[2], row[1], font_size_pt=10)
        set_cell_runs(cells[3], row[2], font_size_pt=10)
        set_cell_runs(cells[4], row[3], font_size_pt=10)
        set_cell_runs(cells[5], row[4], font_size_pt=10)
        cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Closing paragraph
    p_close = doc.add_paragraph()
    set_paragraph_format(p_close, justify=True, line_spacing=1.5,
                         first_line_indent_cm=1.27, space_before=12)
    add_runs_with_italic(p_close, CLOSING)

    doc.save(OUT)
    print(f"[OK] Built standalone testing docx: {OUT}")


if __name__ == "__main__":
    main()
