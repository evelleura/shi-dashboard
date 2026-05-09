# /// script
# requires-python = ">=3.11"
# dependencies = ["pymupdf>=1.24", "python-docx>=1.1.2", "Pillow>=10"]
# ///
"""Extract BAB V (Implementasi & Hasil) from naskah PDF -> .docx.
Pages 86-146 inclusive. Embeds images at their Gambar captions.
Output: Times New Roman, justified body, structured headings."""
from __future__ import annotations
import io
import os
import re
import fitz  # pymupdf
from PIL import Image
from docx import Document
from docx.shared import Pt, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

PDF_PATH = r"D:\__CODING\05-MyProjects\__IRENE\shi-crm\naskah\5220411315_RestyanditoVegaManuhara_BAB 6.pdf"
OUT_PATH = r"D:\__CODING\05-MyProjects\__IRENE\shi-crm\naskah\BAB_V_Implementasi.docx"
IMG_DIR = r"D:\__CODING\05-MyProjects\__IRENE\shi-crm\naskah\_bab5_images"

PAGE_START = 86  # 1-indexed inclusive
PAGE_END = 146   # 1-indexed inclusive

FONT = "Times New Roman"

HEADING_LEVEL_RE = re.compile(r"^(\d)\.(\d)(?:\.(\d))?\.?\s+(.+)$")
GAMBAR_RE = re.compile(r"^Gambar\s+5\.(\d+)\s+(.+)$", re.IGNORECASE)
TABEL_RE = re.compile(r"^Tabel\s+5\.(\d+)\s+(.+)$", re.IGNORECASE)
BAB_HEADING_RE = re.compile(r"^BAB\s+V\b")

# Common PDF unicode mojibake fixes
UNICODE_FIXES = [
    ("�", "-"),      # replacement char -> dash
    (" ", " "),      # non-breaking space
    ("–", "-"),      # en-dash -> hyphen
    ("—", "-"),      # em-dash -> hyphen
    ("‘", "'"),      # left single quote
    ("’", "'"),      # right single quote
    ("“", '"'),      # left double quote
    ("”", '"'),      # right double quote
    ("…", "..."),    # ellipsis
    ("­", ""),       # soft hyphen
]


def fix_unicode(s: str) -> str:
    for bad, good in UNICODE_FIXES:
        s = s.replace(bad, good)
    return s


def style_run(run, *, bold: bool = False, italic: bool = False, size: int = 12,
              mono: bool = False) -> None:
    font_name = "Consolas" if mono else FONT
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), font_name)


def add_para(doc, text: str, *, bold: bool = False, italic: bool = False,
             size: int = 12, justify: bool = False, center: bool = False,
             indent_first: float = 0.0, indent_left: float = 0.0,
             space_before: int = 0, space_after: int = 6, mono: bool = False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    if indent_first:
        pf.first_line_indent = Cm(indent_first)
    if indent_left:
        pf.left_indent = Cm(indent_left)
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        style_run(run, bold=bold, italic=italic, size=size, mono=mono)
    return p


def add_image(doc, img_bytes: bytes, max_width_cm: float = 14.0) -> None:
    """Center-align embedded image, scale to max width."""
    img = Image.open(io.BytesIO(img_bytes))
    w_px, h_px = img.size
    # Convert at 96 DPI
    w_cm = w_px / 96 * 2.54
    h_cm = h_px / 96 * 2.54
    if w_cm > max_width_cm:
        scale = max_width_cm / w_cm
        w_cm = max_width_cm
        h_cm = h_cm * scale
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    bio = io.BytesIO(img_bytes)
    run.add_picture(bio, width=Cm(w_cm))


def add_caption(doc, text: str) -> None:
    """Italic centered figure/table caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    style_run(run, italic=True, size=11)


def add_heading(doc, level: int, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    size = {0: 16, 1: 14, 2: 13, 3: 12}.get(level, 12)
    style_run(run, bold=True, size=size)


def is_heading_5_x(text: str) -> tuple[int, str] | None:
    """Return (level, text) if line is a 5.x heading."""
    m = HEADING_LEVEL_RE.match(text.strip())
    if not m:
        return None
    chap, sec, sub, rest = m.group(1), m.group(2), m.group(3), m.group(4)
    if chap != "5":
        return None
    if sub:
        return (3, f"5.{sec}.{sub} {rest}")
    return (2, f"5.{sec} {rest}")


def is_code_line(line: str) -> bool:
    """Heuristic: code lines have specific patterns."""
    s = line.strip()
    if not s:
        return False
    code_markers = (
        "Future<", "void ", "Widget ", "class ", "import ", "final ",
        "const ", "return ", "if (", "} else", "});", "=> ",
        "@override", "Navigator.", "setState(", "FirebaseAuth",
        "FirebaseFirestore", "context,", "MaterialPageRoute",
    )
    return any(m in s for m in code_markers)


def extract_page_images(page: fitz.Page) -> list[tuple[float, bytes]]:
    """Return [(y_top, image_bytes), ...] sorted by y-coord."""
    images = []
    for xref_info in page.get_images(full=True):
        xref = xref_info[0]
        try:
            img_dict = page.parent.extract_image(xref)
            img_bytes = img_dict["image"]
            # Get bbox via page.get_image_rects()
            rects = page.get_image_rects(xref)
            if rects:
                y_top = rects[0].y0
            else:
                y_top = 0
            images.append((y_top, img_bytes, img_dict.get("ext", "png")))
        except Exception:
            continue
    return sorted(images, key=lambda x: x[0])


def extract_pages(doc_pdf, page_start: int, page_end: int):
    """Yield (page_num, text_blocks, images) for each page in range."""
    for pn in range(page_start - 1, page_end):
        if pn >= doc_pdf.page_count:
            break
        page = doc_pdf[pn]
        # Use blocks for ordered reading
        blocks = page.get_text("blocks")
        # blocks: (x0, y0, x1, y1, text, block_no, block_type)
        # block_type 0 = text, 1 = image
        sorted_blocks = sorted(blocks, key=lambda b: (round(b[1]), b[0]))
        # Get images with y-coord
        page_imgs = extract_page_images(page)
        yield pn + 1, sorted_blocks, page_imgs


def main() -> None:
    os.makedirs(IMG_DIR, exist_ok=True)
    pdf = fitz.open(PDF_PATH)
    print(f"PDF pages: {pdf.page_count}")

    doc = Document()

    # Set default style font
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT)

    # Page margins
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)  # academic Indo style
    section.right_margin = Cm(2.5)

    # Title
    add_heading(doc, 0, "BAB V")
    add_heading(doc, 0, "IMPLEMENTASI DAN HASIL SERTA PEMBAHASAN")

    figures_seen = 0
    tables_seen = 0
    embedded_images = []  # track to emit at captions

    # Pre-extract all images per page
    pages_data = list(extract_pages(pdf, PAGE_START, PAGE_END))

    # Pending images per page (consumed when matched to caption)
    page_image_queue: dict[int, list[tuple[float, bytes, str]]] = {
        pn: list(imgs) for pn, _, imgs in pages_data
    }

    for pn, blocks, page_imgs in pages_data:
        for b in blocks:
            if len(b) < 6 or b[6] != 0:
                continue  # not text block
            text = fix_unicode(b[4]).strip()
            if not text:
                continue
            # Skip page numbers / running headers
            if re.fullmatch(r"\d{1,3}", text):
                continue
            # BAB V title page - skip if found again
            if BAB_HEADING_RE.match(text) and len(text) < 60:
                continue
            if "IMPLEMENTASI DAN HASIL" in text and len(text) < 100:
                continue

            # Process line-by-line within block
            lines = text.split("\n")
            buffer = []

            def flush_buffer():
                if not buffer:
                    return
                joined = " ".join(line.strip() for line in buffer if line.strip())
                if not joined:
                    buffer.clear()
                    return
                # Decide style: code vs prose
                if any(is_code_line(line) for line in buffer):
                    add_para(doc, joined, mono=True, size=10, space_after=4)
                else:
                    add_para(doc, joined, justify=True, indent_first=1.27)
                buffer.clear()

            for line in lines:
                stripped = line.strip()
                if not stripped:
                    flush_buffer()
                    continue

                # Heading 5.x
                head = is_heading_5_x(stripped)
                if head:
                    flush_buffer()
                    level, txt = head
                    add_heading(doc, level, txt)
                    continue

                # Gambar caption
                gm = GAMBAR_RE.match(stripped)
                if gm:
                    flush_buffer()
                    fig_num = int(gm.group(1))
                    fig_title = gm.group(2)
                    # Pop next image from this or earlier pages
                    img_data = None
                    for search_pn in range(pn, pn + 1):
                        q = page_image_queue.get(search_pn, [])
                        if q:
                            img_data = q.pop(0)
                            break
                    # If no image on current page, look at previous page's leftover
                    if img_data is None:
                        for search_pn in range(max(PAGE_START, pn - 1), pn):
                            q = page_image_queue.get(search_pn, [])
                            if q:
                                img_data = q.pop(0)
                                break
                    if img_data:
                        _, img_bytes, _ext = img_data
                        try:
                            add_image(doc, img_bytes)
                        except Exception as e:
                            print(f"  img embed fail page {pn} fig {fig_num}: {e}")
                    add_caption(doc, f"Gambar 5.{fig_num} {fig_title}")
                    figures_seen += 1
                    continue

                # Tabel caption
                tm = TABEL_RE.match(stripped)
                if tm:
                    flush_buffer()
                    add_caption(doc, f"Tabel 5.{tm.group(1)} {tm.group(2)}")
                    tables_seen += 1
                    continue

                buffer.append(stripped)

            flush_buffer()

    # Save
    doc.save(OUT_PATH)
    print(f"OK {OUT_PATH}")
    print(f"  Figures: {figures_seen} | Tables: {tables_seen}")
    print(f"  Pages processed: {len(pages_data)}")


if __name__ == "__main__":
    main()
